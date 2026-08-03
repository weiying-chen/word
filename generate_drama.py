#!/usr/bin/env python3

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from docx_utils import set_run_font_family, set_style_font_family
from style_tokens import (
    BODY_TEXT_SIZE_PT,
    DEFAULT_DOCX_ASCII_FONT_NAME,
    DEFAULT_DOCX_EAST_ASIA_FONT_NAME,
)
from template_styles import ensure_base_styles


SCHEMA_VERSION = "2.0"
DOCUMENT_TYPE = "script_translation_working_file"
SUPPORTED_ELEMENT_TYPES = {
    "action",
    "dialogue",
    "production_note",
    "translation_note",
}
INFORMATIONAL_REVIEW_FLAGS = {"merged_visual_continuation"}
SOURCE_BOOKMARK_PREFIX = "src_"
RAW_PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


class DramaValidationError(ValueError):
    pass


@dataclass(frozen=True)
class DramaFormat:
    page_width_inches: float = 8.5
    page_height_inches: float = 11
    top_margin_inches: float = 1
    bottom_margin_inches: float = 1
    left_margin_inches: float = 1.25
    right_margin_inches: float = 1.25
    header_distance_inches: float = 0.5
    footer_distance_inches: float = 0.5
    font_name: str = DEFAULT_DOCX_ASCII_FONT_NAME
    east_asia_font_name: str = DEFAULT_DOCX_EAST_ASIA_FONT_NAME
    body_size_pt: float = BODY_TEXT_SIZE_PT
    heading_size_pt: float = BODY_TEXT_SIZE_PT
    title_size_pt: float = BODY_TEXT_SIZE_PT
    space_before_pt: float | None = None
    space_after_pt: float | None = None
    line_spacing: float | None = None


@dataclass(frozen=True)
class GenerationResult:
    output_path: Path
    generated_ids: list[str]
    missing_fields: list[str]

    @property
    def missing_count(self) -> int:
        return len(self.missing_fields)


def load_payload(path: Path) -> dict:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise DramaValidationError(f"Unable to read drama JSON: {exc}") from exc
    if not isinstance(payload, dict):
        raise DramaValidationError("Drama JSON must be a top-level object.")
    return payload


def resolve_output_path(
    input_path: Path,
    output_path: Path | None,
    *,
    preview: bool,
) -> Path:
    if output_path is not None:
        return output_path
    base_name = input_path.stem
    if base_name.endswith(".translation"):
        base_name = base_name[: -len(".translation")]
    preview_marker = ".preview" if preview else ""
    return input_path.with_name(f"{base_name}_英文{preview_marker}.docx")


def _nonempty(value: object) -> bool:
    return isinstance(value, str) and bool(value.strip())


def _all_records(payload: dict) -> list[tuple[str, int, dict]]:
    records: list[tuple[str, int, dict]] = []
    for item in payload.get("front_matter", []):
        if isinstance(item, dict):
            records.append((str(item.get("id") or ""), item.get("order"), item))
    for scene in payload.get("scenes", []):
        if not isinstance(scene, dict):
            continue
        heading_record = {
            "type": "scene_heading",
            "translation_en": scene.get("heading_en"),
            "review": scene.get("review"),
        }
        records.append(
            (
                str(scene.get("heading_id") or ""),
                scene.get("heading_order"),
                heading_record,
            )
        )
        for item in scene.get("elements", []):
            if isinstance(item, dict):
                records.append((str(item.get("id") or ""), item.get("order"), item))
    return records


def _review_flags(record: dict) -> list[str]:
    review = record.get("review")
    if not isinstance(review, dict):
        return []
    flags = review.get("flags")
    if not isinstance(flags, list):
        return []
    return [str(flag).strip() for flag in flags if str(flag).strip()]


def _blocking_review_flags(record: dict) -> list[str]:
    return [
        flag
        for flag in _review_flags(record)
        if flag not in INFORMATIONAL_REVIEW_FLAGS
    ]


def collect_missing_english(payload: dict) -> list[str]:
    missing: list[str] = []
    for index, item in enumerate(payload.get("front_matter", []), start=1):
        if not _nonempty(item.get("translation_en")):
            missing.append(f"front_matter[{index}].translation_en")
    for scene in payload.get("scenes", []):
        number = scene.get("scene_number")
        if not _nonempty(scene.get("heading_en")):
            missing.append(f"scene {number}.heading_en")
        for item in scene.get("elements", []):
            record_id = item.get("id") or "missing-id"
            if not _nonempty(item.get("translation_en")):
                missing.append(f"{record_id}.translation_en")
            if item.get("type") != "dialogue":
                continue
            if _nonempty(item.get("speaker_zh")) and not _nonempty(
                item.get("speaker_en")
            ):
                missing.append(f"{record_id}.speaker_en")
            if _nonempty(item.get("parenthetical_zh")) and not _nonempty(
                item.get("parenthetical_en")
            ):
                missing.append(f"{record_id}.parenthetical_en")
    return missing


def validate_payload(payload: dict, *, final: bool) -> list[str]:
    errors: list[str] = []
    if str(payload.get("schema_version")) != SCHEMA_VERSION:
        errors.append(f"Input must use Schema Version {SCHEMA_VERSION}.")
    if payload.get("document_type") != DOCUMENT_TYPE:
        errors.append(f"document_type must be {DOCUMENT_TYPE!r}.")

    front_matter = payload.get("front_matter")
    scenes = payload.get("scenes")
    statistics = payload.get("statistics")
    if not isinstance(front_matter, list):
        errors.append("front_matter must be an array.")
        front_matter = []
    if not isinstance(scenes, list):
        errors.append("scenes must be an array.")
        scenes = []
    if not isinstance(statistics, dict):
        errors.append("statistics must be an object.")
        statistics = {}

    scene_numbers = [scene.get("scene_number") for scene in scenes if isinstance(scene, dict)]
    expected_scene_numbers = list(range(1, len(scenes) + 1))
    if scene_numbers != expected_scene_numbers:
        errors.append(
            "Invalid scene numbering: expected consecutive scenes "
            f"{expected_scene_numbers[:1]} through {expected_scene_numbers[-1:]}, "
            f"got {scene_numbers}."
        )
    if statistics.get("scene_count") != len(scenes):
        errors.append("statistics.scene_count does not match the scenes array.")
    if scenes:
        if statistics.get("first_scene") != 1:
            errors.append("statistics.first_scene must be 1.")
        if statistics.get("last_scene") != len(scenes):
            errors.append("statistics.last_scene does not match scene numbering.")
    if statistics.get("front_matter_count") != len(front_matter):
        errors.append("statistics.front_matter_count does not match front_matter.")
    for scene in scenes:
        if not isinstance(scene, dict) or not _nonempty(scene.get("heading_en")):
            continue
        number = scene.get("scene_number")
        match = re.match(r"^SCENE\s+(\d+)\.", scene["heading_en"].strip(), re.IGNORECASE)
        if match is None or int(match.group(1)) != number:
            errors.append(
                f"Scene {number} heading scene number must begin with "
                f"'SCENE {number}.'."
            )

    records = _all_records(payload)
    ids = [record_id for record_id, _, _ in records]
    missing_ids = [index + 1 for index, value in enumerate(ids) if not value]
    if missing_ids:
        errors.append(f"Records have absent IDs at positions {missing_ids}.")
    duplicate_ids = sorted({value for value in ids if value and ids.count(value) > 1})
    if duplicate_ids:
        errors.append(f"Found duplicate ID values: {', '.join(duplicate_ids)}.")

    orders = [order for _, order, _ in records]
    expected_orders = list(range(1, len(records) + 1))
    if orders != expected_orders:
        errors.append(
            "Broken global order: record orders must be consecutive from 1 "
            f"through {len(records)}."
        )
    if statistics.get("element_count") != len(records):
        errors.append("statistics.element_count does not match all source records.")

    for record_id, _, record in records:
        record_type = record.get("type")
        if record_type not in SUPPORTED_ELEMENT_TYPES | {
            "scene_heading",
            "title_page",
        }:
            errors.append(
                f"{record_id or 'missing-id'} has unsupported type {record_type!r}."
            )

    missing_english = collect_missing_english(payload)
    if final:
        errors.extend(f"Missing required English field: {field}." for field in missing_english)
        for record_id, _, record in records:
            for flag in _blocking_review_flags(record):
                errors.append(
                    f"{record_id or 'missing-id'} has unresolved review flag: {flag}."
                )

    if errors:
        raise DramaValidationError("\n".join(errors))
    return missing_english


def verify_source_pdf(payload: dict, source_pdf: Path) -> None:
    expected = str(payload.get("metadata", {}).get("source_sha256") or "").lower()
    if not expected:
        raise DramaValidationError("metadata.source_sha256 is required.")
    try:
        actual = hashlib.sha256(source_pdf.read_bytes()).hexdigest()
    except OSError as exc:
        raise DramaValidationError(f"Unable to read source PDF: {exc}") from exc
    if actual != expected:
        raise DramaValidationError(
            f"Source PDF SHA-256 mismatch: expected {expected}, got {actual}."
        )


def _inches(value, fallback: float) -> float:
    return float(value.inches) if value is not None else fallback


def inspect_reference_format(reference_path: Path) -> DramaFormat:
    try:
        doc = Document(str(reference_path))
    except (OSError, ValueError, BadZipFile) as exc:
        raise DramaValidationError(f"Unable to read reference DOCX: {exc}") from exc
    section = doc.sections[0]
    return DramaFormat(
        page_width_inches=_inches(section.page_width, 8.5),
        page_height_inches=_inches(section.page_height, 11),
        top_margin_inches=_inches(section.top_margin, 1),
        bottom_margin_inches=_inches(section.bottom_margin, 1),
        left_margin_inches=_inches(section.left_margin, 1.25),
        right_margin_inches=_inches(section.right_margin, 1.25),
        header_distance_inches=_inches(section.header_distance, 0.5),
        footer_distance_inches=_inches(section.footer_distance, 0.5),
    )


def _add_style(
    doc: Document,
    name: str,
    format_spec: DramaFormat,
    *,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
):
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_style_font_family(
        style,
        ascii_font_name=format_spec.font_name,
        east_asia_font_name=format_spec.east_asia_font_name,
    )
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    paragraph_format = style.paragraph_format
    if format_spec.space_before_pt is not None:
        paragraph_format.space_before = Pt(format_spec.space_before_pt)
    if format_spec.space_after_pt is not None:
        paragraph_format.space_after = Pt(format_spec.space_after_pt)
    if format_spec.line_spacing is not None:
        paragraph_format.line_spacing = format_spec.line_spacing
    return style


def _configure_document(doc: Document, format_spec: DramaFormat) -> None:
    ensure_base_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(format_spec.page_width_inches)
    section.page_height = Inches(format_spec.page_height_inches)
    section.top_margin = Inches(format_spec.top_margin_inches)
    section.bottom_margin = Inches(format_spec.bottom_margin_inches)
    section.left_margin = Inches(format_spec.left_margin_inches)
    section.right_margin = Inches(format_spec.right_margin_inches)
    section.header_distance = Inches(format_spec.header_distance_inches)
    section.footer_distance = Inches(format_spec.footer_distance_inches)
    _add_style(
        doc,
        "Drama Title",
        format_spec,
        size_pt=format_spec.title_size_pt,
        bold=True,
    )
    _add_style(
        doc,
        "Drama Scene Heading",
        format_spec,
        size_pt=format_spec.heading_size_pt,
        bold=True,
    ).paragraph_format.keep_with_next = True
    _add_style(
        doc,
        "Drama Action",
        format_spec,
        size_pt=format_spec.body_size_pt,
        italic=True,
    )
    _add_style(
        doc,
        "Drama Dialogue",
        format_spec,
        size_pt=format_spec.body_size_pt,
    )
    _add_style(
        doc,
        "Drama Note",
        format_spec,
        size_pt=format_spec.body_size_pt,
    )
    preview = _add_style(
        doc,
        "Drama Preview",
        format_spec,
        size_pt=format_spec.body_size_pt,
        bold=True,
    )
    preview.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _bookmark_name(source_id: str) -> str:
    return SOURCE_BOOKMARK_PREFIX + source_id.encode("utf-8").hex()


def _add_source_bookmark(paragraph, source_id: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), _bookmark_name(source_id))
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def _add_source_bookmark_around_run(run, source_id: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), _bookmark_name(source_id))
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    run._r.addprevious(start)
    run._r.addnext(end)


def _add_empty_source_bookmark(paragraph, source_id: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), _bookmark_name(source_id))
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(start)
    paragraph._p.append(end)


def _set_run_font(run, format_spec: DramaFormat, *, size_pt: float | None = None) -> None:
    set_run_font_family(
        run,
        ascii_font_name=format_spec.font_name,
        east_asia_font_name=format_spec.east_asia_font_name,
    )
    run.font.size = Pt(size_pt or format_spec.body_size_pt)


def _add_title_page(
    doc: Document,
    item: dict,
    format_spec: DramaFormat,
    *,
    first: bool,
):
    paragraph = doc.add_paragraph(style="Drama Title" if first else "Drama Note")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = str(item["translation_en"]).splitlines()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        _set_run_font(
            run,
            format_spec,
            size_pt=format_spec.title_size_pt if first else format_spec.body_size_pt,
        )
        run.bold = first
        if index < len(lines) - 1:
            run.add_break()
    return paragraph


def _add_scene_heading(doc: Document, scene: dict, format_spec: DramaFormat):
    paragraph = doc.add_paragraph(style="Drama Scene Heading")
    run = paragraph.add_run(str(scene["heading_en"]).strip())
    _set_run_font(run, format_spec, size_pt=format_spec.heading_size_pt)
    run.bold = True
    return paragraph


def _add_action(doc: Document, item: dict, format_spec: DramaFormat):
    paragraph = doc.add_paragraph(style="Drama Action")
    run = paragraph.add_run(str(item["translation_en"]).strip())
    _set_run_font(run, format_spec)
    run.italic = True
    return paragraph


def _add_dialogue(doc: Document, item: dict, format_spec: DramaFormat):
    paragraph = doc.add_paragraph(style="Drama Dialogue")
    speaker = str(item.get("speaker_en") or "").strip()
    qualifier = str(item.get("qualifier") or "").strip()
    if qualifier and qualifier not in speaker:
        speaker = f"{speaker} ({qualifier})"
    speaker_run = paragraph.add_run(speaker)
    _set_run_font(speaker_run, format_spec)
    speaker_run.bold = True
    parenthetical = str(item.get("parenthetical_en") or "").strip()
    if parenthetical:
        parenthetical_run = paragraph.add_run(f"\n({parenthetical})")
        _set_run_font(parenthetical_run, format_spec)
    dialogue_run = paragraph.add_run(f"\n{str(item['translation_en']).strip()}")
    _set_run_font(dialogue_run, format_spec)
    dialogue_run.bold = False
    return paragraph


def _add_note(doc: Document, item: dict, format_spec: DramaFormat):
    paragraph = doc.add_paragraph(style="Drama Note")
    text = str(item["translation_en"]).strip()
    marker_match = re.match(r"^(\([^()]+\)|[^:]+:)(.*)$", text, re.DOTALL)
    if marker_match:
        marker = paragraph.add_run(marker_match.group(1))
        _set_run_font(marker, format_spec)
        marker.bold = True
        remainder = paragraph.add_run(marker_match.group(2))
        _set_run_font(remainder, format_spec)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run, format_spec)
        run.bold = item.get("type") == "production_note"
    return paragraph


def _add_upright_text(
    doc: Document,
    item: dict,
    format_spec: DramaFormat,
    *,
    bold: bool = False,
):
    paragraph = doc.add_paragraph(style="Drama Note")
    run = paragraph.add_run(str(item["translation_en"]).strip())
    _set_run_font(run, format_spec)
    run.bold = bold
    run.italic = False
    return paragraph


def classify_presentation_types(
    records: list[tuple[str, str, dict]],
) -> list[str]:
    presentation_types: list[str] = []
    title_card_mode = False
    previous_record_type: str | None = None
    for _, record_type, record in records:
        if record_type == "scene_heading":
            title_card_mode = False

        text = str(
            record.get("heading_en")
            if record_type == "scene_heading"
            else record.get("translation_en") or ""
        ).strip()
        presentation_type = record_type
        if record_type == "production_note" and text.upper().startswith(
            "TITLE CARDS:"
        ):
            title_card_mode = True
        elif record_type == "action" and title_card_mode:
            presentation_type = "title_card_text"
        elif record_type == "action" and text.upper() == "DELETE.":
            presentation_type = "production_directive"
        elif (
            record_type == "action"
            and text[:1].islower()
            and previous_record_type in {"production_note", "translation_note"}
        ):
            presentation_type = "note_continuation"
        elif title_card_mode and record_type != "action":
            title_card_mode = False

        presentation_types.append(presentation_type)
        previous_record_type = record_type
    return presentation_types


def _iter_preview_records(payload: dict) -> Iterable[tuple[str, str, dict]]:
    for item in payload["front_matter"]:
        if _nonempty(item.get("translation_en")):
            yield str(item["id"]), "title_page", item
    for scene in payload["scenes"]:
        if not _nonempty(scene.get("heading_en")):
            continue
        yield str(scene["heading_id"]), "scene_heading", scene
        for item in scene["elements"]:
            has_required_english = _nonempty(item.get("translation_en"))
            if item.get("type") == "dialogue":
                if _nonempty(item.get("speaker_zh")):
                    has_required_english = has_required_english and _nonempty(
                        item.get("speaker_en")
                    )
                if _nonempty(item.get("parenthetical_zh")):
                    has_required_english = has_required_english and _nonempty(
                        item.get("parenthetical_en")
                    )
            if has_required_english:
                yield str(item["id"]), str(item["type"]), item


def _iter_final_records(payload: dict) -> Iterable[tuple[str, str, dict]]:
    for item in payload["front_matter"]:
        yield str(item["id"]), "title_page", item
    for scene in payload["scenes"]:
        yield str(scene["heading_id"]), "scene_heading", scene
        for item in scene["elements"]:
            yield str(item["id"]), str(item["type"]), item


def read_document_mapping(path: Path) -> list[str]:
    try:
        with ZipFile(path) as package:
            xml = package.read("word/document.xml")
    except (OSError, BadZipFile, KeyError) as exc:
        raise DramaValidationError(f"Invalid DOCX package: {exc}") from exc
    root = ElementTree.fromstring(xml)
    mapped: list[str] = []
    for bookmark in root.iter(qn("w:bookmarkStart")):
        name = bookmark.get(qn("w:name"), "")
        if not name.startswith(SOURCE_BOOKMARK_PREFIX):
            continue
        encoded = name[len(SOURCE_BOOKMARK_PREFIX) :]
        try:
            mapped.append(bytes.fromhex(encoded).decode("utf-8"))
        except (ValueError, UnicodeDecodeError) as exc:
            raise DramaValidationError(f"Invalid source bookmark {name!r}.") from exc
    return mapped


def validate_generated_document(path: Path, expected_ids: list[str]) -> None:
    try:
        with ZipFile(path) as package:
            bad_member = package.testzip()
            if bad_member:
                raise DramaValidationError(f"Corrupt DOCX member: {bad_member}.")
        doc = Document(str(path))
    except (OSError, BadZipFile, ValueError) as exc:
        raise DramaValidationError(f"Generated DOCX cannot be opened: {exc}") from exc
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if RAW_PLACEHOLDER_RE.search(text):
        raise DramaValidationError("Generated DOCX contains a raw template placeholder.")
    mapped_ids = read_document_mapping(path)
    if mapped_ids != expected_ids:
        raise DramaValidationError(
            "Generated document source mapping is incomplete, duplicated, or out of order."
        )


def generate_drama(
    input_path: Path,
    reference_path: Path,
    output_path: Path,
    *,
    preview: bool,
    verify_source_path: Path | None = None,
) -> GenerationResult:
    payload = load_payload(input_path)
    missing_fields = validate_payload(payload, final=not preview)
    if verify_source_path is not None:
        verify_source_pdf(payload, verify_source_path)
    format_spec = inspect_reference_format(reference_path)
    doc = Document()
    _configure_document(doc, format_spec)
    doc.core_properties.title = "Drama shooting script"
    doc.core_properties.subject = (
        "INCOMPLETE PREVIEW" if preview else "Final English shooting script"
    )

    if preview:
        banner = doc.add_paragraph(style="Drama Preview")
        banner.add_run("PREVIEW — INCOMPLETE TRANSLATION")
        detail = doc.add_paragraph(style="Drama Preview")
        detail.add_run(
            f"{len(missing_fields)} required English fields remain untranslated."
        )

    records = list(
        _iter_preview_records(payload) if preview else _iter_final_records(payload)
    )
    presentation_types = classify_presentation_types(records)
    generated_ids: list[str] = []
    front_count = sum(1 for _, record_type, _ in records if record_type == "title_page")
    front_seen = 0
    title_paragraph = None
    bookmark_id = 1
    for (source_id, record_type, record), presentation_type in zip(
        records, presentation_types
    ):
        if record_type == "title_page":
            if front_seen == 0:
                paragraph = _add_title_page(
                    doc,
                    record,
                    format_spec,
                    first=True,
                )
                title_paragraph = paragraph
            else:
                if title_paragraph is None:
                    raise DramaValidationError("Unable to map title-page metadata.")
                _add_empty_source_bookmark(
                    title_paragraph,
                    source_id,
                    bookmark_id,
                )
                bookmark_id += 1
                generated_ids.append(source_id)
                front_seen += 1
                if front_seen == front_count:
                    title_paragraph.paragraph_format.page_break_after = True
                continue
            front_seen += 1
            if front_seen == front_count:
                paragraph.paragraph_format.page_break_after = True
        elif record_type == "scene_heading":
            paragraph = _add_scene_heading(doc, record, format_spec)
        elif presentation_type == "title_card_text":
            paragraph = _add_upright_text(doc, record, format_spec)
        elif presentation_type == "note_continuation":
            continuation = str(record["translation_en"]).strip()
            paragraph = doc.paragraphs[-1]
            previous_run = next(
                (run for run in reversed(paragraph.runs) if run.text),
                None,
            )
            separator = "" if paragraph.text.endswith(("-", "—", "/")) else " "
            run = paragraph.add_run(separator + continuation)
            _set_run_font(run, format_spec)
            run.bold = previous_run.bold if previous_run is not None else False
            run.italic = previous_run.italic if previous_run is not None else False
            _add_source_bookmark_around_run(run, source_id, bookmark_id)
            bookmark_id += 1
            generated_ids.append(source_id)
            continue
        elif presentation_type == "production_directive":
            paragraph = _add_upright_text(doc, record, format_spec, bold=True)
        elif record_type == "action":
            paragraph = _add_action(doc, record, format_spec)
        elif record_type == "dialogue":
            paragraph = _add_dialogue(doc, record, format_spec)
        else:
            paragraph = _add_note(doc, record, format_spec)
        _add_source_bookmark(paragraph, source_id, bookmark_id)
        bookmark_id += 1
        generated_ids.append(source_id)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    validate_generated_document(output_path, generated_ids)
    return GenerationResult(output_path, generated_ids, missing_fields)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generate an English shooting-script DOCX from Schema 2.0 JSON."
    )
    parser.add_argument("--input", required=True, type=Path, help="Bilingual JSON input.")
    parser.add_argument(
        "--reference",
        required=True,
        type=Path,
        help="Manager-produced English DOCX formatting reference.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        help="Output DOCX path; defaults beside the input with an _英文 suffix.",
    )
    parser.add_argument(
        "--preview",
        action="store_true",
        help="Generate translated records only and mark the DOCX as incomplete.",
    )
    parser.add_argument(
        "--verify-source",
        type=Path,
        help="Optional source PDF whose SHA-256 must match JSON metadata.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    output_path = resolve_output_path(args.input, args.output, preview=args.preview)
    try:
        result = generate_drama(
            args.input,
            args.reference,
            output_path,
            preview=args.preview,
            verify_source_path=args.verify_source,
        )
    except DramaValidationError as exc:
        print(f"[error] {exc}", file=sys.stderr)
        return 2
    if args.preview:
        print(
            f"[PREVIEW] {result.missing_count} required English fields remain "
            "untranslated."
        )
    print(f"[created] {result.output_path}")
    print(f"[mapped] {len(result.generated_ids)} source records")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
