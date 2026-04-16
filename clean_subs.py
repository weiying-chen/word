#!/usr/bin/env python3

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


SOURCE_LINK_RE = re.compile(r"^https?://\S+$")
SUBTITLE_LINE_RE = re.compile(
    r"^(?:[^\t]+\t)?\d{2}:\d{2}:\d{2}:\d{2}\t\d{2}:\d{2}:\d{2}:\d{2}\t"
)
INTRO_LABELS = {"簡介：", "簡介:"}
SUBTITLE_LABELS = {"字幕：", "字幕:"}
SECTION_LABELS = {
    "建議YT標題：",
    "建議YT標題:",
    "建議標題：",
    "建議標題:",
    "簡介：",
    "簡介:",
    "選圖：",
    "選圖:",
    "字幕：",
    "字幕:",
}


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _is_url_paragraph(paragraph) -> bool:
    return bool(SOURCE_LINK_RE.match(paragraph.text.strip()))


def _is_subtitle_paragraph(paragraph) -> bool:
    return bool(SUBTITLE_LINE_RE.match(paragraph.text.strip()))


def _section_kind(paragraph) -> str | None:
    text = paragraph.text.strip()
    if text in INTRO_LABELS:
        return "intro"
    if text in SUBTITLE_LABELS:
        return "subs"
    if text in SECTION_LABELS:
        return "other"
    return None


def remove_sources_from_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))
    paragraphs = list(doc.paragraphs)
    remove_indexes: set[int] = set()
    idx = 0
    current_section: str | None = None

    while idx < len(paragraphs):
        paragraph = paragraphs[idx]
        section = _section_kind(paragraph)
        if section is not None:
            current_section = section
            idx += 1
            continue

        if current_section not in {"intro", "subs"}:
            idx += 1
            continue
        if not _is_url_paragraph(paragraph):
            idx += 1
            continue

        remove_indexes.add(idx)
        idx += 1

        while idx < len(paragraphs):
            current = paragraphs[idx]
            stripped = current.text.strip()
            next_section = _section_kind(current)
            if next_section is not None:
                break
            if _is_subtitle_paragraph(current):
                break
            remove_indexes.add(idx)
            idx += 1
            if not stripped:
                break

    for index in sorted(remove_indexes, reverse=True):
        _remove_paragraph(paragraphs[index])

    paragraphs = list(doc.paragraphs)
    remove_blank_indexes: set[int] = set()
    current_section: str | None = None
    for idx, paragraph in enumerate(paragraphs):
        section = _section_kind(paragraph)
        if section is not None:
            current_section = section
            continue
        if current_section != "subs":
            continue
        if paragraph.text.strip():
            continue
        if idx == 0 or idx + 1 >= len(paragraphs):
            continue
        prev_paragraph = paragraphs[idx - 1]
        next_paragraph = paragraphs[idx + 1]
        if _is_subtitle_paragraph(prev_paragraph) and _is_subtitle_paragraph(next_paragraph):
            remove_blank_indexes.add(idx)

    for index in sorted(remove_blank_indexes, reverse=True):
        _remove_paragraph(paragraphs[index])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _default_output_path(path: Path) -> Path:
    return path.with_name(f"{path.stem}_nosource{path.suffix}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Remove source URL blocks from generated DOCX files."
    )
    parser.add_argument("docx_files", nargs="+", help="Input DOCX file paths.")
    parser.add_argument(
        "--in-place",
        action="store_true",
        help="Overwrite input files instead of creating _nosource copies.",
    )
    args = parser.parse_args()

    for raw_path in args.docx_files:
        input_path = Path(raw_path)
        output_path = input_path if args.in_place else _default_output_path(input_path)
        remove_sources_from_docx(input_path, output_path)


if __name__ == "__main__":
    main()
