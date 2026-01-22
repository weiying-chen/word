#!/usr/bin/env python3

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches


TITLE_PLACEHOLDER = "{{TITLE_EN}}"
PEOPLE_PLACEHOLDER = "{{PEOPLE}}"
OVERVIEW_PLACEHOLDER = "{{OVERVIEW_EN}}"


def load_payload(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def find_paragraph_by_text(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def replace_multiline(paragraph: Paragraph, lines: list[str]) -> None:
    paragraph.text = ""
    if not lines:
        remove_paragraph(paragraph)
        return
    paragraph.add_run(lines[0])
    current = paragraph
    for line in lines[1:]:
        current = insert_paragraph_after(current, line)


def build_people_lines(people: list[dict]) -> list[str]:
    lines: list[str] = []
    for idx, person in enumerate(people):
        label_zh = person.get("label_zh")
        if not label_zh:
            role_zh = person.get("role_zh", "").strip()
            name_zh = person.get("name_zh", "").strip()
            if role_zh and name_zh:
                label_zh = f"{role_zh}｜{name_zh}"
            else:
                label_zh = role_zh or name_zh
        lines.append(label_zh or "")
        lines.append(person.get("name_en", ""))
        lines.append(person.get("role_en", ""))
        if idx < len(people) - 1:
            lines.append("")
    return lines


def apply_default_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def render_meta(template_path: Path, payload_path: Path, output_path: Path) -> None:
    data = load_payload(payload_path)
    doc = Document(str(template_path))
    apply_default_margins(doc)

    title_placeholder = find_paragraph_by_text(doc, TITLE_PLACEHOLDER)
    if title_placeholder:
        title_placeholder.text = data.get("title_en") or data.get("title_zh", "")

    people_placeholder = find_paragraph_by_text(doc, PEOPLE_PLACEHOLDER)
    if people_placeholder:
        replace_multiline(people_placeholder, build_people_lines(data.get("people", [])))

    overview_placeholder = find_paragraph_by_text(doc, OVERVIEW_PLACEHOLDER)
    if overview_placeholder:
        overview_placeholder.text = data.get("overview_en") or data.get("summary_zh", "")

    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Render meta.docx from JSON data.")
    parser.add_argument(
        "--template",
        default="meta_template.docx",
        help="Path to the meta DOCX template.",
    )
    parser.add_argument(
        "--input",
        default="meta_filled.json",
        help="Path to the filled JSON payload.",
    )
    parser.add_argument(
        "--output",
        default="meta.docx",
        help="Path to write the rendered meta DOCX.",
    )
    args = parser.parse_args()

    render_meta(Path(args.template), Path(args.input), Path(args.output))


if __name__ == "__main__":
    main()
