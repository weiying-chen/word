from pathlib import Path
import os
import tempfile
import warnings
import zipfile

from docx import Document
from lxml import etree

from docx_utils import add_hyperlink
import generate_news


def _write_source_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Community Clinic Brings Care to Coastal Town")
    link_paragraph = doc.add_paragraph("")
    add_hyperlink(link_paragraph, "https://example.com/news/story", "https://example.com/news/story")
    doc.add_paragraph("Volunteers organized a two-day clinic to support families in a coastal town.")
    doc.add_paragraph("(  11/16~17 )")
    doc.add_paragraph("")
    doc.add_paragraph("<")
    doc.add_paragraph("old body line")
    doc.save(path)


def _write_source_docx_with_linked_title(path: Path) -> None:
    doc = Document()
    title_para = doc.add_paragraph("")
    add_hyperlink(
        title_para,
        "Nobel Laureate visits Jing Si Abode",
        "https://www.daai.tv/news/taiwan/593768",
    )
    doc.add_paragraph("")
    doc.add_paragraph("Summary before marker.")
    doc.add_paragraph("<")
    doc.add_paragraph("old body line")
    doc.save(path)


def _write_template_docx(path: Path, marker: str = "{{BODY}}") -> None:
    doc = Document()
    doc.add_paragraph(marker)
    doc.save(path)


def test_parse_input_extracts_body_section(tmp_path: Path) -> None:
    input_path = tmp_path / "news_input.txt"
    input_path.write_text(
        "\n".join(
            [
                "TITLE_TEXT: Ignored title",
                "SUMMARY:",
                "Ignored summary",
                "",
                "BODY:",
                "1_0014",
                "中文內文。",
                "English line.",
            ]
        ),
        encoding="utf-8",
    )

    data = generate_news.parse_input(input_path)

    assert data["BODY"] == "1_0014\n中文內文。\nEnglish line."


def test_parse_input_uses_whole_file_when_no_body_label(tmp_path: Path) -> None:
    input_path = tmp_path / "news_input.txt"
    input_path.write_text("1_0014\n中文內文。\nEnglish line.\n", encoding="utf-8")

    data = generate_news.parse_input(input_path)

    assert data["BODY"] == "1_0014\n中文內文。\nEnglish line."


def test_parse_input_fallback_encoding_warns_and_rewrites_utf8(tmp_path: Path) -> None:
    input_path = tmp_path / "news_input.txt"
    input_path.write_bytes("BODY: Café News\n".encode("cp1252"))

    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        data = generate_news.parse_input(input_path)

    assert data["BODY"] == "Café News"
    assert caught
    assert "fallback encoding" in str(caught[0].message).lower()
    assert "cp1252" in str(caught[0].message).lower()
    assert input_path.read_bytes() == "BODY: Café News\n".encode("utf-8")


def test_generate_news_preserves_header_and_replaces_body(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    input_path = tmp_path / "news_input.txt"
    output_path = tmp_path / "news_output.docx"

    _write_source_docx(source_docx)
    _write_template_docx(template_docx, marker="<")
    input_path.write_text(
        "\n".join(
            [
                "BODY:",
                "1_0014",
                "Two days of free screenings brought steady lines of local residents.",
                "",
                "2_0025",
                "Residents arrived early to receive free screenings and follow-up advice.",
            ]
        ),
        encoding="utf-8",
    )

    generate_news.generate_news(template_docx, source_docx, input_path, output_path)

    doc = Document(output_path)
    texts = [p.text for p in doc.paragraphs]
    assert texts == [
        "<",
        "",
        "Community Clinic Brings Care to Coastal Town",
        "https://example.com/news/story",
        "Volunteers organized a two-day clinic to support families in a coastal town.",
        "(  11/16~17 )",
        "",
        "<",
        "",
        "1_0014",
        "Two days of free screenings brought steady lines of local residents.",
        "",
        "2_0025",
        "Residents arrived early to receive free screenings and follow-up advice.",
    ]
    assert "old body line" not in texts
    assert any(
        run.font.highlight_color and run.font.highlight_color.name == "BRIGHT_GREEN"
        for run in doc.paragraphs[9].runs
    )
    assert any(
        run.font.highlight_color and run.font.highlight_color.name == "BRIGHT_GREEN"
        for run in doc.paragraphs[12].runs
    )

    with zipfile.ZipFile(output_path) as zf:
        document_xml = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = document_xml.findall(".//w:body/w:p", ns)
    assert paragraphs
    hyperlink_count = sum(
        len(paragraph.findall("w:hyperlink", ns)) for paragraph in paragraphs
    )
    assert hyperlink_count == 1


def test_generate_news_from_sources_uses_body_text_file(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    source_txt = tmp_path / "source.txt"
    output_path = tmp_path / "news_output.docx"

    _write_source_docx(source_docx)
    _write_template_docx(template_docx, marker="<")
    source_txt.write_text("1_0001\n中文內文。\nEnglish body line.\n", encoding="utf-8")

    generate_news.generate_news_from_sources(
        template_docx, source_docx, source_txt, output_path
    )

    texts = [p.text for p in Document(output_path).paragraphs]
    assert texts[-5:] == ["<", "", "1_0001", "中文內文。", "English body line."]


def test_generate_news_omits_standalone_tilde_placeholder_lines(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    source_txt = tmp_path / "source.txt"
    output_path = tmp_path / "news_output.docx"

    _write_source_docx(source_docx)
    _write_template_docx(template_docx, marker="<")
    source_txt.write_text(
        "1_0001\n中文內文。\n~\nEnglish body line.\n",
        encoding="utf-8",
    )

    generate_news.generate_news_from_sources(
        template_docx, source_docx, source_txt, output_path
    )

    texts = [p.text for p in Document(output_path).paragraphs]
    assert "~" not in texts
    assert texts[-5:] == ["<", "", "1_0001", "中文內文。", "English body line."]


def test_generate_news_keeps_single_blank_after_marker_when_body_starts_blank(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    input_path = tmp_path / "news_input.txt"
    output_path = tmp_path / "news_output.docx"

    _write_source_docx(source_docx)
    _write_template_docx(template_docx, marker="<")
    input_path.write_text("\nLeading body line.\n", encoding="utf-8")

    generate_news.generate_news(template_docx, source_docx, input_path, output_path)

    texts = [p.text for p in Document(output_path).paragraphs]
    marker_index = texts.index("<")
    assert texts[marker_index + 1 : marker_index + 4] == [
        "",
        "Community Clinic Brings Care to Coastal Town",
        "https://example.com/news/story",
    ]


def test_generate_news_allows_source_docx_without_marker(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    input_path = tmp_path / "news_input.txt"
    output_path = tmp_path / "news_output.docx"

    doc = Document()
    doc.add_paragraph("Header without marker")
    doc.save(source_docx)
    _write_template_docx(template_docx, marker="<")
    input_path.write_text("BODY:\n1_0001\nBody line.\n", encoding="utf-8")

    generate_news.generate_news(template_docx, source_docx, input_path, output_path)

    texts = [p.text for p in Document(output_path).paragraphs]
    assert texts == ["<", "", "1_0001", "Body line."]


def test_generate_news_with_body_placeholder_template(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    input_path = tmp_path / "news_input.txt"
    output_path = tmp_path / "news_output.docx"

    _write_source_docx(source_docx)
    _write_template_docx(template_docx, marker="{{BODY}}")
    input_path.write_text("BODY:\n1_0001\nBody line.\n", encoding="utf-8")

    generate_news.generate_news(template_docx, source_docx, input_path, output_path)

    doc = Document(output_path)
    texts = [p.text for p in doc.paragraphs]
    assert texts == [
        "Community Clinic Brings Care to Coastal Town",
        "https://example.com/news/story",
        "Volunteers organized a two-day clinic to support families in a coastal town.",
        "(  11/16~17 )",
        "",
        "<",
        "",
        "1_0001",
        "Body line.",
    ]
    with zipfile.ZipFile(output_path) as zf:
        document_xml = etree.fromstring(zf.read("word/document.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = document_xml.findall(".//w:body/w:p", ns)
    hyperlinks = paragraphs[1].findall("w:hyperlink", ns)
    assert len(hyperlinks) == 1


def test_generate_news_preserves_non_url_header_hyperlink(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    input_path = tmp_path / "news_input.txt"
    output_path = tmp_path / "news_output.docx"

    _write_source_docx_with_linked_title(source_docx)
    _write_template_docx(template_docx, marker="{{BODY}}")
    input_path.write_text("BODY:\n1_0001\nBody line.\n", encoding="utf-8")

    generate_news.generate_news(template_docx, source_docx, input_path, output_path)

    doc = Document(output_path)
    texts = [p.text for p in doc.paragraphs]
    assert texts[:6] == [
        "Nobel Laureate visits Jing Si Abode",
        "",
        "Summary before marker.",
        "<",
        "",
        "1_0001",
    ]
    rel_targets = [
        rel._target
        for rel in doc.part.rels.values()
        if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    ]
    assert "https://www.daai.tv/news/taiwan/593768" in rel_targets


def test_generate_news_body_placeholder_works_without_source_marker(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    template_docx = tmp_path / "template.docx"
    input_path = tmp_path / "news_input.txt"
    output_path = tmp_path / "news_output.docx"

    doc = Document()
    doc.add_paragraph("Header without marker")
    doc.save(source_docx)
    _write_template_docx(template_docx, marker="{{BODY}}")
    input_path.write_text("BODY:\n1_0001\nBody line.\n", encoding="utf-8")

    generate_news.generate_news(template_docx, source_docx, input_path, output_path)

    texts = [p.text for p in Document(output_path).paragraphs]
    assert texts == ["1_0001", "Body line."]


def test_default_output_path_uses_source_stem_with_final_suffix(tmp_path: Path) -> None:
    source = tmp_path / "coastal_story.docx"
    output_dir = tmp_path / "output"
    output = generate_news.default_output_path(source, output_dir)
    assert output == output_dir / "coastal_story_final.docx"


def test_default_output_path_preserves_existing_final_suffix(tmp_path: Path) -> None:
    source = tmp_path / "coastal_story_final.docx"
    output_dir = tmp_path / "output"
    output = generate_news.default_output_path(source, output_dir)
    assert output == output_dir / "coastal_story_final.docx"


def test_resolve_template_path_uses_script_directory_for_relative_paths() -> None:
    previous_cwd = Path.cwd()
    with tempfile.TemporaryDirectory() as tmpdir:
        os.chdir(tmpdir)
        try:
            template = generate_news.resolve_template_path(
                Path("templates/news_template.docx")
            )
        finally:
            os.chdir(previous_cwd)
    assert template == Path(__file__).resolve().parent.parent / "templates" / "news_template.docx"
