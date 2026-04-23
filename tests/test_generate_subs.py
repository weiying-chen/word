from pathlib import Path
import zipfile
import warnings

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from lxml import etree

from docx_utils import add_hyperlink
import generate_subs


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def _write_source_docx(
    path: Path,
    *,
    header_paragraphs: list[str] | None = None,
    body_paragraphs: list[str] | None = None,
) -> None:
    doc = Document()
    header_items = header_paragraphs if header_paragraphs is not None else [
        "Original title",
        "https://example.com/story",
        "",
        "Original summary.",
        "",
    ]
    for text in header_items:
        doc.add_paragraph(text)
    body_items = body_paragraphs if body_paragraphs is not None else [
        "00:00:00:00\t00:00:02:00\tFirst line.",
        "First translation.",
        "",
    ]
    for text in body_items:
        doc.add_paragraph(text)
    doc.save(path)


def _write_png(path: Path) -> None:
    path.write_bytes(
        bytes.fromhex(
            "89504E470D0A1A0A"
            "0000000D49484452000000010000000108060000001F15C489"
            "0000000D49444154789C63F8FFFFFF7F0009FB03FD2A86E38A"
            "0000000049454E44AE426082"
        )
    )


def _write_formatted_source_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Source title")
    link_paragraph = doc.add_paragraph("")
    add_hyperlink(
        link_paragraph,
        "https://example.com/source",
        "https://example.com/source",
    )
    doc.add_paragraph("")
    doc.add_paragraph("Source summary.")
    timing_paragraph = doc.add_paragraph("")
    timing_paragraph.add_run("07:34-09:41 (2分7秒)").font.highlight_color = (
        WD_COLOR_INDEX.YELLOW
    )
    subtitle_paragraph = doc.add_paragraph("")
    subtitle_paragraph.add_run(
        "00:00:00:00\t00:00:02:00\tFirst line."
    ).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    doc.add_paragraph("First translation.")
    doc.save(path)


def _rewrite_hyperlink_style_id(path: Path, style_id: str) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = ns["w"]
    with zipfile.ZipFile(path, "r") as zin:
        styles_xml = etree.fromstring(zin.read("word/styles.xml"))
        doc_xml = etree.fromstring(zin.read("word/document.xml"))
        style_nodes = styles_xml.xpath(
            ".//w:style[w:name[@w:val='Hyperlink']]",
            namespaces=ns,
        )
        if style_nodes:
            style_nodes[0].set("{%s}styleId" % w_ns, style_id)
        else:
            style = etree.SubElement(
                styles_xml,
                "{%s}style" % w_ns,
                attrib={
                    "{%s}type" % w_ns: "character",
                    "{%s}styleId" % w_ns: style_id,
                },
            )
            etree.SubElement(style, "{%s}name" % w_ns, attrib={"{%s}val" % w_ns: "Hyperlink"})
            r_pr = etree.SubElement(style, "{%s}rPr" % w_ns)
            etree.SubElement(
                r_pr,
                "{%s}color" % w_ns,
                attrib={"{%s}val" % w_ns: "0563C1", "{%s}themeColor" % w_ns: "hyperlink"},
            )
            etree.SubElement(r_pr, "{%s}u" % w_ns, attrib={"{%s}val" % w_ns: "single"})

        for node in doc_xml.xpath(".//w:hyperlink//w:rStyle", namespaces=ns):
            node.set("{%s}val" % w_ns, style_id)

        tmp_path = path.with_suffix(path.suffix + ".tmp")
        with zipfile.ZipFile(tmp_path, "w") as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/styles.xml":
                    data = etree.tostring(
                        styles_xml,
                        encoding="UTF-8",
                        standalone=True,
                        xml_declaration=True,
                    )
                elif info.filename == "word/document.xml":
                    data = etree.tostring(
                        doc_xml,
                        encoding="UTF-8",
                        standalone=True,
                        xml_declaration=True,
                    )
                zout.writestr(info, data)
    tmp_path.replace(path)


def _rewrite_hyperlink_as_field_code(path: Path) -> None:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    w_ns = ns["w"]
    r_ns = ns["r"]
    xml_ns = "http://www.w3.org/XML/1998/namespace"

    with zipfile.ZipFile(path, "r") as zin:
        doc_xml = etree.fromstring(zin.read("word/document.xml"))
        rels_xml = etree.fromstring(zin.read("word/_rels/document.xml.rels"))

        hyperlink = doc_xml.find(".//w:hyperlink", ns)
        if hyperlink is None:
            raise AssertionError("Expected source document to contain a hyperlink element.")

        relation_id = hyperlink.get("{%s}id" % r_ns)
        if not relation_id:
            raise AssertionError("Expected hyperlink to contain a relationship id.")
        relation = rels_xml.xpath(
            ".//pr:Relationship[@Id=$rid]",
            namespaces=ns,
            rid=relation_id,
        )
        if not relation:
            raise AssertionError("Expected relationship for hyperlink id.")
        target_url = relation[0].get("Target")
        if not target_url:
            raise AssertionError("Expected hyperlink relationship target URL.")

        display_text = "".join(t.text or "" for t in hyperlink.findall(".//w:t", ns))
        run_style = hyperlink.find(".//w:rPr/w:rStyle", ns)
        run_style_val = run_style.get("{%s}val" % w_ns) if run_style is not None else None

        parent = hyperlink.getparent()
        insert_at = parent.index(hyperlink)
        parent.remove(hyperlink)

        def _append_run_with_fld_char(fld_type: str) -> None:
            run = etree.Element("{%s}r" % w_ns)
            fld_char = etree.SubElement(run, "{%s}fldChar" % w_ns)
            fld_char.set("{%s}fldCharType" % w_ns, fld_type)
            parent.insert(insert_at + len(inserted_runs), run)
            inserted_runs.append(run)

        inserted_runs = []
        _append_run_with_fld_char("begin")

        run_instr = etree.Element("{%s}r" % w_ns)
        instr_text = etree.SubElement(run_instr, "{%s}instrText" % w_ns)
        instr_text.set("{%s}space" % xml_ns, "preserve")
        instr_text.text = f' HYPERLINK "{target_url}" '
        parent.insert(insert_at + len(inserted_runs), run_instr)
        inserted_runs.append(run_instr)

        _append_run_with_fld_char("separate")

        run_display = etree.Element("{%s}r" % w_ns)
        run_properties = etree.SubElement(run_display, "{%s}rPr" % w_ns)
        if run_style_val:
            run_style_node = etree.SubElement(run_properties, "{%s}rStyle" % w_ns)
            run_style_node.set("{%s}val" % w_ns, run_style_val)
        text_node = etree.SubElement(run_display, "{%s}t" % w_ns)
        text_node.text = display_text
        parent.insert(insert_at + len(inserted_runs), run_display)
        inserted_runs.append(run_display)

        _append_run_with_fld_char("end")

        tmp_path = path.with_suffix(path.suffix + ".tmp")
        with zipfile.ZipFile(tmp_path, "w") as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    data = etree.tostring(
                        doc_xml,
                        encoding="UTF-8",
                        standalone=True,
                        xml_declaration=True,
                    )
                zout.writestr(info, data)
    tmp_path.replace(path)


def _assert_yt_title_replaced_for_encoded_input(
    tmp_path: Path,
    encoded_text: bytes,
    expected_title: str,
) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["{{YT_TITLE_SUGGESTED}}"])
    _write_source_docx(source_docx, header_paragraphs=[])
    input_path.write_bytes(encoded_text)

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    doc = Document(output_path)
    assert doc.paragraphs[0].text == expected_title


def test_parse_input_multiline_intro(tmp_path: Path) -> None:
    input_path = tmp_path / "input.txt"
    input_path.write_text(
        "\n".join(
            [
                "YT_TITLE_SUGGESTED: Sample Title",
                "INTRO:",
                "First intro line.",
                "Second intro line.",
                "TITLE_SUGGESTED: Another Title",
                "",
            ]
        ),
        encoding="utf-8",
    )

    data = generate_subs.parse_input(input_path)
    assert data["INTRO"] == "First intro line.\nSecond intro line."


def test_yt_title_replacement_utf8_no_bom(tmp_path: Path) -> None:
    _assert_yt_title_replaced_for_encoded_input(
        tmp_path,
        "YT_TITLE_SUGGESTED: UTF8 Title\n".encode("utf-8"),
        "UTF8 Title",
    )


def test_yt_title_replacement_utf8_bom(tmp_path: Path) -> None:
    _assert_yt_title_replaced_for_encoded_input(
        tmp_path,
        b"\xef\xbb\xbfYT_TITLE_SUGGESTED: UTF8 BOM Title\n",
        "UTF8 BOM Title",
    )


def test_yt_title_replacement_utf16_le_bom(tmp_path: Path) -> None:
    _assert_yt_title_replaced_for_encoded_input(
        tmp_path,
        "YT_TITLE_SUGGESTED: UTF16 BOM Title\n".encode("utf-16"),
        "UTF16 BOM Title",
    )


def test_parse_input_fallback_encoding_warns_and_rewrites_utf8(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "input.txt"
    input_path.write_bytes("YT_TITLE_SUGGESTED: Café News\n".encode("cp1252"))

    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        data = generate_subs.parse_input(input_path)

    assert data["YT_TITLE_SUGGESTED"] == "Café News"
    assert caught
    assert "fallback encoding" in str(caught[0].message).lower()
    assert "cp1252" in str(caught[0].message).lower()
    assert input_path.read_bytes() == "YT_TITLE_SUGGESTED: Café News\n".encode("utf-8")


def test_yt_title_replacement_fallback_encoding(tmp_path: Path) -> None:
    with pytest.warns(UserWarning, match="fallback encoding 'cp1252'"):
        _assert_yt_title_replaced_for_encoded_input(
            tmp_path,
            "YT_TITLE_SUGGESTED: Café from fallback\n".encode("cp1252"),
            "Café from fallback",
        )


def test_generate_subs_copies_source_header_before_generated_sections(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["建議YT標題：", "{{YT_TITLE_SUGGESTED}}", "字幕："])
    _write_source_docx(
        source_docx,
        header_paragraphs=[
            "Source title",
            "https://example.com/source",
            "",
            "Source summary.",
            "",
        ],
    )
    input_path.write_text("YT_TITLE_SUGGESTED: Suggested title\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    texts = [p.text for p in Document(output_path).paragraphs]

    assert texts[:8] == [
        "Source title",
        "https://example.com/source",
        "",
        "Source summary.",
        "",
        "建議YT標題：",
        "",
        "Suggested title",
    ]


def test_generate_subs_removes_empty_intro_paragraph(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["{{INTRO}}", "After intro."])
    _write_source_docx(source_docx)
    input_path.write_text("YT_TITLE_SUGGESTED: Sample Title\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    doc = Document(output_path)
    texts = [p.text for p in doc.paragraphs if p.text]
    assert "After intro." in texts
    assert "Original title" in texts


def test_generate_subs_inserts_blank_after_labels(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["簡介：", "Line after label."])
    _write_source_docx(source_docx)
    input_path.write_text("YT_TITLE_SUGGESTED: Sample Title\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    doc = Document(output_path)
    label_idx = [p.text for p in doc.paragraphs].index("簡介：")
    assert not doc.paragraphs[label_idx + 1].text.strip()
    assert doc.paragraphs[label_idx + 2].text.strip() == "Line after label."


def test_generate_subs_inserts_blank_after_other_chinese_labels(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["建議YT標題：", "{{YT_TITLE_SUGGESTED}}"])
    _write_source_docx(source_docx, header_paragraphs=[])
    input_path.write_text("YT_TITLE_SUGGESTED: Suggested title\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    texts = [p.text for p in Document(output_path).paragraphs]
    label_idx = texts.index("建議YT標題：")
    assert texts[label_idx + 1] == ""
    assert texts[label_idx + 2] == "Suggested title"


def test_generate_subs_inserts_blank_after_last_label(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["字幕："])
    _write_source_docx(source_docx)
    input_path.write_text("YT_TITLE_SUGGESTED: Sample Title\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    doc = Document(output_path)
    label_idx = [p.text for p in doc.paragraphs].index("字幕：")
    assert not doc.paragraphs[label_idx + 1].text.strip()
    assert doc.paragraphs[label_idx + 2].text == "00:00:00:00\t00:00:02:00\tFirst line."


def test_generate_subs_adds_thumbnail_credit_after_image(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"
    image_path = tmp_path / "thumbnail.png"

    _write_docx(template_path, ["選圖：", "{{THUMBNAIL}}", "After thumbnail."])
    _write_source_docx(source_docx)
    _write_png(image_path)
    input_path.write_text(
        "\n".join(
            [
                f"THUMBNAIL: {image_path.name}",
                "THUMBNAIL_CREDIT: Image created by ChatGPT",
                "",
            ]
        ),
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    with zipfile.ZipFile(output_path) as zf:
        doc = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []
    for p in doc.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in p.findall(".//w:t", ns)).strip()
        styles = {
            r_style.get("{%s}val" % ns["w"])
            for run in p.findall("w:r", ns)
            for r_style in [run.find("w:rPr/w:rStyle", ns)]
            if r_style is not None
        }
        has_drawing = p.find(".//w:drawing", ns) is not None
        paragraphs.append((text, styles, has_drawing))

    image_idx = next(i for i, (_, _, has_drawing) in enumerate(paragraphs) if has_drawing)
    credit_text, credit_styles, _ = paragraphs[image_idx + 1]

    assert credit_text == "Image created by ChatGPT"
    assert "Annotation" in credit_styles
    assert paragraphs[image_idx + 2][0] == "After thumbnail."


def test_generate_subs_replaces_box_drawing_horizontal(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["{{TITLE_SUGGESTED}}"])
    _write_source_docx(source_docx)
    input_path.write_text(
        "TITLE_SUGGESTED: 為什麼要蓋醫院─貧中帶病拖垮家庭\n",
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    doc = Document(output_path)
    assert "為什麼要蓋醫院 - 貧中帶病拖垮家庭" in [p.text for p in doc.paragraphs]


def test_generate_subs_preserves_commas_in_title(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["{{TITLE_SUGGESTED}}"])
    _write_source_docx(source_docx)
    input_path.write_text(
        "TITLE_SUGGESTED: 大愛真健康 - 改善長者走路與平衡，溫和髖關節保養\n",
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    doc = Document(output_path)
    assert "大愛真健康 - 改善長者走路與平衡，溫和髖關節保養" in [
        p.text for p in doc.paragraphs
    ]


def test_generate_subs_preserves_source_hyperlink_and_highlight_formatting(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["字幕："])
    _write_formatted_source_docx(source_docx)
    input_path.write_text("INTRO: Intro\n", encoding="utf-8")

    with zipfile.ZipFile(source_docx) as zf:
        source_xml = etree.fromstring(zf.read("word/document.xml"))

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    with zipfile.ZipFile(output_path) as zf:
        doc = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    source_timing_highlight = None
    source_subtitle_highlight = None
    for p in source_xml.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in p.findall(".//w:t", ns))
        if "07:34-09:41 (2分7秒)" in text:
            source_timing_highlight = p.find("w:r/w:rPr/w:highlight", ns).get(
                "{%s}val" % ns["w"]
            )
        if "First line." in text:
            source_subtitle_highlight = p.find("w:r/w:rPr/w:highlight", ns).get(
                "{%s}val" % ns["w"]
            )

    url_paragraph = None
    timing_paragraph = None
    subtitle_paragraph = None
    for p in doc.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in p.findall(".//w:t", ns))
        if "https://example.com/source" in text:
            url_paragraph = p
        if "07:34-09:41 (2分7秒)" in text:
            timing_paragraph = p
        if "First line." in text:
            subtitle_paragraph = p

    assert url_paragraph is not None
    assert timing_paragraph is not None
    assert subtitle_paragraph is not None

    hyperlinks = url_paragraph.findall("w:hyperlink", ns)
    assert len(hyperlinks) == 1

    h_runs = hyperlinks[0].findall("w:r", ns)
    assert h_runs
    timing_run = timing_paragraph.find("w:r", ns)
    assert timing_run is not None
    timing_highlight = timing_run.find("w:rPr/w:highlight", ns)
    assert timing_highlight is not None
    assert timing_highlight.get("{%s}val" % ns["w"]) == source_timing_highlight

    subtitle_run = subtitle_paragraph.find("w:r", ns)
    assert subtitle_run is not None
    subtitle_highlight = subtitle_run.find("w:rPr/w:highlight", ns)
    assert subtitle_highlight is not None
    assert subtitle_highlight.get("{%s}val" % ns["w"]) == source_subtitle_highlight


def test_generate_subs_remaps_cloned_hyperlink_style_id_to_template_style(
    tmp_path: Path,
) -> None:
    template_path = Path("templates/subs_template.docx")
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_formatted_source_docx(source_docx)
    _rewrite_hyperlink_style_id(source_docx, "a7")
    input_path.write_text("INTRO: Intro\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    with zipfile.ZipFile(output_path) as zf:
        styles_xml = etree.fromstring(zf.read("word/styles.xml"))
        doc_xml = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    hyperlink_style_nodes = styles_xml.xpath(
        ".//w:style[w:name[@w:val='Hyperlink']]",
        namespaces=ns,
    )
    assert hyperlink_style_nodes
    hyperlink_style_id = hyperlink_style_nodes[0].get("{%s}styleId" % ns["w"])
    assert hyperlink_style_id

    hyperlink_style_run = None
    for paragraph in doc_xml.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in paragraph.findall(".//w:t", ns))
        if "https://example.com/source" not in text:
            continue
        hyperlink_style_run = paragraph.find(".//w:hyperlink//w:rPr/w:rStyle", ns)
        break

    assert hyperlink_style_run is not None
    assert hyperlink_style_run.get("{%s}val" % ns["w"]) == hyperlink_style_id


def test_generate_subs_remaps_field_code_hyperlink_style_id_to_template_style(
    tmp_path: Path,
) -> None:
    template_path = Path("templates/subs_template.docx")
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_formatted_source_docx(source_docx)
    _rewrite_hyperlink_style_id(source_docx, "a7")
    _rewrite_hyperlink_as_field_code(source_docx)
    input_path.write_text("INTRO: Intro\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    with zipfile.ZipFile(output_path) as zf:
        styles_xml = etree.fromstring(zf.read("word/styles.xml"))
        doc_xml = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    hyperlink_style_nodes = styles_xml.xpath(
        ".//w:style[w:name[@w:val='Hyperlink']]",
        namespaces=ns,
    )
    assert hyperlink_style_nodes
    hyperlink_style_id = hyperlink_style_nodes[0].get("{%s}styleId" % ns["w"])
    assert hyperlink_style_id

    field_paragraph = None
    for paragraph in doc_xml.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in paragraph.findall(".//w:t", ns))
        if "https://example.com/source" in text:
            field_paragraph = paragraph
            break

    assert field_paragraph is not None
    assert field_paragraph.findall(".//w:instrText", ns)

    run_styles = [
        node.get("{%s}val" % ns["w"])
        for node in field_paragraph.findall(".//w:rPr/w:rStyle", ns)
    ]
    assert hyperlink_style_id in run_styles

def test_generate_subs_layout_matches_current_output_structure(tmp_path: Path) -> None:
    template_path = Path("templates/subs_template.docx")
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_source_docx(
        source_docx,
        header_paragraphs=[
            "Sample Title",
            "https://example.com/watch?v=1",
            "",
            "Summary line A.",
            "",
            "Summary line B.",
            "",
        ],
        body_paragraphs=[
            "00:00:01:00\t00:00:02:00\t測試",
            "Test",
            "",
        ],
    )
    input_path.write_text(
        "\n".join(
            [
                "YT_TITLE_SUGGESTED: YT title",
                "TITLE_SUGGESTED: Title only",
                "INTRO:",
                "Intro EN.",
                "",
                "Intro ZH.",
                "THUMBNAIL: missing.png",
                "",
            ]
        ),
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    doc = Document(output_path)
    texts = [p.text for p in doc.paragraphs]

    # Key structure order should match current output behavior.
    idx_summary_a = texts.index("Summary line A.")
    idx_summary_b = texts.index("Summary line B.")
    idx_yt_label = texts.index("建議YT標題：")
    idx_title_label = texts.index("建議標題：")
    idx_intro_label = texts.index("簡介：")
    idx_thumb_label = texts.index("選圖：")
    idx_thumb_value = texts.index("missing.png")
    idx_subtitle_label = texts.index("字幕：")
    idx_body_src = texts.index("00:00:01:00\t00:00:02:00\t測試")

    assert idx_summary_a < idx_summary_b < idx_yt_label
    assert idx_yt_label < idx_title_label < idx_intro_label < idx_thumb_label < idx_subtitle_label
    assert idx_subtitle_label < idx_body_src
    assert idx_subtitle_label == idx_thumb_value + 2


def test_generate_subs_treats_bom_prefixed_first_subtitle_as_body(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["建議YT標題：", "{{YT_TITLE_SUGGESTED}}", "字幕："])
    _write_source_docx(
        source_docx,
        header_paragraphs=[
            "Source title",
            "Source summary.",
            "",
        ],
        body_paragraphs=[
            "\ufeff00:00:21:10\t00:00:23:24\tFirst subtitle line.",
            "Second subtitle line.",
        ],
    )
    input_path.write_text("YT_TITLE_SUGGESTED: Suggested title\n", encoding="utf-8")

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    texts = [p.text for p in Document(output_path).paragraphs]

    assert "\ufeff00:00:21:10\t00:00:23:24\tFirst subtitle line." not in texts[:5]
    label_idx = texts.index("字幕：")
    first_subtitle_idx = texts.index("\ufeff00:00:21:10\t00:00:23:24\tFirst subtitle line.")
    assert first_subtitle_idx == label_idx + 2


def test_generate_subs_intro_symbol_font_applies_only_to_icon_runs(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["{{INTRO}}"])
    _write_source_docx(source_docx)
    input_path.write_text(
        "\n".join(
            [
                "INTRO:",
                "📌 本集重點：",
                "✔ 在安全與信任的環境中",
                "",
            ]
        ),
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    with zipfile.ZipFile(output_path) as zf:
        doc = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def para_by_text(target: str):
        for p in doc.findall(".//w:p", ns):
            text = "".join(t.text or "" for t in p.findall(".//w:t", ns))
            if text == target:
                return p
        return None

    p = para_by_text("📌 本集重點：")
    assert p is not None
    runs = p.findall("w:r", ns)
    run_texts = [
        "".join(t.text or "" for t in r.findall(".//w:t", ns))
        for r in runs
    ]
    assert run_texts == ["📌", " 本集重點："]

    icon_fonts = runs[0].find("w:rPr/w:rFonts", ns)
    assert icon_fonts is not None
    assert icon_fonts.get("{%s}ascii" % ns["w"]) == "Segoe UI Symbol"
    assert icon_fonts.get("{%s}hAnsi" % ns["w"]) == "Segoe UI Symbol"
    assert icon_fonts.get("{%s}cs" % ns["w"]) == "Segoe UI Symbol"

    text_fonts = runs[1].find("w:rPr/w:rFonts", ns)
    assert text_fonts is None


def test_generate_subs_middle_dot_uses_cjk_font(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["{{INTRO}}"])
    _write_source_docx(source_docx)
    input_path.write_text(
        "\n".join(
            [
                "INTRO:",
                "示例‧姓名",
                "",
            ]
        ),
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    with zipfile.ZipFile(output_path) as zf:
        doc = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    target_run = None
    for p in doc.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in p.findall(".//w:t", ns))
        if text == "示例‧姓名":
            for r in p.findall("w:r", ns):
                run_text = "".join(t.text or "" for t in r.findall(".//w:t", ns))
                if "‧" in run_text:
                    target_run = r
                    break
            break

    assert target_run is not None
    fonts = target_run.find("w:rPr/w:rFonts", ns)
    assert fonts is not None
    assert fonts.get("{%s}ascii" % ns["w"]) == "新細明體"
    assert fonts.get("{%s}hAnsi" % ns["w"]) == "新細明體"
    assert fonts.get("{%s}cs" % ns["w"]) == "新細明體"


def test_with_subs_output_suffix_appends_al_once() -> None:
    base = Path("output/sample.docx")
    already = Path("output/sample_al.docx")

    assert generate_subs.with_subs_output_suffix(base) == Path("output/sample_al.docx")
    assert generate_subs.with_subs_output_suffix(already) == already


def test_generate_subs_uses_body_from_input_when_present(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["字幕："])
    _write_source_docx(
        source_docx,
        body_paragraphs=[
            "00:00:00:00\t00:00:02:00\tSource body line.",
            "Source translation.",
            "",
        ],
    )
    input_path.write_text(
        "\n".join(
            [
                "BODY:",
                "00:01:00:00\t00:01:02:00\tInput body line.",
                "Input translation.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    texts = [p.text for p in Document(output_path).paragraphs]

    assert "00:01:00:00\t00:01:02:00\tInput body line." in texts
    assert "Input translation." in texts
    assert "00:00:00:00\t00:00:02:00\tSource body line." not in texts
    assert "Source translation." not in texts


def test_generate_subs_keeps_blank_line_after_subtitle_label_with_input_body(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["字幕："])
    _write_source_docx(source_docx)
    input_path.write_text(
        "\n".join(
            [
                "BODY:",
                "00:01:00:00\t00:01:02:00\tInput body line.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)
    doc = Document(output_path)
    texts = [p.text for p in doc.paragraphs]
    label_idx = texts.index("字幕：")
    assert texts[label_idx + 1] == ""
    assert texts[label_idx + 2] == "00:01:00:00\t00:01:02:00\tInput body line."


def test_generate_subs_treats_xxx_prefixed_timecode_as_subtitle_line(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    source_docx = tmp_path / "source.docx"
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.docx"

    _write_docx(template_path, ["字幕："])
    _write_source_docx(source_docx)
    input_path.write_text(
        "\n".join(
            [
                "BODY:",
                "https://en.wikibooks.org/wiki/Traditional_Chinese_Medicine/Prescriptions",
                "*Five-Juice Drink*",
                "*五汁饮*",
                "",
                "XXX\t00:08:29:00\t00:08:43:11\t像梨子汁啊 然後荸薺啊",
                "Juices made from pear and water chestnut.",
                "XXX",
                "and asparagus can all help with hydration.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    generate_subs.generate_subs(template_path, source_docx, input_path, output_path)

    with zipfile.ZipFile(output_path) as zf:
        doc = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    timing_paragraph = None
    url_paragraph = None
    for paragraph in doc.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in paragraph.findall(".//w:t", ns))
        if text.startswith("XXX00:08:29:0000:08:43:11"):
            timing_paragraph = paragraph
        if text == "https://en.wikibooks.org/wiki/Traditional_Chinese_Medicine/Prescriptions":
            url_paragraph = paragraph

    assert timing_paragraph is not None
    assert url_paragraph is not None
    assert url_paragraph.findall("w:hyperlink", ns)
    assert timing_paragraph.find("w:r/w:rPr/w:highlight", ns) is None
