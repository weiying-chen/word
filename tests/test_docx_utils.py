from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.oxml.ns import qn

from docx_utils import add_hyperlink, apply_font_size_to_runs, apply_highlight_to_runs


def test_apply_highlight_to_runs_does_not_change_font_size() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Label")
    run.font.size = Pt(14)

    apply_highlight_to_runs(paragraph, highlight_color=WD_COLOR_INDEX.YELLOW)

    assert paragraph.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW
    assert paragraph.runs[0].font.size == Pt(14)


def test_apply_font_size_to_runs_changes_font_size() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Label")

    apply_font_size_to_runs(paragraph, font_size_pt=10)

    assert run.font.size == Pt(10)


def test_highlighted_hyperlink_uses_valid_run_property_order() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()

    add_hyperlink(
        paragraph,
        "https://example.com/",
        "https://example.com/",
        highlight=True,
        highlight_color="brightGreen",
    )

    run_properties = paragraph._p.find(".//" + qn("w:rPr"))
    assert run_properties is not None
    assert [child.tag for child in run_properties] == [
        qn("w:rStyle"),
        qn("w:color"),
        qn("w:sz"),
        qn("w:highlight"),
        qn("w:u"),
    ]
