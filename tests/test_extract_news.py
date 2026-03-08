from pathlib import Path

from docx import Document

from docx_utils import add_hyperlink
import extract_news


def _write_docx_with_hyperlink(path: Path, title: str, url: str, summary: str, body_lines: list[str]) -> None:
    doc = Document()
    doc.add_paragraph(title)
    p = doc.add_paragraph('Video link')
    add_hyperlink(p, 'Open', url)
    doc.add_paragraph(summary)
    for line in body_lines:
        doc.add_paragraph(line)
    doc.save(path)


def test_extract_news_writes_news_native_sections(tmp_path: Path) -> None:
    input_docx = tmp_path / 'input.docx'
    source_docx = tmp_path / 'source.docx'
    source_txt = tmp_path / 'source.txt'

    body_lines = [
        '<',
        '(6． Uyanda烏漾達)',
        '/*SUPER:',
        '慈濟志工｜烏漾達//',
        '我一點都不累//',
        '*/',
        '正文第一行。',
    ]

    _write_docx_with_hyperlink(
        input_docx,
        '南非校園品格教育 心田開出美善花朵｜大愛新聞 #南非德本',
        'https://youtu.be/example123',
        '摘要段落。',
        body_lines,
    )

    extract_news.extract_news(input_docx, source_docx, source_txt)

    text = source_txt.read_text(encoding='utf-8')
    assert 'TITLE_TEXT: 南非校園品格教育 心田開出美善花朵' in text
    assert 'TITLE_URL: https://youtu.be/example123' in text
    assert 'SUMMARY:\n摘要段落。' in text
    assert 'SUPER_PEOPLE:\n慈濟志工 | 烏漾達\nUyanda' in text
    assert 'BODY:\n<' in text
    assert '正文第一行。' in text


def test_extract_news_preserves_input_docx_as_source_docx(tmp_path: Path) -> None:
    input_docx = tmp_path / 'input.docx'
    source_docx = tmp_path / 'source.docx'
    source_txt = tmp_path / 'source.txt'

    _write_docx_with_hyperlink(
        input_docx,
        'Sample title',
        'https://example.com/x',
        'Summary line',
        ['<', 'Body line'],
    )

    extract_news.extract_news(input_docx, source_docx, source_txt)

    assert source_docx.read_bytes() == input_docx.read_bytes()
