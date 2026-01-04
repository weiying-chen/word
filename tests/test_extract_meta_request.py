import tempfile
import unittest
from pathlib import Path

from docx import Document

from extract_meta_request import extract_from_docx


class ExtractMetaRequestTests(unittest.TestCase):
    def _build_docx(self, path: Path) -> None:
        doc = Document()
        doc.add_paragraph("奧莫克眼科義診 重啟光明安生安心")
        doc.add_paragraph("1_0001")
        doc.add_paragraph("(NS)")
        doc.add_paragraph("海燕颱風重創菲律賓後，慈濟在奧莫克援建大愛村。")
        doc.add_paragraph("篩檢186人，其中超過60例白內障與25例有結膜異常增生問題。")
        doc.add_paragraph("/*SUPER:")
        doc.add_paragraph("病患│羅伯托//")
        doc.add_paragraph("我的左眼看不見 也要接受治療了//")
        doc.add_paragraph("*/")
        doc.add_paragraph("/*SUPER:")
        doc.add_paragraph("醫師│毛伊//")
        doc.add_paragraph("未來 如果慈濟繼續幫助我們//")
        doc.add_paragraph("*/")
        doc.add_paragraph("( 13 Roberto )")
        doc.add_paragraph("( 14 Dr. Maui )")
        doc.save(str(path))

    def test_extracts_title_people_facts(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "main.docx"
            self._build_docx(docx_path)
            payload = extract_from_docx(docx_path)

        self.assertEqual(payload["title_zh"], "奧莫克眼科義診 重啟光明安生安心")
        self.assertEqual(len(payload["people"]), 2)
        self.assertEqual(payload["people"][0]["name_zh"], "羅伯托")
        self.assertEqual(payload["people"][0]["role_zh"], "病患")
        self.assertEqual(payload["people"][1]["name_zh"], "毛伊")
        self.assertEqual(payload["people"][1]["role_zh"], "醫師")
        self.assertEqual(payload["people"][0]["name_en"], "Roberto")
        self.assertEqual(payload["people"][1]["name_en"], "Dr. Maui")
        self.assertEqual(
            payload["summary_zh"],
            "海燕颱風重創菲律賓後，慈濟在奧莫克援建大愛村。",
        )


if __name__ == "__main__":
    unittest.main()
