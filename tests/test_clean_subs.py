from pathlib import Path

from docx import Document

import clean_subs


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def test_remove_sources_blocks_only_in_intro_and_subs_sections(tmp_path: Path) -> None:
    source_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _write_docx(
        source_path,
        [
            "https://www.youtube.com/watch?v=top-header-link",
            "Top header title",
            "簡介：",
            "https://en.wikibooks.org/wiki/Traditional_Chinese_Medicine/Prescriptions",
            "*Five-Juice Drink*",
            "*五汁饮*",
            "Intro final line",
            "",
            "字幕：",
            "00:08:25:09\t00:08:28:02\t就是像我們中醫常常會有一個",
            "00:08:28:02\t00:08:29:00\t五汁飲",
            "https://health.businessweekly.com.tw/article/ARTL003018041",
            "Source note A",
            "Source note B",
            "",
            "XXX\t00:08:29:00\t00:08:43:11\t像梨子汁啊",
            "Juices made from pear and water chestnut.",
        ],
    )

    clean_subs.remove_sources_from_docx(source_path, output_path)

    texts = [p.text for p in Document(output_path).paragraphs]
    assert "https://www.youtube.com/watch?v=top-header-link" in texts
    assert "https://en.wikibooks.org/wiki/Traditional_Chinese_Medicine/Prescriptions" not in texts
    assert "https://health.businessweekly.com.tw/article/ARTL003018041" not in texts
    assert "*Five-Juice Drink*" not in texts
    assert "*五汁饮*" not in texts
    assert "Source note A" not in texts
    assert "Source note B" not in texts
    assert "Intro final line" not in texts
    assert "XXX\t00:08:29:00\t00:08:43:11\t像梨子汁啊" in texts
    assert "Juices made from pear and water chestnut." in texts


def test_remove_blank_lines_between_subtitle_timestamps(tmp_path: Path) -> None:
    source_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _write_docx(
        source_path,
        [
            "字幕：",
            "00:00:01:00\t00:00:02:00\t第一句",
            "",
            "00:00:02:00\t00:00:03:00\t第二句",
            "",
            "Normal line",
        ],
    )

    clean_subs.remove_sources_from_docx(source_path, output_path)

    texts = [p.text for p in Document(output_path).paragraphs]
    assert texts == [
        "字幕：",
        "00:00:01:00\t00:00:02:00\t第一句",
        "00:00:02:00\t00:00:03:00\t第二句",
        "",
        "Normal line",
    ]
