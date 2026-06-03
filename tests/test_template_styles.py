from pathlib import Path

from docx import Document

import template_styles


def test_ensure_base_styles_sets_normal_and_shared_character_styles(tmp_path: Path) -> None:
    path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("x")
    doc.save(path)

    doc = Document(path)
    template_styles.ensure_base_styles(doc)
    doc.save(path)

    out = Document(path)
    assert out.styles["Normal"].font.size.pt == 12
    assert out.styles["SectionLabelSmall"].font.size.pt == 10
    assert str(out.styles["SectionLabelSmall"].font.color.rgb) == "0070C0"
    assert out.styles["SectionLabelLarge"].font.size.pt == 12
    assert str(out.styles["SectionLabelLarge"].font.color.rgb) == "0070C0"
    assert out.styles["ReferenceLink"].font.size.pt == 10
    assert str(out.styles["ReferenceLink"].font.color.rgb) == "0563C1"


def test_sync_review_template_styles_applies_blue_style_to_section_labels(tmp_path: Path) -> None:
    path = tmp_path / "review_template.docx"
    doc = Document()
    table = doc.add_table(rows=14, cols=4)
    table.cell(12, 0).text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    table.cell(13, 0).text = "其他工作:"
    doc.save(path)

    doc = Document(path)
    template_styles.ensure_base_styles(doc)
    template_styles.sync_review_template_styles(doc)
    doc.save(path)

    out = Document(path)
    run_a = out.tables[0].cell(12, 0).paragraphs[0].runs[0]
    run_b = out.tables[0].cell(13, 0).paragraphs[0].runs[0]
    assert run_a.style.name == "SectionLabelSmall"
    assert run_b.style.name == "SectionLabelSmall"


def test_sync_review_template_styles_skips_tables_without_expected_grid(tmp_path: Path) -> None:
    path = tmp_path / "review_template.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1)
    doc.save(path)

    doc = Document(path)
    template_styles.ensure_base_styles(doc)
    template_styles.sync_review_template_styles(doc)
