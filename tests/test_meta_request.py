import tempfile
import unittest
from pathlib import Path

from docx import Document

from meta.extract_meta_request import extract_from_docx


class ExtractMetaRequestTests(unittest.TestCase):
    def _build_docx(self, path: Path) -> None:
        doc = Document()
        doc.add_paragraph("沿海義診守護居民健康")
        doc.add_paragraph("1_0001")
        doc.add_paragraph("(NS)")
        doc.add_paragraph("志工與醫護團隊在沿海小鎮舉辦義診，提供居民所需協助。")
        doc.add_paragraph("兩天內完成多項檢查，協助居民安排後續治療。")
        doc.add_paragraph("/*SUPER:")
        doc.add_paragraph("居民│陳先生//")
        doc.add_paragraph("我的左眼幾乎看不見了//")
        doc.add_paragraph("*/")
        doc.add_paragraph("/*SUPER:")
        doc.add_paragraph("醫師│林醫師//")
        doc.add_paragraph("未來 如果團隊持續提供協助//")
        doc.add_paragraph("*/")
        doc.add_paragraph("( 13 Mr. Chen )")
        doc.add_paragraph("( 14 Dr. Lin )")
        doc.save(str(path))

    def test_extracts_title_people_facts(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "main.docx"
            self._build_docx(docx_path)
            payload = extract_from_docx(docx_path)

        self.assertEqual(payload["title_zh"], "沿海義診守護居民健康")
        self.assertEqual(len(payload["people"]), 2)
        self.assertEqual(payload["people"][0]["name_zh"], "陳先生")
        self.assertEqual(payload["people"][0]["role_zh"], "居民")
        self.assertEqual(payload["people"][1]["name_zh"], "林醫師")
        self.assertEqual(payload["people"][1]["role_zh"], "醫師")
        self.assertEqual(payload["people"][0]["name_en"], "Mr. Chen")
        self.assertEqual(payload["people"][1]["name_en"], "Dr. Lin")
        self.assertEqual(
            payload["summary_zh"],
            "志工與醫護團隊在沿海小鎮舉辦義診，提供居民所需協助。",
        )


if __name__ == "__main__":
    unittest.main()
