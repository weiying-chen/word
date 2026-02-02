from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from datetime import date

from generate_posts import (
    extract_post_entries,
    extract_post_titles,
    normalize_title,
    _build_hashtags,
    build_hashtags_from_title_line,
)


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def _add_hyperlink_paragraph(doc: Document, display_text: str, url: str) -> None:
    paragraph = doc.add_paragraph("")
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = display_text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_normalize_title_strips_translator_tag() -> None:
    title = "大愛醫生館 - 怎麼坐才算有“坐相”？st/rc"
    assert normalize_title(title) == "大愛醫生館 怎麼坐才算有“坐相”？"


def test_extract_post_titles_from_schedule(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    _write_docx(
        schedule_path,
        [
            "節目6則",
            "1. elijah",
            "節目甲 - 測試標題 em/el",
            "https://example.com/1",
            "2. alex",
            "大愛醫生館 - 怎麼坐才算有“坐相”？st/rc",
            "https://example.com/2",
            "搭配",
            "https://example.com/news",
            "新聞標題",
            "3. alex",
            "大愛真健康 - 5分鐘高效有氧 | 上下肢肌耐力 | 肩腿| 背腿 | 核心 nick/cc",
            "https://example.com/3",
            "--------------------------------",
            "FB小編文box裡面的新聞已用掉下面5則",
        ],
    )

    titles = extract_post_titles(schedule_path)
    assert titles == [
        "大愛醫生館 怎麼坐才算有“坐相”",
        "大愛真健康 5分鐘高效有氧 上下肢肌耐力 肩腿 背腿 核心",
    ]


def test_extract_post_titles_from_alex_blocks(tmp_path: Path) -> None:
    schedule_path = tmp_path / "alex_blocks.docx"
    _write_docx(
        schedule_path,
        [
            "1",
            "參考資料:",
            "https://example.com/news",
            "26/1/23",
            "新聞標題",
            "要用的影片:",
            "https://example.com/video",
            "Program - Test Title (大愛醫生館 - 中文標題)",
            "English prompt line",
            "中文提示",
        ],
    )

    titles = extract_post_titles(schedule_path)
    assert titles == ["大愛醫生館 中文標題"]


def test_blocks_date_prefix_and_multiline_ref(tmp_path: Path) -> None:
    schedule_path = tmp_path / "blocks.docx"
    _write_docx(
        schedule_path,
        [
            "1",
            "參考資料:",
            "https://example.com/news",
            "1/27",
            "News title",
            "Extra line",
            "要用的影片:",
            "https://example.com/video",
            "Program - Test Title (大愛醫生館 - 中文標題)",
        ],
    )

    entries = extract_post_entries(schedule_path)
    assert entries[0]["ref_title"] == "News title\nExtra line"
    expected_prefix = f"{date.today().year % 100:02d}0127_"
    assert entries[0]["filename_prefix_override"] == expected_prefix


def test_bodhi_date_prefix(tmp_path: Path) -> None:
    schedule_path = tmp_path / "bodhi.docx"
    _write_docx(
        schedule_path,
        [
            "菩提7則",
            "1. alex",
            "1/20首播 互愛共善造大福",
            "https://www.daai.tv/master/life-wisdom/P90230145",
            "--------------------------------",
        ],
    )

    entries = extract_post_entries(schedule_path)
    assert len(entries) == 1
    entry = entries[0]
    assert entry["video_title"] == "人間菩提 (1/20首播 互愛共善造大福)"
    assert entry["ref_url"] == entry["video_url"]
    assert entry["ref_title"] == ""
    expected_prefix = f"{date.today().year % 100:02d}0120_"
    assert entry["filename_prefix_override"] == expected_prefix


def test_bodhi_section_does_not_leak(tmp_path: Path) -> None:
    schedule_path = tmp_path / "mixed.docx"
    _write_docx(
        schedule_path,
        [
            "菩提7則",
            "1. alex",
            "1/20首播 互愛共善造大福",
            "https://example.com/bodhi",
            "節目1則",
            "1. alex",
            "Normal Show - Title",
            "https://example.com/normal",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )

    entries = extract_post_entries(schedule_path)
    assert len(entries) == 2
    assert entries[0].get("reference_only")
    assert not entries[1].get("reference_only")


def test_extracts_full_url_from_truncated_hyperlink(tmp_path: Path) -> None:
    schedule_path = tmp_path / "alex_blocks_hyperlink.docx"
    doc = Document()
    doc.add_paragraph("1")
    doc.add_paragraph("參考資料:")
    _add_hyperlink_paragraph(
        doc,
        "https://tw.news.yahoo.com/%E8%B5%B0%E8%B7%AF-%E6%9C%80%E8%A2%AB%E4%BD%8E%E4%BC%B0%E7%9A%84%E9%95%B7…",
        "https://tw.news.yahoo.com/%E8%B5%B0%E8%B7%AF-%E6%9C%80%E8%A2%AB%E4%BD%8E%E4%BC%B0%E7%9A%84%E9%95%B7%E5%A3%BD%E9%81%8B%E5%8B%95-073646342.html",
    )
    doc.add_paragraph("25/12/29")
    doc.add_paragraph("新聞標題")
    doc.add_paragraph("要用的影片:")
    doc.add_paragraph("https://example.com/video")
    doc.add_paragraph("Program - Test Title (大愛醫生館 - 測試標題)")
    doc.save(schedule_path)

    entries = extract_post_entries(schedule_path)
    assert entries[0]["ref_url"].endswith("…")
    assert entries[0]["ref_url_target"].endswith("073646342.html")


def test_build_hashtags_strips_quotes_and_spaces() -> None:
    program = "大愛醫生館"
    title = "怎麼坐才算有“坐相”？"
    assert _build_hashtags(program, title, pascal_case=False) == "#大愛醫生館 #怎麼坐才算有坐相"
    assert _build_hashtags(program, title, pascal_case=True) == "#大愛醫生館 #怎麼坐才算有坐相"


def test_build_hashtags_splits_english_and_chinese_parenthetical() -> None:
    title = "Learn Something New Every Day - Every Step Matters (日日有新知 - 走越多 活越久)"
    hashtags_en, hashtags_zh = build_hashtags_from_title_line(title)
    assert hashtags_en == "#LearnSomethingNewEveryDay #EveryStepMatters"
    assert hashtags_zh == "#日日有新知 #走越多活越久"
