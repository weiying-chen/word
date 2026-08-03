import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import generate_drama


def _record(
    record_id: str,
    order: int,
    record_type: str,
    translation: str,
    **extra,
) -> dict:
    return {
        "id": record_id,
        "order": order,
        "type": record_type,
        "source_text": f"source {order}",
        "text_zh": f"source {order}",
        "translation_en": translation,
        "layout": {"page": 1},
        "review": {"status": "reviewed", "flags": []},
        **extra,
    }


def _payload() -> dict:
    return {
        "schema_version": "2.0",
        "document_type": "script_translation_working_file",
        "metadata": {
            "source_pdf": "source.pdf",
            "source_sha256": hashlib.sha256(b"source pdf").hexdigest(),
            "page_count": 1,
        },
        "translation_config": {},
        "statistics": {
            "scene_count": 2,
            "element_count": 8,
            "front_matter_count": 1,
            "first_scene": 1,
            "last_scene": 2,
        },
        "front_matter": [
            _record("title-1", 1, "title_page", "WORKING TITLE"),
        ],
        "scenes": [
            {
                "scene_number": 1,
                "heading_id": "heading-1",
                "heading_order": 2,
                "heading_zh": "1. 內景",
                "heading_en": "SCENE 1. INT. ROOM – DAY",
                "source_text": "1. 內景",
                "layout": {"page": 1},
                "review": {"status": "reviewed", "flags": []},
                "elements": [
                    _record("action-1", 3, "action", "A door opens."),
                    _record(
                        "dialogue-1",
                        4,
                        "dialogue",
                        "Hello.",
                        speaker_zh="甲",
                        speaker_en="PERSON A",
                        parenthetical_zh="微笑",
                        parenthetical_en="smiling",
                        qualifier="V.O.",
                    ),
                    _record(
                        "note-1",
                        5,
                        "translation_note",
                        "(TRANSLATION) A translated sign.",
                    ),
                ],
            },
            {
                "scene_number": 2,
                "heading_id": "heading-2",
                "heading_order": 6,
                "heading_zh": "2. 外景",
                "heading_en": "",
                "source_text": "2. 外景",
                "layout": {"page": 1},
                "review": {"status": "unreviewed", "flags": []},
                "elements": [
                    _record("action-2", 7, "action", ""),
                    _record("production-1", 8, "production_note", ""),
                ],
            },
        ],
    }


def _write_reference(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.6)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(14)
    heading.paragraph_format.line_spacing = 1
    run = heading.add_run("SCENE 1. REFERENCE – DAY")
    run.font.name = "Times New Roman"
    run.font.size = Pt(13.5)
    run.bold = True

    action = doc.add_paragraph()
    action.paragraph_format.space_before = Pt(14)
    action.paragraph_format.space_after = Pt(14)
    action.paragraph_format.line_spacing = 1
    run = action.add_run("Reference action.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True

    dialogue = doc.add_paragraph()
    speaker = dialogue.add_run("PERSON A")
    speaker.font.name = "Times New Roman"
    speaker.font.size = Pt(12)
    speaker.bold = True
    body = dialogue.add_run("\nReference dialogue.")
    body.font.name = "Times New Roman"
    body.font.size = Pt(12)
    doc.save(path)


def test_preview_generates_translated_records_in_source_order(tmp_path: Path) -> None:
    input_path = tmp_path / "drama.json"
    reference_path = tmp_path / "reference.docx"
    output_path = tmp_path / "preview.docx"
    input_path.write_text(json.dumps(_payload()), encoding="utf-8")
    _write_reference(reference_path)

    result = generate_drama.generate_drama(
        input_path,
        reference_path,
        output_path,
        preview=True,
    )

    assert result.generated_ids == [
        "title-1",
        "heading-1",
        "action-1",
        "dialogue-1",
        "note-1",
    ]
    assert result.missing_count == 3

    doc = Document(output_path)
    assert doc.paragraphs[0].text.startswith("PREVIEW")
    heading = next(p for p in doc.paragraphs if p.text.startswith("SCENE 1."))
    assert heading.runs[0].bold is True
    assert heading.runs[0].font.name == "Calibri"
    assert heading.runs[0].font.size == Pt(12)
    assert heading.paragraph_format.space_before is None
    assert heading.paragraph_format.space_after is None
    assert heading.paragraph_format.line_spacing is None

    action = next(p for p in doc.paragraphs if p.text == "A door opens.")
    assert action.runs[0].italic is True
    assert action.runs[0].font.name == "Calibri"
    assert action.runs[0].font.size == Pt(12)

    dialogue = next(p for p in doc.paragraphs if "Hello." in p.text)
    assert dialogue.text == "PERSON A (V.O.)\n(smiling)\nHello."
    assert dialogue.runs[0].bold is True
    assert dialogue.runs[-1].bold is False

    note = next(p for p in doc.paragraphs if "(TRANSLATION)" in p.text)
    assert note.runs[0].bold is True
    assert doc.sections[0].page_width == Inches(8.5)
    assert doc.sections[0].left_margin == Inches(1.25)
    assert doc.sections[0].header_distance == Inches(0.4)
    assert doc.sections[0].footer_distance == Inches(0.6)

    with ZipFile(output_path) as package:
        assert package.testzip() is None
    assert generate_drama.read_document_mapping(output_path) == result.generated_ids


def test_final_rejects_missing_english_and_unresolved_flags(tmp_path: Path) -> None:
    input_path = tmp_path / "drama.json"
    reference_path = tmp_path / "reference.docx"
    output_path = tmp_path / "final.docx"
    payload = _payload()
    payload["scenes"][0]["elements"][0]["review"]["flags"] = ["needs_review"]
    input_path.write_text(json.dumps(payload), encoding="utf-8")
    _write_reference(reference_path)

    with pytest.raises(generate_drama.DramaValidationError) as exc_info:
        generate_drama.generate_drama(
            input_path,
            reference_path,
            output_path,
            preview=False,
        )

    message = str(exc_info.value)
    assert "heading_en" in message
    assert "translation_en" in message
    assert "needs_review" in message
    assert not output_path.exists()


def test_final_allows_confirmed_visual_merge_metadata() -> None:
    payload = _payload()
    payload["scenes"][1]["heading_en"] = "SCENE 2. EXT. STREET – DAY"
    payload["scenes"][1]["elements"][0]["translation_en"] = "A car passes."
    payload["scenes"][1]["elements"][1]["translation_en"] = "END"
    payload["scenes"][0]["elements"][0]["review"]["flags"] = [
        "merged_visual_continuation"
    ]

    assert generate_drama.validate_payload(payload, final=True) == []


def test_preview_skips_dialogue_with_missing_required_speaker(tmp_path: Path) -> None:
    input_path = tmp_path / "drama.json"
    reference_path = tmp_path / "reference.docx"
    output_path = tmp_path / "preview.docx"
    payload = _payload()
    payload["scenes"][0]["elements"][1]["speaker_en"] = ""
    input_path.write_text(json.dumps(payload), encoding="utf-8")
    _write_reference(reference_path)

    result = generate_drama.generate_drama(
        input_path,
        reference_path,
        output_path,
        preview=True,
    )

    assert "dialogue-1" not in result.generated_ids
    assert "dialogue-1.speaker_en" in result.missing_fields
    assert "Hello." not in "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )


@pytest.mark.parametrize(
    ("mutate", "expected"),
    [
        (lambda data: data.update(schema_version="1.0"), "Schema Version 2.0"),
        (
            lambda data: data["scenes"][1].update(scene_number=3),
            "scene numbering",
        ),
        (
            lambda data: data["scenes"][0]["elements"][0].update(id="title-1"),
            "duplicate ID",
        ),
        (
            lambda data: data["scenes"][0]["elements"][0].update(order=20),
            "global order",
        ),
        (
            lambda data: data["scenes"][0].update(
                heading_en="SCENE 9. INT. ROOM – DAY"
            ),
            "heading scene number",
        ),
    ],
)
def test_schema_validation_rejects_structural_errors(mutate, expected) -> None:
    payload = _payload()
    mutate(payload)

    with pytest.raises(generate_drama.DramaValidationError, match=expected):
        generate_drama.validate_payload(payload, final=False)


def test_source_hash_verification_uses_metadata_hash(tmp_path: Path) -> None:
    source_pdf = tmp_path / "source.pdf"
    source_pdf.write_bytes(b"source pdf")
    payload = _payload()

    generate_drama.verify_source_pdf(payload, source_pdf)

    source_pdf.write_bytes(b"changed")
    with pytest.raises(generate_drama.DramaValidationError, match="SHA-256"):
        generate_drama.verify_source_pdf(payload, source_pdf)


def test_default_output_path_uses_english_suffix() -> None:
    input_path = Path("/work/靜曦拍攝本V.4.2.translation.json")

    final_path = generate_drama.resolve_output_path(input_path, None, preview=False)
    preview_path = generate_drama.resolve_output_path(input_path, None, preview=True)

    assert final_path == Path("/work/靜曦拍攝本V.4.2_英文.docx")
    assert preview_path == Path("/work/靜曦拍攝本V.4.2_英文.preview.docx")


def test_presentation_types_distinguish_actions_from_on_screen_text() -> None:
    records = [
        ("heading-1", "scene_heading", {"heading_en": "SCENE 1. ROOM"}),
        ("action-1", "action", {"translation_en": "A person enters."}),
        ("delete-1", "action", {"translation_en": "DELETE."}),
        ("cards-1", "production_note", {"translation_en": "TITLE CARDS:"}),
        ("card-1", "action", {"translation_en": "First on-screen card."}),
        ("card-2", "action", {"translation_en": "Second on-screen card."}),
        ("heading-2", "scene_heading", {"heading_en": "SCENE 2. STREET"}),
        (
            "note-1",
            "production_note",
            {"translation_en": "INS. A sentence continues"},
        ),
        ("continuation-1", "action", {"translation_en": "on the next line."}),
    ]

    assert generate_drama.classify_presentation_types(records) == [
        "scene_heading",
        "action",
        "production_directive",
        "production_note",
        "title_card_text",
        "title_card_text",
        "scene_heading",
        "production_note",
        "note_continuation",
    ]


def test_note_continuation_renders_as_one_paragraph(tmp_path: Path) -> None:
    payload = _payload()
    scene = payload["scenes"][0]
    scene["elements"] = [
        _record(
            "note-1",
            3,
            "production_note",
            "INS. The sentence continues",
        ),
        _record("continuation-1", 4, "action", "on the next visual line."),
    ]
    payload["scenes"] = [scene]
    payload["statistics"].update(
        {
            "scene_count": 1,
            "last_scene": 1,
            "element_count": 4,
        }
    )
    input_path = tmp_path / "drama.json"
    reference_path = tmp_path / "reference.docx"
    output_path = tmp_path / "output.docx"
    input_path.write_text(json.dumps(payload), encoding="utf-8")
    _write_reference(reference_path)

    result = generate_drama.generate_drama(
        input_path,
        reference_path,
        output_path,
        preview=True,
    )

    doc = Document(output_path)
    matches = [
        p for p in doc.paragraphs if p.text.startswith("INS. The sentence")
    ]
    assert len(matches) == 1
    assert matches[0].text == "INS. The sentence continues on the next visual line."
    assert all(run.bold is True for run in matches[0].runs if run.text)
    assert result.generated_ids[-2:] == ["note-1", "continuation-1"]


def test_title_page_keeps_only_department_title_notice_visible(tmp_path: Path) -> None:
    payload = _payload()
    payload["front_matter"][0]["translation_en"] = (
        "[OFFICIAL ENGLISH TITLE TO BE DETERMINED BY THE DEPARTMENT]"
    )
    payload["scenes"][0]["elements"][-1]["translation_en"] = (
        "TITLE: [OFFICIAL ENGLISH TITLE TO BE DETERMINED]"
    )
    payload["front_matter"].extend(
        [
            _record(
                "title-2",
                2,
                "title_page",
                "SHOOTING SCRIPT V.4\nWritten by Someone\nProduced by Company",
            ),
            _record("title-3", 3, "title_page", "Date: 2025/11/9"),
        ]
    )
    for scene in payload["scenes"]:
        scene["heading_order"] += 2
        for element in scene["elements"]:
            element["order"] += 2
    payload["statistics"]["front_matter_count"] = 3
    payload["statistics"]["element_count"] = 10
    input_path = tmp_path / "drama.json"
    reference_path = tmp_path / "reference.docx"
    output_path = tmp_path / "output.docx"
    input_path.write_text(json.dumps(payload), encoding="utf-8")
    _write_reference(reference_path)

    result = generate_drama.generate_drama(
        input_path,
        reference_path,
        output_path,
        preview=True,
    )

    doc = Document(output_path)
    visible_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "【正式英文片名由部門決定】" in visible_text
    assert "片名：【正式英文片名由部門決定】" in visible_text
    assert "OFFICIAL ENGLISH TITLE TO BE DETERMINED" not in visible_text
    assert "SHOOTING SCRIPT V.4" not in visible_text
    assert "Written by Someone" not in visible_text
    assert "Produced by Company" not in visible_text
    assert "Date: 2025/11/9" not in visible_text
    assert generate_drama.read_document_mapping(output_path) == result.generated_ids
    assert result.generated_ids[:4] == [
        "title-1",
        "title-2",
        "title-3",
        "heading-1",
    ]


def test_cli_preview_reports_incomplete_translation(
    tmp_path: Path, capsys
) -> None:
    input_path = tmp_path / "drama.json"
    reference_path = tmp_path / "reference.docx"
    output_path = tmp_path / "preview.docx"
    input_path.write_text(json.dumps(_payload()), encoding="utf-8")
    _write_reference(reference_path)

    exit_code = generate_drama.main(
        [
            "--input",
            str(input_path),
            "--reference",
            str(reference_path),
            "--output",
            str(output_path),
            "--preview",
        ]
    )

    output = capsys.readouterr().out
    assert exit_code == 0
    assert "PREVIEW" in output
    assert "3 required English fields remain untranslated" in output
    assert str(output_path) in output
