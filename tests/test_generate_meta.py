import tempfile
import unittest
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from generate_meta import (
    build_people_lines,
    default_output_path,
    generate_meta,
    parse_input,
    resolve_template_path,
)
from style_tokens import BODY_TEXT_SIZE_PT


class RenderMetaTests(unittest.TestCase):
    def _build_template(self, path: Path) -> None:
        doc = Document()
        doc.add_paragraph("重點標")
        doc.add_paragraph("{{TITLE_EN}}")
        doc.add_paragraph("名字職銜")
        doc.add_paragraph("")
        doc.add_paragraph("{{PEOPLE}}")
        doc.add_paragraph("")
        doc.add_paragraph("YT簡介")
        doc.add_paragraph("{{OVERVIEW_EN}}")
        doc.save(str(path))

    def test_renders_title_people_overview(self) -> None:
        source_text = "\n".join(
            [
                "TITLE_TEXT: 中文標題",
                "SUMMARY:",
                "中文摘要",
                "",
                "TITLE: English Title",
                "OVERVIEW: English overview.",
                "",
                "BODY:",
                "(  13   Alice )",
                "/*SUPER:",
                "病患│甲//",
                "引言一//",
                "*/",
                "Quote one.",
                "",
                "(  14   Bob )",
                "/*SUPER:",
                "醫師│乙//",
                "引言二//",
                "*/",
                "Quote two.",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            payload_path = tmpdir_path / "news_input.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            payload_path.write_text(source_text, encoding="utf-8")

            generate_meta(template_path, payload_path, output_path)

            doc = Document(str(output_path))
            texts = [p.text for p in doc.paragraphs]

        self.assertEqual(
            texts,
            [
                "重點標",
                "",
                "English Title",
                "名字職銜",
                "",
                "",
                "病患｜甲",
                "Alice",
                "{{病患}}",
                "",
                "醫師｜乙",
                "Bob",
                "{{醫師}}",
                "",
                "YT簡介",
                "",
                "English overview.",
            ],
        )

    def test_generate_meta_raises_when_title_and_overview_missing(self) -> None:
        source_text = "\n".join(
            [
                "BODY:",
                "(  13   Alice )",
                "/*SUPER:",
                "病患│甲//",
                "引言一//",
                "*/",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            payload_path = tmpdir_path / "news_input.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            payload_path.write_text(source_text, encoding="utf-8")

            with self.assertRaisesRegex(
                ValueError,
                r"\[error\] Missing required field: TITLE\n\[error\] Missing required field: OVERVIEW",
            ):
                generate_meta(template_path, payload_path, output_path)

    def test_parse_input_extracts_meta_fields_and_people(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "TITLE_TEXT: 中文標題",
                        "SUMMARY:",
                        "中文摘要",
                        "",
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                        "BODY:",
                        "(  13   Alice )",
                        "/*SUPER:",
                        "病患│甲//",
                        "引言一//",
                        "*/",
                        "Quote one.",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(data["title_zh"], "中文標題")
        self.assertEqual(data["summary_zh"], "中文摘要")
        self.assertEqual(data["title_en"], "English Title")
        self.assertEqual(data["overview_en"], "English overview.")
        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "甲",
                    "name_en": "Alice",
                    "role_zh": "病患",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_extracts_english_name_from_generic_super_cues(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                        "(22秒，Alois L. Sikuka)",
                        "/*SUPER:",
                        "國會議員│阿洛斯．史庫卡//",
                        "引言二//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "烏漾達",
                    "name_en": "Uyanda",
                    "role_zh": "慈濟志工",
                    "role_en": "",
                    "org_en": "",
                },
                {
                    "name_zh": "阿洛斯．史庫卡",
                    "name_en": "Alois L. Sikuka",
                    "role_zh": "國會議員",
                    "role_en": "",
                    "org_en": "",
                },
            ],
        )

    def test_parse_input_extracts_english_name_from_multi_parenthetical_cue(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(SB) (Anabel) (17秒)",
                        "/*SUPER:",
                        "社工│安娜貝爾//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "安娜貝爾",
                    "name_en": "Anabel",
                    "role_zh": "社工",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_extracts_accented_name_from_mixed_parentheses(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(19秒  Paulino Eusébio Sande）",
                        "/*SUPER:",
                        "恩佳小學校長│保利諾·尤西比奧//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "保利諾·尤西比奧",
                    "name_en": "Paulino Eusébio Sande",
                    "role_zh": "恩佳小學校長",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_extracts_fullwidth_latin_name_from_parentheses(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        '(3" Ｍaya)',
                        "/*SUPER:",
                        "杜爾加瓦蒂的姊姊｜瑪雅//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "瑪雅",
                    "name_en": "Maya",
                    "role_zh": "杜爾加瓦蒂的姊姊",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_ignores_instructional_parenthetical_english_and_keeps_full_name(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(11，Dr. Margarita “Maui” Bondoc-Hermosa 毛伊醫師)",
                        "/*SUPER:",
                        "慈濟人醫會眼科醫師│毛伊//",
                        "引言一//",
                        "*/",
                        "",
                        "(志工說以Maul 來為名)",
                        "(7，Allyza Jane Alinsub Sergida)",
                        "/*SUPER:",
                        "小患者│艾莉莎//",
                        "引言二//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "毛伊",
                    "name_en": 'Dr. Margarita "Maui" Bondoc-Hermosa',
                    "role_zh": "慈濟人醫會眼科醫師",
                    "role_en": "",
                    "org_en": "",
                },
                {
                    "name_zh": "艾莉莎",
                    "name_en": "Allyza Jane Alinsub Sergida",
                    "role_zh": "小患者",
                    "role_en": "",
                    "org_en": "",
                },
            ],
        )

    def test_parse_input_ignores_super1_blocks_for_people(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                        "(8．)",
                        "/*SUPER1:",
                        "我們愛我們的新學校 耶",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "烏漾達",
                    "name_en": "Uyanda",
                    "role_zh": "慈濟志工",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_keeps_repeated_super_speaker_in_order(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(Michael Shiu)",
                        "/*SUPER:",
                        "市議員｜邵浩然//",
                        "第一段//",
                        "*/",
                        "",
                        "(16，粵語。Michael Shiu)",
                        "/*SUPER:",
                        "市議員｜邵浩然//",
                        "第二段//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "邵浩然",
                    "name_en": "Michael Shiu",
                    "role_zh": "市議員",
                    "role_en": "",
                    "org_en": "",
                },
                {
                    "name_zh": "邵浩然",
                    "name_en": "",
                    "role_zh": "市議員",
                    "role_en": "",
                    "org_en": "",
                },
            ],
        )

    def test_parse_input_accepts_super_block_without_colon(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        '(15" Altaf Husen Khan)',
                        "/*SUPER",
                        "悉達多基礎學校校長｜阿達夫//",
                        "我們非常高興//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "阿達夫",
                    "name_en": "Altaf Husen Khan",
                    "role_zh": "悉達多基礎學校校長",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_strips_trailing_date_parenthesis_for_role_only_super(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "news_input.txt"
            input_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(Venerable Master Cheng Yen)",
                        "/*SUPER:",
                        "證嚴上人開示(2026.1.21)｜//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(input_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "",
                    "name_en": "Venerable Master Cheng Yen",
                    "role_zh": "證嚴上人開示",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_supports_separate_meta_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"

            body_path.write_text(
                "\n".join(
                    [
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(body_path, meta_path)

        self.assertEqual(data["title_en"], "English Title")
        self.assertEqual(data["overview_en"], "English overview.")
        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "烏漾達",
                    "name_en": "Uyanda",
                    "role_zh": "慈濟志工",
                    "role_en": "",
                    "org_en": "",
                }
            ],
        )

    def test_parse_input_ignores_unknown_meta_keys_in_separate_meta_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"

            body_path.write_text(
                "\n".join(
                    [
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "OLD_TITLE: Legacy title",
                        "OLD_OVERVIEW: Legacy overview",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(body_path, meta_path)

        self.assertEqual(data["title_en"], "")
        self.assertEqual(data["overview_en"], "")

    def test_parse_input_supports_separate_meta_people_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"

            body_path.write_text(
                "\n".join(
                    [
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                        "PEOPLE:",
                        "慈濟志工｜烏漾達",
                        "Uyanda",
                        "Volunteer",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(body_path, meta_path)

        self.assertEqual(data["title_en"], "English Title")
        self.assertEqual(data["overview_en"], "English overview.")
        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "烏漾達",
                    "name_en": "Uyanda",
                    "role_zh": "慈濟志工",
                    "role_en": "Volunteer",
                    "org_en": "",
                    "label_zh": "慈濟志工｜烏漾達",
                }
            ],
        )

    def test_parse_input_keeps_non_person_blocks_from_meta_people(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"

            body_path.write_text(
                "\n".join(
                    [
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                        "PEOPLE:",
                        "慈濟志工｜烏漾達",
                        "Uyanda",
                        "Volunteer",
                        "",
                        "0027",
                        "(最後上字在畫面上：慈濟 我願意 我們愛你)",
                        "Tzu Chi, we love you.",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(body_path, meta_path)

        blocks = data["people_meta_blocks"]
        free_block = next(b for b in blocks if b.get("kind") == "free")
        self.assertEqual(
            free_block.get("lines"),
            [
                "0027",
                "(最後上字在畫面上：慈濟 我願意 我們愛你)",
                "Tzu Chi, we love you.",
            ],
        )

    def test_build_people_lines_preserves_meta_people_block_order(self) -> None:
        people = [
            {
                "name_zh": "李麗華",
                "name_en": "Lee Leyhua",
                "role_zh": "馬來西亞慈濟志工",
                "role_en": "",
                "org_en": "",
                "label_zh": "馬來西亞慈濟志工｜李麗華",
            },
            {
                "name_zh": "瑪麗亞",
                "name_en": "Maria Jose",
                "role_zh": "西班牙志工",
                "role_en": "Tzu Chi volunteer",
                "org_en": "Spain",
                "label_zh": "西班牙志工｜瑪麗亞",
            },
        ]
        ordered_blocks = [
            {
                "kind": "free",
                "lines": [
                    "0032",
                    "(接法香",
                    "把你的手碰觸到水)",
                    "Take a leaf.",
                    "Touch the water with your hand.",
                ],
            },
            {
                "kind": "person",
                "entry": {
                    "label_zh": "馬來西亞慈濟志工｜李麗華",
                    "name_zh": "李麗華",
                    "role_zh": "馬來西亞慈濟志工",
                    "name_en": "Lee Leyhua",
                    "role_en": "",
                    "org_en": "",
                },
            },
            {
                "kind": "free",
                "lines": ["0099", "(中段文字)", "Middle block"],
            },
            {
                "kind": "person",
                "entry": {
                    "label_zh": "西班牙志工｜瑪麗亞",
                    "name_zh": "瑪麗亞",
                    "role_zh": "西班牙志工",
                    "name_en": "Maria Jose",
                    "role_en": "Tzu Chi volunteer",
                    "org_en": "Spain",
                },
            },
        ]

        lines = build_people_lines(people, ordered_blocks=ordered_blocks)
        texts = [line for line in lines if line]

        assert texts.index("0032") < texts.index("馬來西亞慈濟志工｜李麗華")
        assert texts.index("0099") > texts.index("馬來西亞慈濟志工｜李麗華")
        assert texts.index("0099") < texts.index("西班牙志工｜瑪麗亞")

    def test_build_people_lines_uses_ordered_people_blocks_without_fanning_out_duplicates(self) -> None:
        people = [
            {
                "name_zh": "邵浩然",
                "name_en": "MICHAEL",
                "role_zh": "市議員",
                "role_en": "Councillor",
                "org_en": "Richmond Hill",
                "label_zh": "市議員｜邵浩然",
            },
            {
                "name_zh": "邵浩然",
                "name_en": "MICHAEL",
                "role_zh": "市議員",
                "role_en": "Councillor",
                "org_en": "Richmond Hill",
                "label_zh": "市議員｜邵浩然",
            },
            {
                "name_zh": "張肅建",
                "name_en": "THE OTHER",
                "role_zh": "志工",
                "role_en": "Tzu Chi volunteer",
                "org_en": "",
                "label_zh": "志工｜張肅建",
            },
        ]
        ordered_blocks = [
            {
                "kind": "person",
                "entry": {
                    "label_zh": "市議員｜邵浩然",
                    "name_zh": "邵浩然",
                    "role_zh": "市議員",
                    "name_en": "Michael Shiu",
                    "role_en": "Councillor",
                    "org_en": "Richmond Hill",
                },
            },
            {
                "kind": "person",
                "entry": {
                    "label_zh": "志工｜張肅建",
                    "name_zh": "張肅建",
                    "role_zh": "志工",
                    "name_en": "???",
                    "role_en": "Tzu Chi volunteer",
                    "org_en": "",
                },
            },
        ]

        lines = build_people_lines(people, ordered_blocks=ordered_blocks)
        texts = [line for line in lines if line]
        michael_positions = [idx for idx, text in enumerate(texts) if text == "MICHAEL"]
        other_position = texts.index("THE OTHER")
        assert len(michael_positions) == 1
        assert michael_positions[0] < other_position

    def test_build_people_lines_prefers_super_sequence_when_requested(self) -> None:
        people = [
            {
                "name_zh": "班傑納斯",
                "name_en": "Baijnath Barai",
                "role_zh": "藍毘尼志工",
                "role_en": "Tzu Chi volunteer",
                "org_en": "Lumbini, Nepal",
                "label_zh": "藍毘尼志工｜班傑納斯",
            },
            {
                "name_zh": "杜爾加瓦蒂",
                "name_en": "Dargawati",
                "role_zh": "慈濟縫紉班老師",
                "role_en": "Tzu Chi sewing teacher",
                "org_en": "",
                "label_zh": "慈濟縫紉班老師｜杜爾加瓦蒂",
            },
            {
                "name_zh": "瑪雅",
                "name_en": "Maya",
                "role_zh": "杜爾加瓦蒂的姊姊",
                "role_en": "Dargawati's sister",
                "org_en": "",
                "label_zh": "杜爾加瓦蒂的姊姊｜瑪雅",
            },
            {
                "name_zh": "班傑納斯",
                "name_en": "Baijnath Barai",
                "role_zh": "藍毘尼志工",
                "role_en": "Tzu Chi volunteer",
                "org_en": "Lumbini, Nepal",
                "label_zh": "藍毘尼志工｜班傑納斯",
            },
            {
                "name_zh": "杜爾加瓦蒂",
                "name_en": "Dargawati",
                "role_zh": "慈濟縫紉班老師",
                "role_en": "Tzu Chi sewing teacher",
                "org_en": "",
                "label_zh": "慈濟縫紉班老師｜杜爾加瓦蒂",
            },
        ]
        ordered_blocks = [
            {"kind": "person", "entry": {"label_zh": "藍毘尼志工｜班傑納斯"}},
            {"kind": "person", "entry": {"label_zh": "慈濟縫紉班老師｜杜爾加瓦蒂"}},
            {"kind": "person", "entry": {"label_zh": "杜爾加瓦蒂的姊姊｜瑪雅"}},
        ]

        lines = build_people_lines(
            people,
            ordered_blocks=ordered_blocks,
            prefer_people_sequence=True,
        )
        texts = [line for line in lines if line]
        assert texts.count("Baijnath Barai") == 2
        assert texts.count("Dargawati") == 2
        assert texts.count("Maya") == 1
        last_baijnath = len(texts) - 1 - list(reversed(texts)).index("Baijnath Barai")
        assert texts.index("Maya") < last_baijnath


def test_build_people_lines_appends_tail_lines_after_people() -> None:
    lines = build_people_lines(
        [
            {
                "name_zh": "烏漾達",
                "name_en": "Uyanda",
                "role_zh": "慈濟志工",
                "role_en": "Volunteer",
                "org_en": "",
            }
        ],
        ["0027", "(最後上字在畫面上：慈濟 我願意 我們愛你)", "Tzu Chi, we love you."],
    )

    assert lines[-3:] == [
        "0027",
        "(最後上字在畫面上：慈濟 我願意 我們愛你)",
        "Tzu Chi, we love you.",
    ]

    def test_generate_meta_supports_separate_meta_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            body_path.write_text(
                "\n".join(
                    [
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            generate_meta(template_path, body_path, output_path, meta_path=meta_path)

            doc = Document(str(output_path))
            texts = [p.text for p in doc.paragraphs]

        self.assertEqual(
            texts,
            [
                "重點標",
                "",
                "English Title",
                "名字職銜",
                "",
                "",
                "慈濟志工｜烏漾達",
                "Uyanda",
                "{{慈濟志工}}",
                "",
                "YT簡介",
                "",
                "English overview.",
            ],
        )

    def test_generate_meta_does_not_leave_empty_run_in_people_paragraph(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            body_path.write_text(
                "\n".join(
                    [
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                        "PEOPLE:",
                        "慈濟志工｜烏漾達",
                        "Uyanda",
                        "Volunteer",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            generate_meta(template_path, body_path, output_path, meta_path=meta_path)

            doc = Document(str(output_path))
            people_paragraph = doc.paragraphs[6]

        self.assertEqual(people_paragraph.text, "慈濟志工｜烏漾達")
        self.assertEqual(
            [run.text for run in people_paragraph.runs],
            ["慈濟志工｜烏漾達"],
        )

    def test_generate_meta_preserves_label_highlights_from_template(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            body_path = tmpdir_path / "source.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            template_doc = Document(str(template_path))
            for paragraph in template_doc.paragraphs:
                if paragraph.text in {"重點標", "名字職銜", "YT簡介"}:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            template_doc.save(str(template_path))
            body_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                        "BODY:",
                        "(6． Uyanda烏漾達)",
                        "/*SUPER:",
                        "慈濟志工│烏漾達//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            generate_meta(template_path, body_path, output_path)

            doc = Document(str(output_path))
            labels = {
                paragraph.text: [run.font.highlight_color for run in paragraph.runs]
                for paragraph in doc.paragraphs
                if paragraph.text in {"重點標", "名字職銜", "YT簡介"}
            }

        self.assertEqual(labels["重點標"], [WD_COLOR_INDEX.YELLOW])
        self.assertEqual(labels["名字職銜"], [WD_COLOR_INDEX.YELLOW])
        self.assertEqual(labels["YT簡介"], [WD_COLOR_INDEX.YELLOW])

    def test_generate_meta_inserts_blank_after_chinese_labels(self) -> None:
        source_text = "\n".join(
            [
                "TITLE: English Title",
                "OVERVIEW: English overview.",
                "",
                "BODY:",
                "(  13   Alice )",
                "/*SUPER:",
                "病患│甲//",
                "引言一//",
                "*/",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            payload_path = tmpdir_path / "news_input.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            payload_path.write_text(source_text, encoding="utf-8")

            generate_meta(template_path, payload_path, output_path)
            texts = [p.text for p in Document(str(output_path)).paragraphs]

        idx_title = texts.index("重點標")
        idx_people = texts.index("名字職銜")
        idx_overview = texts.index("YT簡介")
        self.assertEqual(texts[idx_title + 1], "")
        self.assertEqual(texts[idx_people + 1], "")
        self.assertEqual(texts[idx_overview + 1], "")

    def test_generate_meta_omits_english_name_from_label_when_repeated_below(self) -> None:
        source_text = "\n".join(
            [
                "BODY:",
                "(7)",
                "/*SUPER:",
                "艾莉莎的父親│Mar Jason B. Sergida//",
                "內容//",
                "*/",
                "",
            ]
        )

        meta_text = "\n".join(
            [
                "TITLE: English Title",
                "OVERVIEW: English overview.",
                "",
                "PEOPLE:",
                "艾莉莎的父親｜Mar Jason B. Sergida",
                "Mar Jason B. Sergida",
                "Allyza's father",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            body_path.write_text(source_text, encoding="utf-8")
            meta_path.write_text(meta_text, encoding="utf-8")

            generate_meta(template_path, body_path, output_path, meta_path=meta_path)

            texts = [p.text for p in Document(str(output_path)).paragraphs]

        self.assertEqual(
            texts,
            [
                "重點標",
                "",
                "English Title",
                "名字職銜",
                "",
                "",
                "艾莉莎的父親",
                "Mar Jason B. Sergida",
                "Allyza's father",
                "",
                "YT簡介",
                "",
                "English overview.",
            ],
        )

    def test_generate_meta_omits_name_placeholder_for_role_only_entries(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            body_path = tmpdir_path / "source.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            body_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                        "BODY:",
                        "(8．)",
                        "/*SUPER:",
                        "家長｜//",
                        "內容//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            generate_meta(template_path, body_path, output_path)

            doc = Document(str(output_path))
            texts = [p.text for p in doc.paragraphs]

        self.assertEqual(
            texts,
            [
                "重點標",
                "",
                "English Title",
                "名字職銜",
                "",
                "",
                "家長",
                "{{家長}}",
                "",
                "YT簡介",
                "",
                "English overview.",
            ],
        )

    def test_generate_meta_raises_when_overview_missing(self) -> None:
        source_text = "\n".join(
            [
                "TITLE_TEXT: 中文標題",
                "SUMMARY:",
                "中文摘要",
                "",
                "TITLE: English Title",
                "",
                "BODY:",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            payload_path = tmpdir_path / "news_input.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            payload_path.write_text(source_text, encoding="utf-8")

            with self.assertRaisesRegex(
                ValueError, r"\[error\] Missing required field: OVERVIEW"
            ):
                generate_meta(template_path, payload_path, output_path)

    def test_renders_from_news_txt_input(self) -> None:
        source_text = "\n".join(
            [
                "TITLE_TEXT: 沿海義診守護居民健康",
                "TITLE_URL: https://example.com/news/story",
                "SUMMARY:",
                "A volunteer team hosted a two-day community clinic in a coastal town.",
                "(  11/16~17 )",
                "",
                "TITLE: Coastal Clinic Restores Access to Care",
                "OVERVIEW: A volunteer medical team brought screenings and referrals to local residents.",
                "",
                "BODY:",
                "1_0014",
                "居民一早就到現場排隊。",
                "Residents lined up early for registration and screening.",
                "",
                "(  13   Mr. Chen )",
                "/*SUPER:",
                "居民│陳先生//",
                "真的很感謝大家的幫忙//",
                "*/",
                "I am truly grateful for everyone's help.",
                "",
                "(  14   )",
                "/*SUPER:",
                "醫師│林醫師//",
                "我們會持續協助需要後續治療的居民//",
                "*/",
                "We will continue helping residents who need follow-up care.",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            payload_path = tmpdir_path / "news_input.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            payload_path.write_text(source_text, encoding="utf-8")

            generate_meta(template_path, payload_path, output_path)

            doc = Document(str(output_path))
            texts = [p.text for p in doc.paragraphs]

        self.assertEqual(
            texts,
            [
                "重點標",
                "",
                "Coastal Clinic Restores Access to Care",
                "名字職銜",
                "",
                "",
                "居民｜陳先生",
                "Mr. Chen",
                "{{居民}}",
                "",
                "醫師｜林醫師",
                "{{林醫師}}",
                "{{醫師}}",
                "",
                "YT簡介",
                "",
                "A volunteer medical team brought screenings and referrals to local residents.",
            ],
        )

    def test_meta_people_overrides_matching_blocks(self) -> None:
        source_text = "\n".join(
            [
                "TITLE_TEXT: 測試標題",
                "SUMMARY:",
                "測試摘要",
                "",
                "TITLE: Test English Title",
                "OVERVIEW: Test English overview.",
                "",
                "PEOPLE:",
                "居民｜受訪者",
                "Guest (Edited)",
                "Resident",
                "",
                "醫師｜林醫師",
                "Dr. Lin",
                "Doctor",
                "Clinic Team",
                "",
                "BODY:",
                "(  13   Guest )",
                "/*SUPER:",
                "居民│受訪者//",
                "內容//",
                "*/",
                "English line.",
                "",
                "(  14   )",
                "/*SUPER:",
                "醫師│林醫師//",
                "內容//",
                "*/",
                "English line.",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            input_path = tmpdir_path / "news_input.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            input_path.write_text(source_text, encoding="utf-8")

            generate_meta(template_path, input_path, output_path)
            doc = Document(str(output_path))
            texts = [p.text for p in doc.paragraphs]

        self.assertEqual(
            texts,
            [
                "重點標",
                "",
                "Test English Title",
                "名字職銜",
                "",
                "",
                "居民｜受訪者",
                "Guest (Edited)",
                "Resident",
                "",
                "醫師｜林醫師",
                "Dr. Lin",
                "Doctor",
                "Clinic Team",
                "",
                "YT簡介",
                "",
                "Test English overview.",
            ],
        )

    def test_meta_people_overrides_support_role_only_labels_by_english_name(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"

            body_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(11, Timothy Yu)",
                        "/*SUPER:",
                        "慈濟人醫會醫師│余俊傑//",
                        "內容//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "PEOPLE:",
                        "慈濟人醫會醫師",
                        "Timothy Yu",
                        "TIMA Doctor",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            data = parse_input(body_path, meta_path)

        self.assertEqual(
            data["people"],
            [
                {
                    "name_zh": "余俊傑",
                    "name_en": "Timothy Yu",
                    "role_zh": "慈濟人醫會醫師",
                    "role_en": "TIMA Doctor",
                    "org_en": "",
                    "label_zh": "慈濟人醫會醫師",
                }
            ],
        )

    def test_generate_meta_omits_role_placeholder_when_only_name_en_is_provided(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            body_path = tmpdir_path / "source.txt"
            meta_path = tmpdir_path / "meta.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            body_path.write_text(
                "\n".join(
                    [
                        "BODY:",
                        "(6． Patient)",
                        "/*SUPER:",
                        "患者│個案A//",
                        "引言一//",
                        "*/",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            meta_path.write_text(
                "\n".join(
                    [
                        "TITLE: English Title",
                        "OVERVIEW: English overview.",
                        "",
                        "PEOPLE:",
                        "患者",
                        "Patient",
                        "",
                    ]
                ),
                encoding="utf-8",
            )

            generate_meta(template_path, body_path, output_path, meta_path=meta_path)
            texts = [p.text for p in Document(str(output_path)).paragraphs]

        assert "Patient" in texts
        assert "{{患者}}" not in texts

    def test_generate_meta_enforces_body_font_size_from_shared_token(self) -> None:
        source_text = "\n".join(
            [
                "TITLE: English Title",
                "OVERVIEW: English overview.",
                "",
                "BODY:",
                "(  13   Alice )",
                "/*SUPER:",
                "病患│甲//",
                "引言一//",
                "*/",
                "",
            ]
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "meta_template.docx"
            payload_path = tmpdir_path / "news_input.txt"
            output_path = tmpdir_path / "meta.docx"

            self._build_template(template_path)
            template_doc = Document(str(template_path))
            template_doc.styles["Normal"].font.size = Pt(10)
            template_doc.save(str(template_path))
            payload_path.write_text(source_text, encoding="utf-8")

            generate_meta(template_path, payload_path, output_path)

            doc = Document(str(output_path))
            title_para = next(p for p in doc.paragraphs if p.text == "English Title")
            overview_para = next(p for p in doc.paragraphs if p.text == "English overview.")

        self.assertTrue(
            all(run.font.size == Pt(BODY_TEXT_SIZE_PT) for run in title_para.runs if run.text)
        )
        self.assertTrue(
            all(run.font.size == Pt(BODY_TEXT_SIZE_PT) for run in overview_para.runs if run.text)
        )


def test_default_output_path_uses_source_stem(tmp_path: Path) -> None:
    source = tmp_path / "sample_story_final.docx"
    output_dir = tmp_path / "output"
    output = default_output_path(source, output_dir)
    assert output == output_dir / "sample_story_標題職銜_final.docx"


def test_default_output_path_adds_final_suffix_when_missing(tmp_path: Path) -> None:
    source = tmp_path / "sample_story.docx"
    output_dir = tmp_path / "output"
    output = default_output_path(source, output_dir)
    assert output == output_dir / "sample_story_標題職銜_final.docx"


def test_resolve_template_path_uses_script_directory_for_relative_paths() -> None:
    previous_cwd = Path.cwd()
    with tempfile.TemporaryDirectory() as tmpdir:
        os.chdir(tmpdir)
        try:
            template = resolve_template_path(Path("templates/meta_template.docx"))
        finally:
            os.chdir(previous_cwd)
    assert template == Path(__file__).resolve().parent.parent / "templates" / "meta_template.docx"


if __name__ == "__main__":
    unittest.main()
