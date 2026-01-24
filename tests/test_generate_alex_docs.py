from pathlib import Path

from docx import Document

import zipfile
import xml.etree.ElementTree as ET

from generate_alex_docs import extract_alex_titles, generate_docs, normalize_title


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def test_normalize_title_strips_translator_tag() -> None:
    title = "大愛醫生館 - 怎麼坐才算有“坐相”？st/rc"
    assert normalize_title(title) == "大愛醫生館 怎麼坐才算有“坐相”？"


def test_extract_alex_titles_from_schedule(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    _write_docx(
        schedule_path,
        [
            "節目6則",
            "1. elijah",
            "節目甲 - 測試標題 em/el",
            "https://example.com/1",
            "2. alex",
            "大愛醫生館 - 怎麼坐才算有“坐相”？st/rc",
            "https://example.com/2",
            "搭配",
            "https://example.com/news",
            "新聞標題",
            "3. alex",
            "大愛真健康 - 5分鐘高效有氧 | 上下肢肌耐力 | 肩腿| 背腿 | 核心 nick/cc",
            "https://example.com/3",
            "--------------------------------",
            "FB小編文box裡面的新聞已用掉下面5則",
        ],
    )

    titles = extract_alex_titles(schedule_path)
    assert titles == [
        "大愛醫生館 怎麼坐才算有“坐相”？",
        "大愛真健康 5分鐘高效有氧 | 上下肢肌耐力 | 肩腿| 背腿 | 核心",
    ]


def test_generated_docx_is_well_formed(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "--------------------------------",
        ],
    )
    _write_docx(template_path, ["Base paragraph."])
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    assert len(output_paths) == 1

    output_path = output_paths[0]
    with zipfile.ZipFile(output_path) as zf:
        xml_text = zf.read("word/document.xml")
    ET.fromstring(xml_text)
    Document(str(output_path))
