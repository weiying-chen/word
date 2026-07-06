from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt

import zipfile
import xml.etree.ElementTree as ET

from generate_posts import generate_docs
from style_tokens import BODY_TEXT_SIZE_PT, REFERENCE_TEXT_SIZE_PT


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def test_generated_docx_is_well_formed(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "{{HEADER_URL}}",
            "{{REF_URL}}",
            "{{REF_TITLE}}",
            "{{VIDEO_URL}}",
            "{{VIDEO_TITLE}}",
        ],
    )
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    assert len(output_paths) == 1

    output_path = output_paths[0]
    doc = Document(str(output_path))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "Program - Test Title"
    assert texts[1] == "https://example.com/video"
    assert texts[2] == "https://example.com/news"
    assert texts[3] == "News title"
    assert texts[4] == "https://example.com/video"
    assert texts[5] == "Program - Test Title"
    with zipfile.ZipFile(output_path) as zf:
        xml_text = zf.read("word/document.xml")
    ET.fromstring(xml_text)


def test_generated_docx_from_alex_blocks_uses_date_prefix(tmp_path: Path) -> None:
    schedule_path = tmp_path / "alex_blocks.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "1",
            "參考資料:",
            "https://example.com/news",
            "26/1/23",
            "News title",
            "要用的影片:",
            "https://example.com/video",
            "Program - Test Title (健康節目 - 測試標題)",
            "English prompt",
            "中文提示",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "{{HEADER_URL}}",
            "{{REF_URL}}",
            "{{REF_TITLE}}",
            "{{VIDEO_URL}}",
            "{{VIDEO_TITLE}}",
            "{{VIDEO_DESC_EN}}",
            "{{VIDEO_DESC_ZH}}",
        ],
    )
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="日期未定_",
        filename_suffix="",
    )
    assert len(output_paths) == 1

    output_path = output_paths[0]
    assert output_path.name.startswith("260123_")
    doc = Document(str(output_path))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "Program - Test Title (健康節目 - 測試標題)"
    assert texts[1] == "https://example.com/video"
    assert texts[2] == "https://example.com/news"
    assert texts[3] == "News title"
    assert texts[4] == "https://example.com/video"
    assert texts[5] == "Program - Test Title (健康節目 - 測試標題)"
    assert texts[6] == "English prompt"
    assert texts[7] == "中文提示"


def test_generated_docx_from_schedule_includes_reference_bilingual_summary_in_source_section(
    tmp_path: Path,
) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "人文講堂 - 養腦 防失智 - 曾文毅 [2]",
            "https://www.youtube.com/watch?v=LA76bSzLVgE&t=280s",
            "搭配",
            "https://example.com/news",
            "Reducing Brain Age to Prevent Dementia (人文講堂 - 養腦 防失智 - 曾文毅 [2])",
            'From his own example of following the "Three Sevens Rule," Prof. Tseng Wen-yih of the National Taiwan University College of Medicine shares how we can effectively reduce brain age to maintain its health and minimize the risk of dementia.',
            "研究發現六十歲以上，若腦年齡退化十歲以上，未來三年內可能會發生認知功能衰退。臺大醫學院兼任教授曾文毅分享逆轉腦齡的七七七法則，告訴大家如何維持腦健康，遠離失智。",
            "--------------------------------",
        ],
    )
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)
    for text in [
        "標題",
        "{{HEADER_TITLE}}",
        "參考資料：",
        "{{REF_URL}}",
        "{{REF_TITLE}}",
        "{{REF_SUMMARY_ZH}}",
        "英文翻譯：",
        "{{REF_TITLE_EN}}",
        "{{REF_SUMMARY_EN}}",
        "要用的影片：",
        "{{VIDEO_URL}}",
        "{{VIDEO_TITLE}}",
        "{{VIDEO_DESC_EN}}",
        "{{VIDEO_DESC_ZH}}",
    ]:
        doc.add_paragraph(text)
    doc.save(template_path)
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    rendered = Document(str(output_paths[0]))
    texts = [p.text.strip() for p in rendered.paragraphs if p.text.strip()]
    raw_texts = [p.text for p in rendered.paragraphs]

    assert "https://www.youtube.com/watch?v=LA76bSzLVgE&t=280s" in texts
    assert "Reducing Brain Age to Prevent Dementia (人文講堂 - 養腦 防失智 - 曾文毅 [2])" in texts
    assert "研究發現六十歲以上，若腦年齡退化十歲以上，未來三年內可能會發生認知功能衰退。臺大醫學院兼任教授曾文毅分享逆轉腦齡的七七七法則，告訴大家如何維持腦健康，遠離失智。" in texts
    assert 'From his own example of following the "Three Sevens Rule," Prof. Tseng Wen-yih of the National Taiwan University College of Medicine shares how we can effectively reduce brain age to maintain its health and minimize the risk of dementia.' in texts
    video_label_idx = raw_texts.index("要用的影片：")
    assert raw_texts[video_label_idx + 1] == ""
    assert raw_texts[video_label_idx + 2] == "https://www.youtube.com/watch?v=LA76bSzLVgE&t=280s"
    assert raw_texts[video_label_idx + 3] == "Reducing Brain Age to Prevent Dementia (人文講堂 - 養腦 防失智 - 曾文毅 [2])"
    assert raw_texts[video_label_idx + 4] == ""
    assert raw_texts[video_label_idx + 5] == 'From his own example of following the "Three Sevens Rule," Prof. Tseng Wen-yih of the National Taiwan University College of Medicine shares how we can effectively reduce brain age to maintain its health and minimize the risk of dementia.'
    assert raw_texts[video_label_idx + 6] == ""
    assert raw_texts[video_label_idx + 7] == "研究發現六十歲以上，若腦年齡退化十歲以上，未來三年內可能會發生認知功能衰退。臺大醫學院兼任教授曾文毅分享逆轉腦齡的七七七法則，告訴大家如何維持腦健康，遠離失智。"

    p_by_text = {p.text.strip(): p for p in rendered.paragraphs if p.text.strip()}
    for text in [
        "Reducing Brain Age to Prevent Dementia (人文講堂 - 養腦 防失智 - 曾文毅 [2])",
        "研究發現六十歲以上，若腦年齡退化十歲以上，未來三年內可能會發生認知功能衰退。臺大醫學院兼任教授曾文毅分享逆轉腦齡的七七七法則，告訴大家如何維持腦健康，遠離失智。",
        'From his own example of following the "Three Sevens Rule," Prof. Tseng Wen-yih of the National Taiwan University College of Medicine shares how we can effectively reduce brain age to maintain its health and minimize the risk of dementia.',
    ]:
        assert all(
            run.font.size == Pt(REFERENCE_TEXT_SIZE_PT)
            for run in p_by_text[text].runs
            if run.text
        ), text


def test_generated_docx_fetches_youtube_video_summary_when_missing(
    tmp_path: Path,
) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "All About Health - When Liver Disease Goes Unnoticed (大愛醫生館 - 雲霧肝癌)",
            "https://www.youtube.com/watch?v=47tbzXquNm8",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "{{HEADER_URL}}",
            "要用的影片：",
            "{{VIDEO_URL}}",
            "{{VIDEO_TITLE}}",
            "{{VIDEO_DESC_EN}}",
            "{{VIDEO_DESC_ZH}}",
        ],
    )
    output_dir.mkdir()

    with patch(
        "generate_posts.fetch_youtube_video_descriptions",
        return_value=(
            "For hepatitis B and C carriers, regular blood tests are important—but they don't tell the whole story. Because the liver has no pain sensation, problems can develop silently. In one case, a tumor wasn't discovered until it had grown to 10 centimeters. How did it go unnoticed for so long? Why are imaging tests so important for early detection? Let's hear what the doctor has to say.",
            "對於B型與C型肝炎帶原者而言，單靠定期的抽血檢查並不足夠。由於肝臟沒有痛覺神經，問題往往不易被察覺。有一個案例，腫瘤長到10公分才被發現。為何會這麼晚才察覺？影像檢查在早期發現中扮演什麼角色？一起來聽聽醫師怎麼說。",
        ),
    ):
        output_paths = generate_docs(
            schedule_path=schedule_path,
            template_path=template_path,
            output_dir=output_dir,
            filename_prefix="",
            filename_suffix="",
        )

    raw_texts = [p.text for p in Document(str(output_paths[0])).paragraphs]
    video_label_idx = raw_texts.index("要用的影片：")

    assert raw_texts[video_label_idx + 1] == ""
    assert raw_texts[video_label_idx + 2] == "https://www.youtube.com/watch?v=47tbzXquNm8"
    assert raw_texts[video_label_idx + 3] == (
        "All About Health - When Liver Disease Goes Unnoticed (大愛醫生館 - 雲霧肝癌)"
    )
    assert raw_texts[video_label_idx + 4] == ""
    assert raw_texts[video_label_idx + 5] == (
        "For hepatitis B and C carriers, regular blood tests are important—but they don't tell the whole story. Because the liver has no pain sensation, problems can develop silently. In one case, a tumor wasn't discovered until it had grown to 10 centimeters. How did it go unnoticed for so long? Why are imaging tests so important for early detection? Let's hear what the doctor has to say."
    )
    assert raw_texts[video_label_idx + 6] == ""
    assert raw_texts[video_label_idx + 7] == (
        "對於B型與C型肝炎帶原者而言，單靠定期的抽血檢查並不足夠。由於肝臟沒有痛覺神經，問題往往不易被察覺。有一個案例，腫瘤長到10公分才被發現。為何會這麼晚才察覺？影像檢查在早期發現中扮演什麼角色？一起來聽聽醫師怎麼說。"
    )


def test_generated_docx_has_highlighted_ref_hyperlink(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{REF_URL}}",
        ],
    )
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    output_path = output_paths[0]

    with zipfile.ZipFile(output_path) as zf:
        doc = ET.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    url_paragraph = None
    for p in doc.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in p.findall(".//w:t", ns))
        if "https://example.com/news" in text:
            url_paragraph = p
            break

    assert url_paragraph is not None
    hyperlinks = url_paragraph.findall("w:hyperlink", ns)
    assert len(hyperlinks) == 1
    h_runs = hyperlinks[0].findall("w:r", ns)
    assert h_runs
    h_rpr = h_runs[0].find("w:rPr", ns)
    h_highlight = h_rpr.find("w:highlight", ns)
    h_size = h_rpr.find("w:sz", ns)
    assert h_highlight.get("{%s}val" % ns["w"]) == "cyan"
    assert h_size.get("{%s}val" % ns["w"]) == "20"


def test_generated_docx_syncs_empty_paragraph_indent(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{VIDEO_TITLE}}",
            "",
            "{{VIDEO_DESC_EN}}",
        ],
    )
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    output_path = output_paths[0]
    doc = Document(str(output_path))

    video_title = doc.paragraphs[0]
    empty_para = doc.paragraphs[1]
    assert video_title.text.strip()
    assert empty_para.text.strip() == ""
    assert empty_para.paragraph_format.left_indent == video_title.paragraph_format.left_indent
    assert empty_para.paragraph_format.first_line_indent == video_title.paragraph_format.first_line_indent
    assert empty_para.paragraph_format.left_indent == Inches(0.5)


def test_generated_docx_sets_source_indent_for_labels(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "要用的影片：",
            "{{VIDEO_URL}}",
        ],
    )
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    output_path = output_paths[0]
    doc = Document(str(output_path))

    label_para = doc.paragraphs[0]
    blank_para = doc.paragraphs[1]
    url_para = doc.paragraphs[2]
    assert label_para.text.strip() == "要用的影片："
    assert not blank_para.text.strip()
    assert url_para.text.strip() == "https://example.com/video"
    assert label_para.paragraph_format.left_indent == Inches(0.5)
    assert label_para.paragraph_format.first_line_indent == 0
    assert url_para.paragraph_format.left_indent == Inches(0.5)
    assert url_para.paragraph_format.first_line_indent == 0
    assert all(run.font.size == Pt(REFERENCE_TEXT_SIZE_PT) for run in label_para.runs if run.text)


def test_generated_bodhi_docx_puts_english_title_under_chinese_title(tmp_path: Path) -> None:
    schedule_path = tmp_path / "bodhi.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "菩提1則",
            "1. alex",
            "1/20首播 廣行環保護人間",
            "https://www.daai.tv/master/life-wisdom/P90230231?more=true",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "{{HEADER_URL}}",
            "參考資料：",
            "{{REF_URL}}",
            "{{REF_TITLE}}",
            "要用的影片：",
            "{{VIDEO_URL}}",
            "{{VIDEO_TITLE}}",
            "#hashtagline",
        ],
    )
    output_dir.mkdir()

    en_title = "32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore"
    with patch("generate_posts.fetch_bodhi_english_subtitle", return_value=en_title):
        output_paths = generate_docs(
            schedule_path=schedule_path,
            template_path=template_path,
            output_dir=output_dir,
            filename_prefix="",
            filename_suffix="",
        )

    doc = Document(str(output_paths[0]))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == (
        "人間菩提\n1/20首播 廣行環保護人間\n"
        "32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore"
    )
    combined = "廣行環保護人間\n32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore"
    assert texts.count(combined) == 1
    assert "#hashtagline" in texts


def test_generated_bodhi_docx_strips_title_placeholder_block(tmp_path: Path) -> None:
    schedule_path = tmp_path / "bodhi.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "菩提1則",
            "1. alex",
            "7/3善用此生厚美德",
            "https://www.daai.tv/master/life-wisdom/P90230307?more=true",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "{{HEADER_URL}}",
            "標題",
            "{{TITLE_LINE_1}}",
            "{{TITLE_LINE_2}}",
            "{{POST_EN}}",
            "{{HASHTAGS_EN}}",
            "{{POST_ZH}}",
            "{{HASHTAGS_ZH}}",
            "參考資料：",
            "{{REF_URL}}",
            "{{REF_TITLE}}",
        ],
    )
    output_dir.mkdir()

    en_title = "Cherishing Every Opportunity to Serve"
    with patch("generate_posts.fetch_bodhi_english_subtitle", return_value=en_title):
        output_paths = generate_docs(
            schedule_path=schedule_path,
            template_path=template_path,
            output_dir=output_dir,
            filename_prefix="",
            filename_suffix="",
        )

    texts = [p.text.strip() for p in Document(str(output_paths[0])).paragraphs]

    assert "標題" not in texts
    assert "{{TITLE_LINE_1}}" not in texts
    assert "{{TITLE_LINE_2}}" not in texts
    assert "{{POST_EN}}" in texts
    assert "{{POST_ZH}}" in texts


def test_generated_bodhi_docx_injects_english_under_fixed_title_lines_not_hashtags(
    tmp_path: Path,
) -> None:
    schedule_path = tmp_path / "bodhi_fixed_title.docx"
    template_path = tmp_path / "template_fixed.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "菩提1則",
            "1. alex",
            "1/20首播 廣行環保護人間",
            "https://www.daai.tv/master/life-wisdom/P90230231?more=true",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "{{REF_TITLE}}",
            "◎標題：廣行環保護人間",
            "廣行環保護人間",
            "#廣行環保護人間",
        ],
    )
    output_dir.mkdir()

    en_title = "32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore"
    with patch("generate_posts.fetch_bodhi_english_subtitle", return_value=en_title):
        output_paths = generate_docs(
            schedule_path=schedule_path,
            template_path=template_path,
            output_dir=output_dir,
            filename_prefix="",
            filename_suffix="",
        )

    doc = Document(str(output_paths[0]))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "◎標題：廣行環保護人間\n32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore" in texts
    assert "廣行環保護人間\n32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore" in texts
    assert "#廣行環保護人間" in texts


def test_bodhi_injection_uses_title_line_not_program_name(tmp_path: Path) -> None:
    schedule_path = tmp_path / "bodhi_injection_lines.docx"
    template_path = tmp_path / "template_injection_lines.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "人間菩提1則",
            "1. alex",
            "5/18膚慰人間菩薩行",
            "Give Comfort and Encouragement on the Bodhisattva Path",
            "https://www.daai.tv/master/life-wisdom/P90230263?more=true",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "◎標題：膚慰人間菩薩行",
            "人間菩提",
            "#人間菩提",
        ],
    )
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )

    doc = Document(str(output_paths[0]))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert (
        "◎標題：膚慰人間菩薩行\nGive Comfort and Encouragement on the Bodhisattva Path"
        in texts
    )
    assert "人間菩提\nGive Comfort and Encouragement on the Bodhisattva Path" not in texts


def test_generated_posts_enforces_body_12pt_and_source_10pt(tmp_path: Path) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)
    for text in ["{{HEADER_TITLE}}", "{{REF_TITLE}}", "{{VIDEO_TITLE}}"]:
        doc.add_paragraph(text)
    doc.save(template_path)
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    rendered = Document(str(output_paths[0]))

    header_para = next(p for p in rendered.paragraphs if p.text.strip() == "Program - Test Title")
    source_para = next(p for p in rendered.paragraphs if p.text.strip() == "News title")

    assert all(run.font.size == Pt(BODY_TEXT_SIZE_PT) for run in header_para.runs if run.text)
    assert all(run.font.size == Pt(REFERENCE_TEXT_SIZE_PT) for run in source_para.runs if run.text)


def test_generated_posts_keeps_header_url_12pt_and_ref_url_10pt(
    tmp_path: Path,
) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)
    for text in ["{{HEADER_URL}}", "{{REF_URL}}"]:
        doc.add_paragraph(text)
    doc.save(template_path)
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    output_path = output_paths[0]

    with zipfile.ZipFile(output_path) as zf:
        doc_xml = ET.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    by_text: dict[str, ET.Element] = {}
    for para in doc_xml.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in para.findall(".//w:t", ns)).strip()
        if text:
            by_text[text] = para

    header_url_para = by_text["https://example.com/video"]
    ref_url_para = by_text["https://example.com/news"]

    header_run = header_url_para.find(".//w:hyperlink/w:r/w:rPr/w:sz", ns)
    ref_run = ref_url_para.find(".//w:hyperlink/w:r/w:rPr/w:sz", ns)
    assert header_run is not None
    assert ref_run is not None
    assert header_run.get("{%s}val" % ns["w"]) == str(BODY_TEXT_SIZE_PT * 2)
    assert ref_run.get("{%s}val" % ns["w"]) == str(REFERENCE_TEXT_SIZE_PT * 2)


def test_generated_bodhi_docx_adds_blank_after_ref_label(tmp_path: Path) -> None:
    schedule_path = tmp_path / "bodhi.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "菩提1則",
            "1. alex",
            "1/20首播 廣行環保護人間",
            "https://www.daai.tv/master/life-wisdom/P90230231?more=true",
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "參考資料：",
            "{{REF_URL}}",
            "{{REF_TITLE}}",
            "要用的影片：",
            "{{VIDEO_URL}}",
            "{{VIDEO_TITLE}}",
        ],
    )
    output_dir.mkdir()

    with patch(
        "generate_posts.fetch_bodhi_english_subtitle",
        return_value="32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore",
    ):
        output_paths = generate_docs(
            schedule_path=schedule_path,
            template_path=template_path,
            output_dir=output_dir,
            filename_prefix="",
            filename_suffix="",
        )

    doc = Document(str(output_paths[0]))
    label_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "參考資料：")
    assert doc.paragraphs[label_idx + 1].text.strip() == ""


def test_generated_bodhi_docx_does_not_duplicate_ref_url_inside_reference_block(
    tmp_path: Path,
) -> None:
    schedule_path = tmp_path / "bodhi.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    url = "https://www.daai.tv/master/life-wisdom/P90230231?more=true"
    _write_docx(
        schedule_path,
        [
            "菩提1則",
            "1. alex",
            "1/20首播 廣行環保護人間",
            url,
            "--------------------------------",
        ],
    )
    _write_docx(
        template_path,
        [
            "{{HEADER_TITLE}}",
            "{{HEADER_URL}}",
            "參考資料：",
            "{{REF_URL}}",
            "{{REF_TITLE}}",
            "要用的影片：",
            "{{VIDEO_URL}}",
            "{{VIDEO_TITLE}}",
        ],
    )
    output_dir.mkdir()

    with patch(
        "generate_posts.fetch_bodhi_english_subtitle",
        return_value="32 Years of Dedication Tzu Chi’s Recycling Efforts in Singapore",
    ):
        output_paths = generate_docs(
            schedule_path=schedule_path,
            template_path=template_path,
            output_dir=output_dir,
            filename_prefix="",
            filename_suffix="",
        )

    doc = Document(str(output_paths[0]))
    texts = [p.text.strip() for p in doc.paragraphs]
    # Header URL + one reference URL only.
    assert texts.count(url) == 2
    ref_idx = next(i for i, p in enumerate(texts) if p == "參考資料：")
    tail = texts[ref_idx:]
    assert tail.count(url) == 1


def test_generated_posts_title_is_12pt_and_source_block_is_10pt(tmp_path: Path) -> None:
    schedule_path = tmp_path / "alex_blocks.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "1",
            "參考資料:",
            "https://example.com/news",
            "26/1/23",
            "News title",
            "News summary zh",
            "News title en",
            "News summary en",
            "要用的影片:",
            "https://example.com/video",
            "Program - Test Title (健康節目 - 測試標題)",
            "Video desc en",
            "Video desc zh",
        ],
    )
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)
    for text in [
        "標題",
        "{{HEADER_TITLE}}",
        "參考資料：",
        "{{REF_URL}}",
        "{{REF_TITLE}}",
        "{{REF_SUMMARY_ZH}}",
        "英文翻譯：",
        "{{REF_TITLE_EN}}",
        "{{REF_SUMMARY_EN}}",
        "要用的影片：",
        "{{VIDEO_URL}}",
        "{{VIDEO_TITLE}}",
        "{{VIDEO_DESC_EN}}",
        "{{VIDEO_DESC_ZH}}",
    ]:
        doc.add_paragraph(text)
    doc.save(template_path)
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )
    rendered = Document(str(output_paths[0]))

    p_by_text = {p.text.strip(): p for p in rendered.paragraphs if p.text.strip()}
    assert all(
        run.font.size == Pt(BODY_TEXT_SIZE_PT)
        for run in p_by_text["標題"].runs
        if run.text
    )

    source_texts = [
        "參考資料：",
        "https://example.com/news",
        "News title\nNews summary zh\nNews title en\nNews summary en",
        "英文翻譯：",
        "要用的影片：",
        "https://example.com/video",
        "Program - Test Title (健康節目 - 測試標題)",
        "Video desc en",
        "Video desc zh",
    ]
    for text in source_texts:
        assert all(
            run.font.size == Pt(REFERENCE_TEXT_SIZE_PT)
            for run in p_by_text[text].runs
            if run.text
        ), text


def test_generated_posts_set_shared_font_metadata_on_generated_content(
    tmp_path: Path,
) -> None:
    schedule_path = tmp_path / "schedule.docx"
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"

    _write_docx(
        schedule_path,
        [
            "節目1則",
            "1. alex",
            "Program - Test Title st/rc",
            "https://example.com/video",
            "搭配",
            "https://example.com/news",
            "News title",
            "--------------------------------",
        ],
    )
    _write_docx(template_path, ["{{HEADER_TITLE}}", "{{REF_TITLE}}"])
    output_dir.mkdir()

    output_paths = generate_docs(
        schedule_path=schedule_path,
        template_path=template_path,
        output_dir=output_dir,
        filename_prefix="",
        filename_suffix="",
    )

    rendered = Document(str(output_paths[0]))
    header_para = next(
        p for p in rendered.paragraphs if p.text.strip() == "Program - Test Title"
    )
    run = next(run for run in header_para.runs if run.text)
    fonts = run._element.find("w:rPr/w:rFonts", run._element.nsmap)
    assert fonts is not None
    assert fonts.get(qn("w:ascii")) == "Calibri"
    assert fonts.get(qn("w:hAnsi")) == "Calibri"
    assert fonts.get(qn("w:cs")) == "Calibri"
    assert fonts.get(qn("w:eastAsia")) == "新細明體"
