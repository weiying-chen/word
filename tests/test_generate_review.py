from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt

import generate_review
from style_tokens import REVIEW_NOTES_TEXT_SIZE_PT, REVIEW_TEXT_SIZE_PT


def _write_review_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=11, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(0, 2).text = "編輯回饋"
    table.cell(0, 3).text = "主管回饋"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(1, 2).text = ""
    table.cell(1, 3).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "(例行)字幕審稿"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(3, 2).text = ""
    table.cell(3, 3).text = ""
    table.cell(4, 0).text = "日期"
    table.cell(4, 1).text = "臨時工作"
    table.cell(5, 0).text = ""
    table.cell(5, 1).text = ""
    table.cell(5, 2).text = ""
    table.cell(5, 3).text = ""
    table.cell(6, 0).text = "本月工作心得:"
    summary_cell = table.cell(7, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("長度:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("其他工作:")
    summary_cell.add_paragraph("英文新聞: ?篇")
    table.cell(8, 0).text = "之前工作紀錄"
    table.cell(9, 0).text = "日期"
    table.cell(9, 1).text = "工作項目"
    doc.save(path)


def _doc_snapshot(doc: Document) -> dict:
    return {
        "paragraphs": [p.text for p in doc.paragraphs],
        "tables": [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in doc.tables
        ],
    }


def _find_summary_text(table) -> str:
    for row in table.rows:
        cell_text = "\n".join(p.text for p in row.cells[0].paragraphs)
        if "本月總翻譯時數(字幕)" in cell_text:
            return cell_text
    return ""


def _stage_task(
    name: str,
    *,
    start_at: str | None = None,
    work_minutes: int | None = None,
    content_seconds: int | None = None,
    notes: list[str] | None = None,
    children: list[dict] | None = None,
    stage_type: str | None = None,
) -> dict:
    stage: dict[str, object] = {}
    if start_at is not None:
        stage["startAt"] = start_at
    if work_minutes is not None:
        stage["workMinutes"] = work_minutes
    if content_seconds is not None:
        stage["contentSeconds"] = content_seconds
    if stage_type is not None:
        stage["type"] = stage_type

    task = {"name": name, "stages": [stage], "children": children or []}
    if notes is not None:
        task["notes"] = notes
    return task


def test_generate_review_renders_header_fields_from_sources(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2022-11-08T00:00:00.000Z",
                    work_minutes=60,
                    content_seconds=120,
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    assert [p.text for p in out_doc.paragraphs] == [
        "外文編譯中心QCD",
        "姓名: 陳威穎",
        "2022年11月",
        "本月精進目標:",
    ]
    goal_label_runs = out_doc.paragraphs[3].runs
    assert goal_label_runs
    assert all(
        run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in goal_label_runs
    )
    assert all(run.font.size == Pt(REVIEW_TEXT_SIZE_PT) for run in goal_label_runs)


def test_resolve_template_path_accepts_relative_repo_template() -> None:
    resolved = generate_review.resolve_template_path(Path("templates/review_template.docx"))
    assert resolved.exists()


def test_resolve_default_output_path_uses_review_month_suffix() -> None:
    tasks = [
        _stage_task(
            "A",
            start_at="2026-05-08T00:00:00.000Z",
            work_minutes=60,
            content_seconds=120,
        )
    ]

    resolved = generate_review.resolve_output_path(None, tasks)

    assert resolved == Path("output/QCD_Alex_2605.docx")


def test_resolve_default_output_path_uses_created_at_when_start_at_missing() -> None:
    tasks = [
        {
            "name": "A",
            "createdAt": "2026-06-08T00:00:00.000Z",
            "workMinutes": 60,
            "contentSeconds": 120,
            "children": [],
        }
    ]

    resolved = generate_review.resolve_output_path(None, tasks)

    assert resolved == Path("output/QCD_Alex_2606.docx")


def test_resolve_output_path_preserves_explicit_path() -> None:
    explicit = Path("custom/review.docx")
    tasks = [
        _stage_task(
            "A",
            start_at="2026-05-08T00:00:00.000Z",
            work_minutes=60,
            content_seconds=120,
        )
    ]

    resolved = generate_review.resolve_output_path(explicit, tasks)

    assert resolved == explicit


def test_generate_review_populates_regular_translation_rows(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "回眸(中翻英)",
                    start_at="2026-05-08T00:00:00.000Z",
                    work_minutes=240,
                    content_seconds=210,
                    notes=["This is a note"],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    table = out_doc.tables[0]
    assert table.cell(1, 0).text.strip() == "5/8"
    assert table.cell(1, 1).text.strip() == "1.\n回眸(中翻英)\n長度:3分30秒\n實際作業時間:4時"
    assert table.cell(1, 2).text.strip() == "• This is a note"


def test_generate_review_inserts_rows_for_multiple_tasks(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-08T00:00:00.000Z",
                    work_minutes=60,
                    content_seconds=120,
                ),
                _stage_task(
                    "B",
                    start_at="2026-05-09T00:00:00.000Z",
                    work_minutes=120,
                    content_seconds=180,
                ),
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    table = out_doc.tables[0]
    assert table.cell(1, 1).text.strip() == "1.\nA\n長度:2分\n實際作業時間:1時"
    assert table.cell(2, 1).text.strip() == "2.\nB\n長度:3分\n實際作業時間:2時"
    assert table.cell(3, 0).text.strip() == "日期"
    assert table.cell(3, 1).text.strip() == "臨時工作"


def test_generate_review_uses_template_font_for_generated_table_content(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "回眸(中翻英)",
                    start_at="2026-05-08T00:00:00.000Z",
                    work_minutes=240,
                    content_seconds=210,
                    notes=["This is a note"],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    row_cell_runs = out_doc.tables[0].cell(1, 1).paragraphs[0].runs
    date_cell_runs = out_doc.tables[0].cell(1, 0).paragraphs[0].runs
    comment_runs = out_doc.tables[0].cell(1, 2).paragraphs[0].runs
    assert row_cell_runs
    assert date_cell_runs
    assert comment_runs
    assert all(run.font.size == Pt(REVIEW_TEXT_SIZE_PT) for run in row_cell_runs)
    assert all(run.font.size == Pt(REVIEW_TEXT_SIZE_PT) for run in date_cell_runs)
    assert all(run.font.size == Pt(REVIEW_NOTES_TEXT_SIZE_PT) for run in comment_runs)


def test_generate_review_sets_cjk_font_metadata_on_generated_table_content(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "回眸(中翻英)",
                    start_at="2026-05-08T00:00:00.000Z",
                    work_minutes=240,
                    content_seconds=210,
                    notes=["這是一則回饋"],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    generated_run = out_doc.tables[0].cell(1, 1).paragraphs[0].runs[1]
    fonts = generated_run._element.find("w:rPr/w:rFonts", generated_run._element.nsmap)
    assert fonts is not None
    assert fonts.get(qn("w:ascii")) == "Calibri"
    assert fonts.get(qn("w:hAnsi")) == "Calibri"
    assert fonts.get(qn("w:cs")) == "Calibri"
    assert fonts.get(qn("w:eastAsia")) == "新細明體"


def test_generate_review_supports_top_level_tasks_list_with_new_field_names(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=210,
                    notes=["n1"],
                ),
                _stage_task(
                    "B",
                    start_at="2026-05-02T04:05:06Z",
                    work_minutes=120,
                    notes=["n2"],
                ),
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    assert out_doc.paragraphs[2].text == "2026年5月"
    table = out_doc.tables[0]
    assert table.cell(1, 0).text.strip() == "5/1"
    assert table.cell(1, 1).text.strip() == "1.\nA\n長度:3分30秒\n實際作業時間:1時"
    assert table.cell(2, 0).text.strip() == "5/2"
    assert table.cell(2, 1).text.strip() == "2.\nB\n實際作業時間:2時"


def test_generate_review_uses_notes_for_editor_feedback(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                    notes=["note one", "note two"],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    table = out_doc.tables[0]
    assert table.cell(1, 2).text.strip() == "• note one\n• note two"


def test_generate_review_populates_temp_work_from_posts_and_news_children(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "主任務",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                    children=[
                        _stage_task(
                            "POST A",
                            start_at="2026-05-18T11:37:41.273370Z",
                            work_minutes=50,
                            content_seconds=120,
                            stage_type="posts",
                        ),
                        _stage_task(
                            "NEWS B",
                            start_at="2026-05-18T11:37:41.273370Z",
                            work_minutes=40,
                            stage_type="news",
                        ),
                    ],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(
        template_path,
        output_path,
        tasks_json,
    )

    out_doc = Document(output_path)
    table = out_doc.tables[0]
    # After removing 字幕審稿 section: row 2 is 臨時工作 header, row 3 is first temp work row
    assert table.cell(3, 0).text.strip() == "5/18"
    assert table.cell(3, 1).text.strip() == "1.\nPOST A\nFB小編文\n長度:2分\n實際作業時間:50分"
    assert table.cell(4, 0).text.strip() == "5/18"
    assert table.cell(4, 1).text.strip() == "2.\nNEWS B\n英文新聞\n實際作業時間:40分"


def test_generate_review_groups_posts_before_news_in_temp_work(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "主任務",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                    children=[
                        _stage_task(
                            "NEWS A",
                            start_at="2026-05-18T11:37:41.273370Z",
                            work_minutes=40,
                            stage_type="news",
                        ),
                        _stage_task(
                            "POST A",
                            start_at="2026-05-18T12:00:00Z",
                            work_minutes=50,
                            content_seconds=120,
                            stage_type="posts",
                        ),
                        _stage_task(
                            "NEWS B",
                            start_at="2026-05-18T12:30:00Z",
                            work_minutes=30,
                            content_seconds=180,
                            stage_type="news",
                        ),
                        _stage_task(
                            "POST B",
                            start_at="2026-05-18T13:00:00Z",
                            work_minutes=45,
                            content_seconds=90,
                            stage_type="posts",
                        ),
                    ],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    table = out_doc.tables[0]
    assert table.cell(3, 1).text.strip() == "1.\nPOST A\nFB小編文\n長度:2分\n實際作業時間:50分"
    assert table.cell(4, 1).text.strip() == "2.\nPOST B\nFB小編文\n長度:1分30秒\n實際作業時間:45分"
    assert table.cell(5, 1).text.strip() == "3.\nNEWS A\n英文新聞\n實際作業時間:40分"
    assert table.cell(6, 1).text.strip() == "4.\nNEWS B\n英文新聞\n長度:3分\n實際作業時間:30分"


def test_generate_review_removes_subtitle_review_summary_block(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=8, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "(例行)字幕審稿"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "日期"
    table.cell(4, 1).text = "臨時工作"
    table.cell(5, 0).text = ""
    table.cell(5, 1).text = ""
    table.cell(6, 0).text = "本月工作心得:"
    summary_cell = table.cell(7, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("英翻中:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("本月總審稿時數(字幕): (影片長度總和 非工作時數)")
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("英翻中:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("其他工作:")
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    text = "\n".join(p.text for p in out_doc.tables[0].cell(5, 0).paragraphs)
    assert "本月總審稿時數(字幕)" not in text


def test_generate_review_removes_translation_english_to_chinese_line(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("英翻中:")
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    text = ""
    for row in out_doc.tables[0].rows:
        cell_text = "\n".join(p.text for p in row.cells[0].paragraphs)
        if "本月總翻譯時數(字幕)" in cell_text:
            text = cell_text
            break
    assert "本月總翻譯時數(字幕)" in text
    assert "長度:2分" in text
    assert "中翻英:" not in text
    assert "英翻中:" not in text


def test_generate_review_sets_other_work_news_count_from_children(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("其他工作:")
    summary_cell.add_paragraph("英文新聞: ?篇")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("行政工作:")
    summary_cell.add_paragraph("PM選稿子:")
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                    children=[
                        _stage_task("N1", stage_type="news"),
                        _stage_task("N2", stage_type="news"),
                        _stage_task("P1", stage_type="posts"),
                    ],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    text = _find_summary_text(out_doc.tables[0])
    assert "其他工作:" in text
    assert "英文新聞: 2篇" in text
    assert "行政工作:" in text
    assert "PM work" in text
    assert "PM選稿子:" not in text


def test_generate_review_uses_10pt_for_summary_block_lines(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("其他工作:")
    summary_cell.add_paragraph("英文新聞: ?篇")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("行政工作:")
    summary_cell.add_paragraph("PM選稿子:")
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                    children=[
                        _stage_task("N1", stage_type="news"),
                    ],
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    summary_texts = {
        "長度:2分",
        "其他工作:",
        "英文新聞: 1篇",
        "行政工作:",
        "PM work",
    }
    found = {}
    for row in out_doc.tables[0].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()
                if text in summary_texts:
                    found[text] = [
                        run.font.size for run in paragraph.runs if run.text.strip()
                    ]
    assert found.keys() == summary_texts
    for sizes in found.values():
        assert sizes
        assert all(size == Pt(REVIEW_TEXT_SIZE_PT) for size in sizes)


def test_generate_review_removes_meeting_lines_from_work_notes(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    work_notes = table.cell(4, 0)
    work_notes.text = "本月工作心得:"
    work_notes.add_paragraph("工作亮點:")
    work_notes.add_paragraph("健忘次數與事項:")
    work_notes.add_paragraph("同心圓會議:")
    work_notes.add_paragraph("部門內部會議:")
    table.cell(5, 0).text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    text = "\n".join(p.text for p in out_doc.tables[0].cell(4, 0).paragraphs)
    assert "同心圓會議:" not in text
    assert "部門內部會議:" not in text


def test_generate_review_normalizes_translation_summary_spacing(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                )
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    text = "\n".join(p.text for p in out_doc.tables[0].cell(5, 0).paragraphs)
    assert "影片長度總和 非工作時數" not in text
    assert "影片長度總和非工作時數" in text


def test_generate_review_sets_translation_total_length_from_all_tasks_and_children(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("英翻中:")
    doc.save(template_path)

    # total = 3600 + 300 = 3900 seconds => 1時5分
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "A",
                    start_at="2026-05-01T01:02:03Z",
                    content_seconds=3600,
                    children=[
                        _stage_task("C", content_seconds=60, stage_type="posts")
                    ],
                ),
                _stage_task(
                    "B",
                    start_at="2026-05-02T01:02:03Z",
                    content_seconds=300,
                ),
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    text = ""
    for row in out_doc.tables[0].rows:
        cell_text = "\n".join(p.text for p in row.cells[0].paragraphs)
        if "本月總翻譯時數(字幕)" in cell_text:
            text = cell_text
            break
    assert "長度:1時5分" in text
    assert "中翻英:" not in text


def test_generate_review_reads_dates_metrics_and_types_from_stages(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(0, 2).text = "編輯回饋"
    table.cell(0, 3).text = "主管回饋"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(1, 2).text = ""
    table.cell(1, 3).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(3, 2).text = ""
    table.cell(3, 3).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("其他工作:")
    summary_cell.add_paragraph("英文新聞: ?篇")
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                {
                    "name": "Main task",
                    "stages": [
                        {
                            "assignedTo": "Alex",
                            "startAt": "2026-05-20T08:40:00Z",
                            "workMinutes": 432,
                            "contentSeconds": 540,
                        }
                    ],
                    "notes": ["note one"],
                    "children": [
                        {
                            "name": "Post child",
                            "stages": [
                                {
                                    "type": "posts",
                                    "startAt": "2026-05-25T01:00:00Z",
                                    "workMinutes": 50,
                                    "contentSeconds": 120,
                                }
                            ],
                            "notes": ["post note"],
                            "children": [],
                        },
                        {
                            "name": "News child",
                            "stages": [
                                {
                                    "type": "news",
                                    "startAt": "2026-05-27T12:59:38.842799Z",
                                    "workMinutes": 130,
                                    "contentSeconds": 8460,
                                }
                            ],
                            "children": [],
                        },
                    ],
                }
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    assert out_doc.paragraphs[2].text == "2026年5月"
    table = out_doc.tables[0]
    assert table.cell(1, 0).text.strip() == "5/20"
    assert table.cell(1, 1).text.strip() == "1.\nMain task\n長度:9分\n實際作業時間:7時12分"
    assert table.cell(1, 2).text.strip() == "• note one"
    assert table.cell(3, 0).text.strip() == "5/25"
    assert table.cell(3, 1).text.strip() == "1.\nPost child\nFB小編文\n長度:2分\n實際作業時間:50分"
    assert table.cell(3, 2).text.strip() == "• post note"
    summary_text = _find_summary_text(table)
    assert "長度:2時30分" in summary_text
    assert "英文新聞: 1篇" in summary_text


def test_generate_review_supports_top_level_content_seconds_and_stage_extensions(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=6, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(0, 2).text = "編輯回饋"
    table.cell(0, 3).text = "主管回饋"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(1, 2).text = ""
    table.cell(1, 3).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(3, 2).text = ""
    table.cell(3, 3).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("其他工作:")
    summary_cell.add_paragraph("英文新聞: ?篇")
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                {
                    "name": "Main task",
                    "contentSeconds": 540,
                    "stages": [
                        {
                            "name": "translate",
                            "startAt": "2026-05-20T08:40:00Z",
                            "workMinutes": 432,
                            "extensions": [
                                {
                                    "name": "Post child",
                                    "type": "posts",
                                    "startAt": "2026-05-25T01:00:00Z",
                                    "workMinutes": 50,
                                    "contentSeconds": 120,
                                    "notes": ["post note"],
                                },
                                {
                                    "name": "News child",
                                    "type": "news",
                                    "startAt": "2026-05-27T12:59:38.842799Z",
                                    "workMinutes": 130,
                                    "contentSeconds": 8460,
                                },
                            ],
                        }
                    ],
                    "notes": ["note one"],
                }
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    assert out_doc.paragraphs[2].text == "2026年5月"
    table = out_doc.tables[0]
    assert table.cell(1, 0).text.strip() == "5/20"
    assert table.cell(1, 1).text.strip() == "1.\nMain task\n長度:9分\n實際作業時間:7時12分"
    assert table.cell(1, 2).text.strip() == "• note one"
    assert table.cell(3, 0).text.strip() == "5/25"
    assert table.cell(3, 1).text.strip() == "1.\nPost child\nFB小編文\n長度:2分\n實際作業時間:50分"
    assert table.cell(3, 2).text.strip() == "• post note"
    summary_text = _find_summary_text(table)
    assert "長度:2時30分" in summary_text
    assert "英文新聞: 1篇" in summary_text


def test_generate_review_uses_last_task_month_for_header(tmp_path: Path) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                _stage_task(
                    "Later by date",
                    start_at="2026-06-30T23:00:00Z",
                    work_minutes=60,
                    content_seconds=120,
                ),
                _stage_task(
                    "Last in file",
                    start_at="2026-05-01T01:02:03Z",
                    work_minutes=60,
                    content_seconds=120,
                ),
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    assert out_doc.paragraphs[2].text == "2026年5月"


def test_generate_review_splits_current_and_previous_month_subs_and_total(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                {
                    "name": "Previous month task",
                    "type": "subs",
                    "contentSeconds": 600,
                    "notes": ["old note"],
                    "stages": [
                        {
                            "startAt": "2026-05-28T08:00:00Z",
                            "workMinutes": 120,
                            "extensions": [
                                {
                                    "name": "Old news",
                                    "type": "news",
                                    "startAt": "2026-05-29T09:00:00Z",
                                    "workMinutes": 30,
                                    "contentSeconds": 180,
                                }
                            ],
                        }
                    ],
                },
                {
                    "name": "Current month task",
                    "type": "subs",
                    "contentSeconds": 1200,
                    "notes": ["current note"],
                    "stages": [
                        {
                            "startAt": "2026-06-03T08:00:00Z",
                            "workMinutes": 240,
                        }
                    ],
                },
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    table = out_doc.tables[0]
    assert out_doc.paragraphs[2].text == "2026年6月"
    assert table.cell(1, 0).text.strip() == "6/3"
    assert table.cell(1, 1).text.strip() == "1.\nCurrent month task\n長度:20分\n實際作業時間:4時"
    summary_text = ""
    previous_row_idx = None
    for idx, row in enumerate(table.rows):
        cell_text = "\n".join(p.text for p in row.cells[0].paragraphs)
        if "本月總翻譯時數(字幕)" in cell_text:
            summary_text = cell_text
        if row.cells[0].text.strip() == "之前工作紀錄":
            previous_row_idx = idx
    assert "長度:20分" in summary_text
    assert "長度:13分" not in summary_text
    assert previous_row_idx is not None
    assert table.cell(previous_row_idx + 2, 0).text.strip() == "5/28"
    assert (
        table.cell(previous_row_idx + 2, 1).text.strip()
        == "1.\nPrevious month task\n長度:10分\n實際作業時間:2時"
    )
    assert table.cell(previous_row_idx + 2, 2).text.strip() == "• old note"


def test_generate_review_splits_months_using_created_at_when_start_at_missing(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    _write_review_template(template_path)
    tasks_json.write_text(
        json.dumps(
            [
                {
                    "name": "Previous month task",
                    "createdAt": "2026-05-28T08:00:00Z",
                    "workMinutes": 120,
                    "contentSeconds": 600,
                    "notes": ["old note"],
                    "children": [],
                },
                {
                    "name": "Current month task",
                    "createdAt": "2026-06-03T08:00:00Z",
                    "workMinutes": 240,
                    "contentSeconds": 1200,
                    "notes": ["current note"],
                    "children": [],
                },
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)
    out_doc = Document(output_path)
    table = out_doc.tables[0]

    assert "2026年6月" in [p.text for p in out_doc.paragraphs]
    assert table.cell(1, 0).text.strip() == "6/3"
    assert table.cell(1, 1).text.strip() == "1.\nCurrent month task\n長度:20分\n實際作業時間:4時"

    previous_row_idx = None
    for idx, row in enumerate(table.rows):
        if row.cells[0].text.strip() == "之前工作紀錄":
            previous_row_idx = idx
            break

    assert previous_row_idx is not None
    assert table.cell(previous_row_idx + 2, 0).text.strip() == "5/28"
    assert (
        table.cell(previous_row_idx + 2, 1).text.strip()
        == "1.\nPrevious month task\n長度:10分\n實際作業時間:2時"
    )


def test_generate_review_ignores_previous_month_extensions_for_temp_work_and_news_count(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "review_template.docx"
    tasks_json = tmp_path / "tasks.json"
    output_path = tmp_path / "review_output.docx"

    doc = Document()
    doc.add_paragraph("外文編譯中心QCD")
    doc.add_paragraph("姓名: {{NAME}}")
    doc.add_paragraph("{{MONTH}}")
    doc.add_paragraph("本月精進目標:")
    table = doc.add_table(rows=11, cols=4)
    table.cell(0, 0).text = "日期"
    table.cell(0, 1).text = "(例行)字幕翻譯"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "日期"
    table.cell(2, 1).text = "臨時工作"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "本月工作心得:"
    summary_cell = table.cell(5, 0)
    summary_cell.text = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    summary_cell.add_paragraph("中翻英:")
    summary_cell.add_paragraph("")
    summary_cell.add_paragraph("其他工作:")
    summary_cell.add_paragraph("英文新聞: ?篇")
    table.cell(6, 0).text = "之前工作紀錄"
    table.cell(7, 0).text = "日期"
    table.cell(7, 1).text = "工作項目"
    doc.save(template_path)

    tasks_json.write_text(
        json.dumps(
            [
                {
                    "name": "Previous month task",
                    "type": "subs",
                    "contentSeconds": 600,
                    "stages": [
                        {
                            "startAt": "2026-05-28T08:00:00Z",
                            "workMinutes": 120,
                            "extensions": [
                                {
                                    "name": "Old post",
                                    "type": "posts",
                                    "startAt": "2026-05-29T09:00:00Z",
                                    "workMinutes": 50,
                                    "contentSeconds": 120,
                                },
                                {
                                    "name": "Old news",
                                    "type": "news",
                                    "startAt": "2026-05-29T10:00:00Z",
                                    "workMinutes": 30,
                                    "contentSeconds": 180,
                                },
                            ],
                        }
                    ],
                },
                {
                    "name": "Current month task",
                    "type": "subs",
                    "contentSeconds": 1200,
                    "stages": [
                        {
                            "startAt": "2026-06-03T08:00:00Z",
                            "workMinutes": 240,
                            "extensions": [
                                {
                                    "name": "Current post",
                                    "type": "posts",
                                    "startAt": "2026-06-04T09:00:00Z",
                                    "workMinutes": 50,
                                    "contentSeconds": 120,
                                },
                                {
                                    "name": "Current news",
                                    "type": "news",
                                    "startAt": "2026-06-04T10:00:00Z",
                                    "workMinutes": 30,
                                    "contentSeconds": 180,
                                },
                            ],
                        }
                    ],
                },
            ],
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    generate_review.generate_review(template_path, output_path, tasks_json)

    out_doc = Document(output_path)
    table = out_doc.tables[0]
    assert table.cell(3, 1).text.strip() == "1.\nCurrent post\nFB小編文\n長度:2分\n實際作業時間:50分"
    assert table.cell(4, 1).text.strip() == "2.\nCurrent news\n英文新聞\n長度:3分\n實際作業時間:30分"
    previous_row_idx = None
    for idx, row in enumerate(table.rows):
        if row.cells[0].text.strip() == "之前工作紀錄":
            previous_row_idx = idx
            break
    assert previous_row_idx is not None
    summary_text = _find_summary_text(table)
    assert "長度:23分" in summary_text
    assert "英文新聞: 1篇" in summary_text
