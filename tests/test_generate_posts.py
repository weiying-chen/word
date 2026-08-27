from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from generate_posts import generate_post, parse_post_text


POST_TEXT = """TITLE: Da Ai Journal - Staying Young at 105 (大愛全紀實 - 人生歌未央 [1])

VIDEO_URL: https://www.youtube.com/watch?v=example

BODY:

Lin You-mao is 105 years old—and he's still playing badminton.

#DaAiJournal #StayingYoungAt105

林友茂今年105歲了，至今仍在打羽球。

#大愛全紀實 #人生歌未央

SOURCES:

source.docx

00:01:31:22\t00:01:34:22\t中文字幕
English translation.
"""


def test_parse_completed_post_text() -> None:
    post = parse_post_text(POST_TEXT)

    assert post["title"] == (
        "Da Ai Journal - Staying Young at 105 (大愛全紀實 - 人生歌未央 [1])"
    )
    assert post["video_url"] == "https://www.youtube.com/watch?v=example"
    assert post["post_en"] == (
        "Lin You-mao is 105 years old—and he's still playing badminton."
    )
    assert post["hashtags_en"] == "#DaAiJournal #StayingYoungAt105"
    assert post["post_zh"] == "林友茂今年105歲了，至今仍在打羽球。"
    assert post["hashtags_zh"] == "#大愛全紀實 #人生歌未央"
    assert post["sources"] == (
        "source.docx\n\n"
        "00:01:31:22\t00:01:34:22\t中文字幕\n"
        "English translation."
    )


def test_parse_completed_post_text_requires_valid_video_url() -> None:
    invalid = POST_TEXT.replace(
        "VIDEO_URL: https://www.youtube.com/watch?v=example",
        "VIDEO_URL: interview.mp4",
    )

    try:
        parse_post_text(invalid)
    except ValueError as exc:
        assert str(exc) == "[error] VIDEO_URL must be an HTTP or HTTPS URL."
    else:
        raise AssertionError("Expected invalid VIDEO_URL to be rejected.")


def test_generate_completed_post_uses_post_template_without_draft_header(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "post.txt"
    output_path = tmp_path / "post.docx"
    input_path.write_text(POST_TEXT, encoding="utf-8")

    with patch(
        "generate_posts.fetch_youtube_video_metadata",
        return_value=(
            "Actual YouTube Video Title",
            "English video summary.",
            "中文影片摘要。",
        ),
    ):
        generate_post(
            input_path=input_path,
            template_path=Path("templates/post_template.docx"),
            output_path=output_path,
        )

    texts = [paragraph.text for paragraph in Document(output_path).paragraphs]
    assert "標題" not in texts
    assert not any(text.startswith("9/20(日)") for text in texts)
    assert texts[0] == (
        "Da Ai Journal - Staying Young at 105 (大愛全紀實 - 人生歌未央 [1])"
    )
    assert texts[1] == "https://www.youtube.com/watch?v=example"
    assert texts[3] == (
        "Lin You-mao is 105 years old—and he's still playing badminton."
    )
    assert sum(
        "Da Ai Journal - Staying Young at 105" in text for text in texts
    ) == 1
    assert "Actual YouTube Video Title" in texts
    assert "source.docx" in texts
    assert "00:01:31:22\t00:01:34:22\t中文字幕" in texts
    assert "English translation." in texts
    assert "English video summary." not in texts
    assert "中文影片摘要。" in texts

    timestamp = next(
        paragraph
        for paragraph in Document(output_path).paragraphs
        if paragraph.text.startswith("00:01:31:22")
    )
    assert timestamp.paragraph_format.left_indent is not None
    assert timestamp.text.count("\t") == 2
    assert all(run.font.size == Pt(10) for run in timestamp.runs if run.text)
    assert all(run.font.highlight_color is not None for run in timestamp.runs if run.text)


def test_generate_completed_post_renders_body_as_real_paragraphs(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "post.txt"
    output_path = tmp_path / "post.docx"
    input_path.write_text(
        POST_TEXT.replace(
            "Lin You-mao is 105 years old—and he's still playing badminton.",
            "English paragraph one.\n\nEnglish paragraph two.",
        ).replace(
            "林友茂今年105歲了，至今仍在打羽球。",
            "中文第一段。\n\n中文第二段。",
        ),
        encoding="utf-8",
    )

    with patch(
        "generate_posts.fetch_youtube_video_metadata",
        return_value=("Video title", "", "中文影片摘要。"),
    ):
        generate_post(
            input_path=input_path,
            template_path=Path("templates/post_template.docx"),
            output_path=output_path,
        )

    doc = Document(output_path)
    texts = [paragraph.text for paragraph in doc.paragraphs]
    assert "English paragraph one." in texts
    assert "English paragraph two." in texts
    assert "中文第一段。" in texts
    assert "中文第二段。" in texts
    for text in {
        "English paragraph one.",
        "English paragraph two.",
        "中文第一段。",
        "中文第二段。",
    }:
        paragraph = next(p for p in doc.paragraphs if p.text == text)
        assert not paragraph._p.findall(".//" + qn("w:br"))
        assert not paragraph._p.findall(".//" + qn("w:cr"))
