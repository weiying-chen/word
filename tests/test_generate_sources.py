import json
from pathlib import Path

from docx import Document

from generate_sources import (
    generate_sources,
    resolve_default_episodes_json,
    resolve_default_sources_dir,
)


def _write_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("")
    doc.save(str(path))


def test_resolve_default_episodes_json_uses_current_folder_file(tmp_path: Path) -> None:
    episodes_path = tmp_path / "episodes.json"
    episodes_path.write_text("[]", encoding="utf-8")

    assert resolve_default_episodes_json(tmp_path) == episodes_path


def test_resolve_default_sources_dir_prefers_subtitles(tmp_path: Path) -> None:
    sources_dir = tmp_path / "sources"
    subtitles_dir = tmp_path / "subtitles"
    sources_dir.mkdir()
    subtitles_dir.mkdir()

    assert resolve_default_sources_dir(tmp_path) == subtitles_dir


def test_resolve_default_sources_dir_falls_back_to_sources(tmp_path: Path) -> None:
    sources_dir = tmp_path / "sources"
    sources_dir.mkdir()

    assert resolve_default_sources_dir(tmp_path) == sources_dir


def test_generate_sources_skips_when_subtitle_file_missing(tmp_path: Path) -> None:
    episodes_path = tmp_path / "episodes.json"
    template_path = tmp_path / "sources_template.docx"
    sources_dir = tmp_path / "sources"
    output_dir = tmp_path / "output"
    sources_dir.mkdir()
    output_dir.mkdir()
    _write_template(template_path)

    episodes = [
        {
            "epId": "6797",
            "titleZh": "肺腺癌先禮後兵",
            "ytId": "P0uiRM2no18",
            "youtubeUrl": "https://www.youtube.com/watch?v=P0uiRM2no18",
            "youtubeTitle": "【大愛醫生館】 肺腺癌先禮後兵 20260520",
            "youtubeDescription": "第一行摘要\n第二行",
        }
    ]
    episodes_path.write_text(json.dumps(episodes, ensure_ascii=False), encoding="utf-8")

    result = generate_sources(
        episodes_json=episodes_path,
        template_path=template_path,
        sources_dir=sources_dir,
        output_dir=output_dir,
    )

    assert result["generated"] == 0
    assert result["skipped"] == 1
    assert not list(output_dir.glob("*.docx"))


def test_generate_sources_writes_docx_for_existing_subtitle_file(tmp_path: Path) -> None:
    episodes_path = tmp_path / "episodes.json"
    template_path = tmp_path / "sources_template.docx"
    sources_dir = tmp_path / "sources"
    output_dir = tmp_path / "output"
    sources_dir.mkdir()
    output_dir.mkdir()
    _write_template(template_path)

    (sources_dir / "大愛醫生館第6797集_ch_肺腺癌先禮後兵.txt").write_text(
        "00:00:01:00\t00:00:03:00\t第一句\nSecond line",
        encoding="utf-8",
    )

    episodes = [
        {
            "epId": "6797",
            "titleZh": "肺腺癌先禮後兵",
            "ytId": "P0uiRM2no18",
            "youtubeUrl": "https://www.youtube.com/watch?v=P0uiRM2no18",
            "youtubeTitle": "【大愛醫生館】 肺腺癌先禮後兵 20260520",
            "youtubeDescription": "五十歲男性長期吸菸、慢性咳嗽。",
            "descriptionLastTimestampLine": "07:27｜肺腺癌先禮後兵",
        }
    ]
    episodes_path.write_text(json.dumps(episodes, ensure_ascii=False), encoding="utf-8")

    result = generate_sources(
        episodes_json=episodes_path,
        template_path=template_path,
        sources_dir=sources_dir,
        output_dir=output_dir,
    )

    assert result["generated"] == 1
    files = list(output_dir.glob("*.docx"))
    assert len(files) == 1
    doc = Document(str(files[0]))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "【大愛醫生館】 肺腺癌先禮後兵 20260520"
    assert texts[1] == "https://www.youtube.com/watch?v=P0uiRM2no18"
    assert texts[2] == "00:00-07:27 (7分27秒)"
    assert texts[3] == "五十歲男性長期吸菸、慢性咳嗽。"
    assert texts[4] == "00:00:01:00\t00:00:03:00\t第一句"
    assert texts[5] == "Second line"


def test_generate_sources_reads_utf16_subtitle_file(tmp_path: Path) -> None:
    episodes_path = tmp_path / "episodes.json"
    template_path = tmp_path / "sources_template.docx"
    sources_dir = tmp_path / "sources"
    output_dir = tmp_path / "output"
    sources_dir.mkdir()
    output_dir.mkdir()
    _write_template(template_path)

    (sources_dir / "大愛醫生館第6797集_ch_肺腺癌先禮後兵.txt").write_text(
        "第一行\n第二行",
        encoding="utf-16",
    )

    episodes = [
        {
            "epId": "6797",
            "titleZh": "肺腺癌先禮後兵",
            "ytId": "P0uiRM2no18",
            "youtubeUrl": "https://www.youtube.com/watch?v=P0uiRM2no18",
            "youtubeTitle": "【大愛醫生館】 肺腺癌先禮後兵 20260520",
            "youtubeDescription": "摘要。",
            "descriptionLastTimestampLine": "07:27｜肺腺癌先禮後兵",
        }
    ]
    episodes_path.write_text(json.dumps(episodes, ensure_ascii=False), encoding="utf-8")

    result = generate_sources(
        episodes_json=episodes_path,
        template_path=template_path,
        sources_dir=sources_dir,
        output_dir=output_dir,
    )
    assert result["generated"] == 1


def test_generate_sources_falls_back_to_hardcoded_timestamp_without_last_line(
    tmp_path: Path,
) -> None:
    episodes_path = tmp_path / "episodes.json"
    template_path = tmp_path / "sources_template.docx"
    sources_dir = tmp_path / "sources"
    output_dir = tmp_path / "output"
    sources_dir.mkdir()
    output_dir.mkdir()
    _write_template(template_path)

    (sources_dir / "大愛醫生館第6797集_ch_肺腺癌先禮後兵.txt").write_text(
        "line1",
        encoding="utf-8",
    )

    episodes = [
        {
            "epId": "6797",
            "titleZh": "肺腺癌先禮後兵",
            "ytId": "P0uiRM2no18",
            "youtubeUrl": "https://www.youtube.com/watch?v=P0uiRM2no18",
            "youtubeTitle": "【大愛醫生館】 肺腺癌先禮後兵 20260520",
            "youtubeDescription": "摘要。",
            "descriptionLastTimestampLine": "",
        }
    ]
    episodes_path.write_text(json.dumps(episodes, ensure_ascii=False), encoding="utf-8")

    generate_sources(
        episodes_json=episodes_path,
        template_path=template_path,
        sources_dir=sources_dir,
        output_dir=output_dir,
    )

    doc = Document(str(next(output_dir.glob("*.docx"))))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[2] == "07:27-09:20 (1分53秒)"


def test_generate_sources_uses_two_star_markers_for_range_and_highlights(
    tmp_path: Path,
) -> None:
    episodes_path = tmp_path / "episodes.json"
    template_path = tmp_path / "sources_template.docx"
    sources_dir = tmp_path / "sources"
    output_dir = tmp_path / "output"
    sources_dir.mkdir()
    output_dir.mkdir()
    _write_template(template_path)

    (sources_dir / "大愛醫生館第6797集_ch_肺腺癌先禮後兵.txt").write_text(
        "00:00:01:00\t00:00:03:00\tA\n"
        "00:00:05:00\t00:00:07:00\tB *\n"
        "00:00:08:00\t00:00:10:00\tC\n"
        "00:00:12:00\t00:00:14:00\tD *\n"
        "00:00:15:00\t00:00:17:00\tE\n",
        encoding="utf-8",
    )

    episodes = [
        {
            "epId": "6797",
            "titleZh": "肺腺癌先禮後兵",
            "ytId": "P0uiRM2no18",
            "youtubeUrl": "https://www.youtube.com/watch?v=P0uiRM2no18",
            "youtubeTitle": "【大愛醫生館】 肺腺癌先禮後兵 20260520",
            "youtubeDescription": "摘要。",
            "descriptionLastTimestampLine": "",
        }
    ]
    episodes_path.write_text(json.dumps(episodes, ensure_ascii=False), encoding="utf-8")

    generate_sources(
        episodes_json=episodes_path,
        template_path=template_path,
        sources_dir=sources_dir,
        output_dir=output_dir,
    )

    doc = Document(str(next(output_dir.glob("*.docx"))))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[2] == "00:05-00:14 (0分9秒)"

    subtitle_paragraphs = [p for p in doc.paragraphs if "\t" in p.text]
    assert subtitle_paragraphs[0].text.endswith("\tA")
    assert subtitle_paragraphs[1].text.endswith("\tB")
    assert subtitle_paragraphs[2].text.endswith("\tC")
    assert subtitle_paragraphs[3].text.endswith("\tD")
    assert subtitle_paragraphs[4].text.endswith("\tE")

    def _is_yellow(p):
        return any(r.font.highlight_color is not None for r in p.runs)

    assert _is_yellow(subtitle_paragraphs[0]) is False
    assert _is_yellow(subtitle_paragraphs[1]) is True
    assert _is_yellow(subtitle_paragraphs[2]) is True
    assert _is_yellow(subtitle_paragraphs[3]) is True
    assert _is_yellow(subtitle_paragraphs[4]) is False


def test_generate_sources_does_not_highlight_touching_boundary_lines(
    tmp_path: Path,
) -> None:
    episodes_path = tmp_path / "episodes.json"
    template_path = tmp_path / "sources_template.docx"
    sources_dir = tmp_path / "sources"
    output_dir = tmp_path / "output"
    sources_dir.mkdir()
    output_dir.mkdir()
    _write_template(template_path)

    (sources_dir / "大愛醫生館第6797集_ch_肺腺癌先禮後兵.txt").write_text(
        "00:09:24:20\t00:09:25:24\tbefore\n"
        "00:09:25:24\t00:09:26:19\tstart *\n"
        "00:09:26:19\t00:09:27:17\tmiddle\n"
        "00:09:28:00\t00:09:29:00\tend *\n"
        "00:09:29:00\t00:09:30:00\tafter\n",
        encoding="utf-8",
    )

    episodes = [
        {
            "epId": "6797",
            "titleZh": "肺腺癌先禮後兵",
            "ytId": "P0uiRM2no18",
            "youtubeUrl": "https://www.youtube.com/watch?v=P0uiRM2no18",
            "youtubeTitle": "【大愛醫生館】 肺腺癌先禮後兵 20260520",
            "youtubeDescription": "摘要。",
            "descriptionLastTimestampLine": "",
        }
    ]
    episodes_path.write_text(json.dumps(episodes, ensure_ascii=False), encoding="utf-8")

    generate_sources(
        episodes_json=episodes_path,
        template_path=template_path,
        sources_dir=sources_dir,
        output_dir=output_dir,
    )

    doc = Document(str(next(output_dir.glob("*.docx"))))
    subtitle_paragraphs = [p for p in doc.paragraphs if "\t" in p.text]

    def _is_yellow(p):
        return any(r.font.highlight_color is not None for r in p.runs)

    assert _is_yellow(subtitle_paragraphs[0]) is False
    assert _is_yellow(subtitle_paragraphs[1]) is True
    assert _is_yellow(subtitle_paragraphs[2]) is True
    assert _is_yellow(subtitle_paragraphs[3]) is True
    assert _is_yellow(subtitle_paragraphs[4]) is False


def test_generate_sources_respects_frame_boundary_after_end_marker(
    tmp_path: Path,
) -> None:
    episodes_path = tmp_path / "episodes.json"
    template_path = tmp_path / "sources_template.docx"
    sources_dir = tmp_path / "sources"
    output_dir = tmp_path / "output"
    sources_dir.mkdir()
    output_dir.mkdir()
    _write_template(template_path)

    (sources_dir / "大愛醫生館第6797集_ch_肺腺癌先禮後兵.txt").write_text(
        "00:06:21:15\t00:06:24:22\tstart *\n"
        "00:09:16:06\t00:09:18:01\tend *\n"
        "00:09:18:01\t00:09:18:17\tafter\n",
        encoding="utf-8",
    )

    episodes = [
        {
            "epId": "6797",
            "titleZh": "肺腺癌先禮後兵",
            "ytId": "P0uiRM2no18",
            "youtubeUrl": "https://www.youtube.com/watch?v=P0uiRM2no18",
            "youtubeTitle": "【大愛醫生館】 肺腺癌先禮後兵 20260520",
            "youtubeDescription": "摘要。",
            "descriptionLastTimestampLine": "",
        }
    ]
    episodes_path.write_text(json.dumps(episodes, ensure_ascii=False), encoding="utf-8")

    generate_sources(
        episodes_json=episodes_path,
        template_path=template_path,
        sources_dir=sources_dir,
        output_dir=output_dir,
    )

    doc = Document(str(next(output_dir.glob("*.docx"))))
    subtitle_paragraphs = [p for p in doc.paragraphs if "\t" in p.text]

    def _is_yellow(p):
        return any(r.font.highlight_color is not None for r in p.runs)

    assert _is_yellow(subtitle_paragraphs[0]) is True
    assert _is_yellow(subtitle_paragraphs[1]) is True
    assert _is_yellow(subtitle_paragraphs[2]) is False
