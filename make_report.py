#!/usr/bin/env python3

from datetime import date

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # blue color
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
doc.add_heading("Daily Report", level=1)

# Date line with highlight
p = doc.add_paragraph()
p.add_run("Date: ")

date_run = p.add_run(str(date.today()))
date_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# Body text
doc.add_paragraph("This document was generated automatically.")

# Hyperlink
link_p = doc.add_paragraph("Reference: ")
add_hyperlink(
    link_p,
    "Project Dashboard",
    "https://example.com",
)

doc.save("daily_report.docx")
