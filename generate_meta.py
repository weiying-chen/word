#!/usr/bin/env python3

from __future__ import annotations

import argparse
import re
import warnings
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
from docx.shared import Inches

from generate_subs import fix_docx_namespaces, normalize_input_text


TITLE_PLACEHOLDER = "{{TITLE_EN}}"
PEOPLE_PLACEHOLDER = "{{PEOPLE}}"
OVERVIEW_PLACEHOLDER = "{{OVERVIEW_EN}}"
HIGHLIGHT_LABELS = {"重點標", "名字職銜", "YT簡介"}
TITLE_KEY = "TITLE"
OVERVIEW_KEY = "OVERVIEW"
PEOPLE_KEY = "PEOPLE"
CJK_RE = re.compile(r"[\u4e00-\u9fff]")
EN_NAME_TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z.\s'-]*")
EN_NAME_HINT_RE = re.compile(
    r'^[\s\d.,，．。:：;；!?！？~\-–—秒分]*'
    r'([A-Za-zÀ-ÖØ-öø-ÿĀ-žḀ-ỹ][A-Za-zÀ-ÖØ-öø-ÿĀ-žḀ-ỹ.\s"“”\'‘’\-]*[A-Za-zÀ-ÖØ-öø-ÿĀ-žḀ-ỹ])'
    r'(?:\s*[\u4e00-\u9fff].*)?$'
)
EN_NAME_VALUE_RE = re.compile(
    r"[A-Za-zÀ-ÖØ-öø-ÿĀ-žḀ-ỹ][A-Za-zÀ-ÖØ-öø-ÿĀ-žḀ-ỹ.\s\"'“”‘’\-]*[A-Za-zÀ-ÖØ-öø-ÿĀ-žḀ-ỹ.]"
)
ALLOWED_KEYS = {
    "TITLE_TEXT",
    "SUMMARY",
    TITLE_KEY,
    OVERVIEW_KEY,
    PEOPLE_KEY,
    "BODY",
}


def _has_section_key(path: Path, key: str) -> bool:
    section_re = re.compile(rf"^\s*{re.escape(key)}\s*:", re.IGNORECASE)
    for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        if section_re.match(raw):
            return True
    return False


def _contains_cjk(text: str) -> bool:
    return bool(CJK_RE.search(text))


def _clean_super_line(text: str) -> str:
    cleaned = text.strip()
    if cleaned.endswith("//"):
        cleaned = cleaned[:-2].rstrip()
    return cleaned


def _extract_english_name_hint(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return ""
    if stripped[0] not in {"(", "（"} or stripped[-1] not in {")", "）"}:
        return ""

    inner = stripped[1:-1]
    match = EN_NAME_HINT_RE.match(inner.strip())
    if match:
        name = match.group(1).strip().rstrip(" .,;:-")
        return (
            name.replace("“", '"')
            .replace("”", '"')
            .replace("‘", "'")
            .replace("’", "'")
        )

    # Fallback: support cues like "(SB) (Anabel) (17秒)".
    chunks = re.findall(r"[（(]([^（）()]*)[）)]", stripped)
    for chunk in chunks:
        candidate = chunk.strip().rstrip(" .,;:-")
        if candidate.isupper() and len(candidate) <= 3:
            continue
        if EN_NAME_VALUE_RE.fullmatch(candidate):
            return (
                candidate.replace("“", '"')
                .replace("”", '"')
                .replace("‘", "'")
                .replace("’", "'")
            )
    return ""


def _parse_super(lines: list[str]) -> dict:
    role_zh = ""
    name_zh = ""
    quotes_zh: list[str] = []
    if lines:
        header = lines[0]
        if "│" in header:
            role_zh, name_zh = [part.strip() for part in header.split("│", 1)]
        elif "｜" in header:
            role_zh, name_zh = [part.strip() for part in header.split("｜", 1)]
        else:
            name_zh = header.strip()
        if len(lines) > 1:
            quotes_zh = [line for line in lines[1:] if line]
    return {
        "role_zh": role_zh,
        "name_zh": name_zh,
        "quotes_zh": quotes_zh,
    }


def _parse_meta_people_blocks(text: str) -> list[dict[str, str]]:
    if not text.strip():
        return []

    blocks: list[list[str]] = []
    current: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)
    if current:
        blocks.append(current)

    entries: list[dict[str, str]] = []
    for block in blocks:
        if not block:
            continue
        label_zh = block[0].strip()
        name_zh = ""
        role_zh = ""
        if "｜" in label_zh:
            role_zh, name_zh = [part.strip() for part in label_zh.split("｜", 1)]
            if role_zh and not name_zh:
                label_zh = role_zh
        else:
            name_zh = label_zh.strip()

        entries.append(
            {
                "label_zh": label_zh,
                "name_zh": name_zh,
                "role_zh": role_zh,
                "name_en": block[1].strip() if len(block) > 1 else "",
                "role_en": block[2].strip() if len(block) > 2 else "",
                "org_en": block[3].strip() if len(block) > 3 else "",
            }
        )
    return entries


def _merge_meta_people_overrides(
    people: list[dict[str, str]],
    overrides: list[dict[str, str]],
) -> list[dict[str, str]]:
    if not overrides:
        return people

    merged = [dict(person) for person in people]
    for person in merged:
        role_zh = person.get("role_zh", "").strip()
        name_zh = person.get("name_zh", "").strip()
        label_zh = f"{role_zh}｜{name_zh}" if role_zh and name_zh else (role_zh or name_zh)

        match = next(
            (
                entry
                for entry in overrides
                if entry.get("label_zh", "").strip()
                and entry.get("label_zh", "").strip() == label_zh
            ),
            None,
        )
        if match is None and name_zh:
            match = next(
                (
                    entry
                    for entry in overrides
                    if entry.get("name_zh", "").strip() == name_zh
                ),
                None,
            )
        if match is None and role_zh:
            # Support role-only override labels, matched by English name line.
            # Example:
            #   慈濟人醫會醫師
            #   Timothy Yu
            person_name_en = person.get("name_en", "").strip()
            match = next(
                (
                    entry
                    for entry in overrides
                    if entry.get("label_zh", "").strip() == role_zh
                    and entry.get("name_en", "").strip()
                    and entry.get("name_en", "").strip().casefold()
                    in {name_zh.casefold(), person_name_en.casefold()}
                ),
                None,
            )
        if match is None and role_zh and not name_zh:
            # Support nameless role overrides written as "角色｜" in meta.txt.
            # Example:
            #   恩佳大愛村民｜
            #   Resident
            #   Ndeja Tzu Chi Great Love Village
            match = next(
                (
                    entry
                    for entry in overrides
                    if entry.get("role_zh", "").strip() == role_zh
                    and not entry.get("name_zh", "").strip()
                ),
                None,
            )
        if match is None:
            continue

        label_override = match.get("label_zh", "").strip()
        if label_override:
            person["label_zh"] = label_override
        for key in ("name_en", "role_en", "org_en"):
            value = match.get(key, "").strip()
            if value:
                person[key] = value

    return merged


def _decode_input_text(path: Path) -> tuple[str, str, bool]:
    raw = path.read_bytes()
    tried_encodings: list[str] = []

    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig"), "utf-8-sig", False

    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16"), "utf-16", False

    for encoding in ("utf-8", "big5", "cp950", "gb18030", "cp1252"):
        tried_encodings.append(encoding)
        try:
            return raw.decode(encoding), encoding, encoding != "utf-8"
        except UnicodeDecodeError:
            continue

    tried = ", ".join(tried_encodings)
    raise UnicodeError(
        f"Unable to decode input file '{path}' with supported encodings: {tried}"
    )


def _parse_news_payload(path: Path, *, allow_body_fallback: bool = False) -> dict[str, str]:
    text, encoding_used, used_fallback = _decode_input_text(path)
    if used_fallback:
        warnings.warn(
            f"Using fallback encoding '{encoding_used}' for {path}; rewriting as UTF-8.",
            stacklevel=2,
        )
        path.write_text(text, encoding="utf-8")

    lines = text.splitlines()
    data: dict[str, str] = {}
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        if ":" not in raw_line:
            idx += 1
            continue

        key, value = raw_line.split(":", 1)
        key = key.lstrip("\ufeff").strip().upper()
        value = value.lstrip()

        if key not in ALLOWED_KEYS:
            idx += 1
            continue

        if key in {"SUMMARY", PEOPLE_KEY, "BODY"}:
            collected: list[str] = []
            if value:
                collected.append(value)
            idx += 1
            while idx < len(lines):
                next_line = lines[idx]
                if ":" in next_line:
                    next_key = next_line.split(":", 1)[0].strip().upper()
                    if next_key in ALLOWED_KEYS:
                        break
                collected.append(next_line)
                idx += 1
            data[key] = normalize_input_text("\n".join(collected).rstrip())
            continue

        data[key] = normalize_input_text(value)
        idx += 1

    data.setdefault("TITLE_TEXT", "")
    data.setdefault("SUMMARY", "")
    data.setdefault(TITLE_KEY, "")
    data.setdefault(OVERVIEW_KEY, "")
    data.setdefault(PEOPLE_KEY, "")
    data.setdefault("BODY", "")
    if allow_body_fallback and not data["BODY"] and not any(
        data[key]
        for key in (
            "TITLE_TEXT",
            "SUMMARY",
            TITLE_KEY,
            OVERVIEW_KEY,
            PEOPLE_KEY,
        )
    ):
        data["BODY"] = normalize_input_text(text.rstrip())
    return data


def parse_input(path: Path, meta_path: Path | None = None) -> dict[str, object]:
    if path.suffix.lower() != ".txt":
        raise ValueError(f"Unsupported input format: {path}")

    data = _parse_news_payload(path, allow_body_fallback=True)
    meta_has_people_section = False
    if meta_path is not None:
        if meta_path.suffix.lower() != ".txt":
            raise ValueError(f"Unsupported meta input format: {meta_path}")
        meta_has_people_section = _has_section_key(meta_path, PEOPLE_KEY)
        meta_data = _parse_news_payload(meta_path, allow_body_fallback=False)
        data = {**data, **{k: v for k, v in meta_data.items() if k != "BODY" and v}}
    body_lines = data.get("BODY", "").splitlines()

    narration_zh: list[str] = []
    supers: list[dict[str, object]] = []
    report_zh: list[str] = []
    english_names: list[str] = []
    super_lines: list[str] = []
    in_super = False
    in_report = False

    for raw in body_lines:
        line = raw.strip()
        english_name = _extract_english_name_hint(line)
        if english_name:
            english_names.append(english_name)
            continue
        if line == "/*SUPER:":
            in_super = True
            super_lines = []
            continue
        if line.startswith("/*REPORT"):
            in_report = True
            continue
        if in_report:
            if line.startswith("*/"):
                in_report = False
            else:
                report_zh.append(_clean_super_line(line))
            continue
        if in_super:
            if line.startswith("*/"):
                supers.append(_parse_super(super_lines))
                in_super = False
            else:
                super_lines.append(_clean_super_line(line))
            continue

        if not _contains_cjk(line):
            continue
        if re.fullmatch(r"[\d_]+", line):
            continue
        if re.fullmatch(r"\(?\s*\d+.*\)?", line) and not _contains_cjk(line):
            continue
        if re.fullmatch(r"\(\s*NS\s*\)", line):
            continue
        if re.fullmatch(r"\(\s*\d+\s+[A-Za-z]+\s*\)", line):
            continue

        line = re.sub(r"^\(\s*[^)]*\)\s*", "", line).strip()
        if not line:
            continue
        narration_zh.append(line)

    people: list[dict[str, str]] = [
        {
            "name_zh": str(s.get("name_zh", "")),
            "name_en": "",
            "role_zh": str(s.get("role_zh", "")),
            "role_en": "",
            "org_en": "",
        }
        for s in supers
    ]
    for idx, en_name in enumerate(english_names):
        if idx >= len(people):
            break
        people[idx]["name_en"] = en_name

    summary = data.get("SUMMARY", "").splitlines()
    meta_people_text = data.get(PEOPLE_KEY, "")
    overrides = _parse_meta_people_blocks(meta_people_text)
    people = _merge_meta_people_overrides(people, overrides)
    if meta_path is not None and not meta_has_people_section:
        people = []
    return {
        "title_zh": data.get("TITLE_TEXT", ""),
        "summary_zh": summary[0] if summary else "",
        "narration_zh": narration_zh,
        "supers_zh": supers,
        "report_zh": report_zh,
        "people": people,
        "title_en": data.get(TITLE_KEY, ""),
        "overview_en": data.get(OVERVIEW_KEY, ""),
    }


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def find_paragraph_by_text(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def replace_multiline(paragraph: Paragraph, lines: list[str]) -> None:
    if not lines:
        remove_paragraph(paragraph)
        return
    clear_paragraph(paragraph)
    paragraph.add_run(lines[0])
    current = paragraph
    for line in lines[1:]:
        current = insert_paragraph_after(current, line)


def replace_or_remove_paragraph_text(paragraph: Paragraph | None, text: str) -> None:
    if paragraph is None:
        return
    if text:
        paragraph.text = text
    else:
        next_element = paragraph._p.getnext()
        if next_element is not None and next_element.tag.endswith("}p"):
            next_paragraph = Paragraph(next_element, paragraph._parent)
            if not next_paragraph.text.strip():
                remove_paragraph(paragraph)
                return
        paragraph.text = ""


def ensure_single_blank_before_label(doc: Document, label_text: str) -> None:
    while True:
        label_para = find_paragraph_by_text(doc, label_text)
        if label_para is None:
            return
        paragraphs = doc.paragraphs
        idx = next((i for i, p in enumerate(paragraphs) if p._p is label_para._p), -1)
        if idx < 2:
            return
        if paragraphs[idx - 1].text.strip() or paragraphs[idx - 2].text.strip():
            return
        remove_paragraph(paragraphs[idx - 2])


def _label_without_repeated_english_name(
    label_zh: str,
    *,
    role_zh: str,
    name_en: str,
) -> str:
    if not label_zh or not role_zh or not name_en:
        return label_zh
    if "｜" not in label_zh:
        return label_zh

    label_role, label_name = [part.strip() for part in label_zh.split("｜", 1)]
    if label_role != role_zh.strip():
        return label_zh
    if label_name.strip().casefold() != name_en.strip().casefold():
        return label_zh
    return role_zh.strip()


def build_people_lines(people: list[dict]) -> list[str]:
    lines: list[str] = []
    if people:
        # Keep one blank line between the "名字職銜" label and the first person entry.
        lines.append("")
    for idx, person in enumerate(people):
        role_zh = person.get("role_zh", "").strip()
        name_zh = person.get("name_zh", "").strip()
        name_en = person.get("name_en", "").strip()
        label_zh = person.get("label_zh")
        if not label_zh:
            if role_zh and name_zh:
                label_zh = f"{role_zh}｜{name_zh}"
            else:
                label_zh = role_zh or name_zh
        label_zh = _label_without_repeated_english_name(
            label_zh or "",
            role_zh=role_zh,
            name_en=name_en,
        )
        lines.append(label_zh or "")
        if name_en:
            lines.append(name_en)
        elif name_zh:
            lines.append(f"{{{{{name_zh}}}}}")
        role_en = person.get("role_en", "").strip()
        if not role_en:
            placeholder_key = role_zh or "ROLE_EN"
            role_en = f"{{{{{placeholder_key}}}}}"
        lines.append(role_en)
        org_en = person.get("org_en", "").strip()
        if org_en:
            lines.append(org_en)
        if idx < len(people) - 1:
            lines.append("")
    return lines


def apply_default_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def default_output_path(source_docx: Path, output_dir: Path) -> Path:
    stem = source_docx.stem
    if stem.endswith("_final"):
        stem = stem[: -len("_final")]
    return output_dir / f"{stem}_標題職銜_final.docx"


def resolve_template_path(template_path: Path) -> Path:
    if template_path.is_absolute() or template_path.exists():
        return template_path
    return Path(__file__).resolve().parent / template_path


def generate_meta(
    template_path: Path,
    input_path: Path,
    output_path: Path,
    meta_path: Path | None = None,
) -> None:
    data = parse_input(input_path, meta_path)
    doc = Document(str(resolve_template_path(template_path)))
    apply_default_margins(doc)

    title_placeholder = find_paragraph_by_text(doc, TITLE_PLACEHOLDER)
    replace_or_remove_paragraph_text(title_placeholder, str(data.get("title_en", "")))

    people_lines = build_people_lines(data.get("people", []))
    people_placeholder = find_paragraph_by_text(doc, PEOPLE_PLACEHOLDER)
    if people_placeholder:
        replace_multiline(people_placeholder, people_lines)

    if not people_lines:
        people_label = find_paragraph_by_text(doc, "名字職銜")
        if people_label is not None:
            remove_paragraph(people_label)
        ensure_single_blank_before_label(doc, "YT簡介")

    overview_placeholder = find_paragraph_by_text(doc, OVERVIEW_PLACEHOLDER)
    replace_or_remove_paragraph_text(
        overview_placeholder, str(data.get("overview_en", ""))
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    fix_docx_namespaces(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Render meta.docx from shared news txt data.")
    parser.add_argument(
        "--template",
        default="templates/meta_template.docx",
        help="Path to the meta DOCX template.",
    )
    parser.add_argument(
        "--source-txt",
        default="source.txt",
        help="Path to the body text input.",
    )
    parser.add_argument(
        "--meta-txt",
        default="",
        help="Optional path to a separate meta txt input.",
    )
    parser.add_argument(
        "--output",
        default="",
        help="Path to write the rendered meta DOCX.",
    )
    parser.add_argument(
        "--source-docx",
        required=True,
        help="Original source DOCX for naming the output.",
    )
    args = parser.parse_args()

    output_path = Path(args.output) if args.output else default_output_path(
        Path(args.source_docx), Path("output")
    )

    meta_path = Path(args.meta_txt) if args.meta_txt else None
    generate_meta(Path(args.template), Path(args.source_txt), output_path, meta_path=meta_path)


if __name__ == "__main__":
    main()
