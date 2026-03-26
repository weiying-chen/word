#!/usr/bin/env python3

from __future__ import annotations

import argparse
import re
import zipfile
import warnings
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from docx_utils import add_hyperlink
from generate_subs import fix_docx_namespaces, normalize_input_text, remove_paragraph


SHOT_ID_RE = re.compile(r"^\d+_\d+$")
BODY_LABEL_LINE_RE = re.compile(r"^\s*(BODY|字幕)\s*[:：]\s*$")
BODY_INLINE_LINE_RE = re.compile(r"^\s*(BODY|字幕)\s*[:：]\s*(.*)$")
SOURCE_LINK_RE = re.compile(r"^https?://\S+$")
MARKER_TEXT = "<"
BODY_PLACEHOLDER = "{{BODY}}"


def _decode_input_text(path: Path) -> tuple[str, str, bool]:
    raw = path.read_bytes()
    tried_encodings: list[str] = []

    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig"), "utf-8-sig", False

    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16"), "utf-16", False

    for encoding in ("utf-8", "big5", "cp950", "gb18030", "cp1252"):
        tried_encodings.append(encoding)
        try:
            return raw.decode(encoding), encoding, encoding != "utf-8"
        except UnicodeDecodeError:
            continue

    tried = ", ".join(tried_encodings)
    raise UnicodeError(
        f"Unable to decode input file '{path}' with supported encodings: {tried}"
    )


def parse_input(path: Path) -> dict[str, str]:
    text, encoding_used, used_fallback = _decode_input_text(path)
    if used_fallback:
        warnings.warn(
            f"Using fallback encoding '{encoding_used}' for {path}; rewriting as UTF-8.",
            stacklevel=2,
        )
        path.write_text(text, encoding="utf-8")

    lines = text.splitlines()

    for idx, raw_line in enumerate(lines):
        stripped = raw_line.strip()

        inline = BODY_INLINE_LINE_RE.match(stripped)
        if inline and not BODY_LABEL_LINE_RE.match(stripped):
            collected = [inline.group(2)] if inline.group(2) else []
            collected.extend(lines[idx + 1 :])
            return {"BODY": normalize_input_text("\n".join(collected).rstrip())}

        if BODY_LABEL_LINE_RE.match(stripped):
            return {
                "BODY": normalize_input_text("\n".join(lines[idx + 1 :]).rstrip())
            }

    return {"BODY": normalize_input_text(text.rstrip())}


def _set_line_in_paragraph(paragraph, text: str) -> None:
    paragraph.text = ""
    if not text:
        return
    if SOURCE_LINK_RE.match(text.strip()):
        link = text.strip()
        add_hyperlink(paragraph, link, link)
        return
    run = paragraph.add_run(text)
    run.font.highlight_color = (
        WD_COLOR_INDEX.BRIGHT_GREEN
        if SHOT_ID_RE.match(text.strip())
        else WD_COLOR_INDEX.WHITE
    )


def _insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    _set_line_in_paragraph(new_para, text)
    return new_para


def _add_plain_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph("")
    _set_line_in_paragraph(paragraph, text)


def _render_multiline_block(doc: Document, text: str) -> None:
    lines = text.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    for line in lines:
        _add_plain_paragraph(doc, line)


def default_output_path(source_docx: Path, output_dir: Path) -> Path:
    stem = source_docx.stem
    if not stem.endswith("_final"):
        stem = f"{stem}_final"
    return output_dir / f"{stem}.docx"


def resolve_template_path(template_path: Path) -> Path:
    if template_path.is_absolute() or template_path.exists():
        return template_path
    return Path(__file__).resolve().parent / template_path


def _marker_info(doc: Document) -> tuple[int, str]:
    for idx, paragraph in enumerate(doc.paragraphs):
        marker = paragraph.text.strip()
        if marker in {MARKER_TEXT, BODY_PLACEHOLDER}:
            return idx, marker
    raise ValueError("template.docx must contain either '<' or '{{BODY}}' marker paragraph.")


def _extract_source_header_lines(source_docx_path: Path) -> list[str]:
    with zipfile.ZipFile(source_docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    lines: list[str] = []
    for p in root.findall(".//w:body/w:p", ns):
        text_parts = [
            t.text or ""
            for t in p.findall(".//w:t", ns)
        ]
        lines.append("".join(text_parts).strip())

    marker_idx = None
    for idx, line in enumerate(lines):
        if line.strip() == MARKER_TEXT:
            marker_idx = idx
            break
    if marker_idx is None:
        return []

    lines = lines[: marker_idx + 1]
    while lines and not lines[0].strip():
        lines.pop(0)
    return lines


def _trim_existing_body(doc: Document, marker_idx: int) -> None:
    for paragraph in list(doc.paragraphs[marker_idx + 1 :]):
        remove_paragraph(paragraph)


def generate_news(
    template_path: Path,
    source_docx_path: Path,
    input_path: Path,
    output_path: Path,
) -> None:
    data = parse_input(input_path)
    generate_news_from_data(template_path, source_docx_path, data, output_path)


def generate_news_from_data(
    template_path: Path,
    source_docx_path: Path,
    data: dict[str, str],
    output_path: Path,
) -> None:
    doc = Document(str(resolve_template_path(template_path)))
    marker_idx, marker = _marker_info(doc)
    _trim_existing_body(doc, marker_idx)
    header_lines = _extract_source_header_lines(source_docx_path)

    body = data.get("BODY", "")
    content_lines: list[str] = []
    if header_lines:
        content_lines.extend(header_lines)
        if body:
            content_lines.append("")
    if body:
        content_lines.extend(body.splitlines())

    lines = content_lines
    if body:
        if marker == BODY_PLACEHOLDER:
            target = doc.paragraphs[marker_idx]
            first = lines[0] if lines else ""
            _set_line_in_paragraph(target, first)
            current = target
            for line in lines[1:]:
                current = _insert_paragraph_after(current, line)
        else:
            if lines:
                doc.add_paragraph("")
                _render_multiline_block(doc, "\n".join(lines))
    elif marker == BODY_PLACEHOLDER:
        target = doc.paragraphs[marker_idx]
        if lines:
            _set_line_in_paragraph(target, lines[0])
            current = target
            for line in lines[1:]:
                current = _insert_paragraph_after(current, line)
        else:
            _set_line_in_paragraph(target, "")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    fix_docx_namespaces(output_path)


def generate_news_from_sources(
    template_path: Path,
    source_docx_path: Path,
    source_txt_path: Path,
    output_path: Path,
) -> None:
    generate_news(template_path, source_docx_path, source_txt_path, output_path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Render a newsroom DOCX by preserving the DOCX header and replacing the body from text."
    )
    parser.add_argument(
        "--source-txt",
        default="source.txt",
        help="Path to the body text source.",
    )
    parser.add_argument(
        "--template",
        default="templates/news_template.docx",
        help="Path to the news DOCX template.",
    )
    parser.add_argument(
        "--output",
        default="",
        help="Path to write the generated DOCX.",
    )
    parser.add_argument(
        "--source-docx",
        required=True,
        help="Original source DOCX whose header is preserved.",
    )
    args = parser.parse_args()

    output_path = Path(args.output) if args.output else default_output_path(
        Path(args.source_docx), Path("output")
    )
    generate_news(
        Path(args.template),
        Path(args.source_docx),
        Path(args.source_txt),
        output_path,
    )


if __name__ == "__main__":
    main()
