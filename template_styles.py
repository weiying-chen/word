from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

from docx_utils import clear_paragraph
from style_tokens import (
    BODY_TEXT_SIZE_PT,
    REFERENCE_LINK_RGB,
    REFERENCE_TEXT_SIZE_PT,
    SECTION_LABEL_BLUE_RGB,
)



@dataclass(frozen=True)
class SectionLabelRule:
    prefix: str
    style_name: str


SECTION_LABEL_RULES = (
    SectionLabelRule("本月總翻譯時數(字幕):", "SectionLabelSmall"),
    SectionLabelRule("本月總審稿時數(字幕):", "SectionLabelSmall"),
    SectionLabelRule("其他工作:", "SectionLabelSmall"),
    SectionLabelRule("行政工作:", "SectionLabelSmall"),
)


def ensure_character_style(doc: Document, name: str, size_pt: int, color: RGBColor) -> None:
    styles = doc.styles
    if name in {style.name for style in styles}:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
    style.font.size = Pt(size_pt)
    style.font.color.rgb = color


def ensure_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.size = Pt(BODY_TEXT_SIZE_PT)

    ensure_character_style(
        doc, "SectionLabelSmall", REFERENCE_TEXT_SIZE_PT, SECTION_LABEL_BLUE_RGB
    )
    ensure_character_style(doc, "SectionLabelLarge", 12, SECTION_LABEL_BLUE_RGB)
    ensure_character_style(doc, "ReferenceLink", REFERENCE_TEXT_SIZE_PT, REFERENCE_LINK_RGB)


def apply_section_label_style(paragraph, style_name: str) -> None:
    text = paragraph.text
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.style = paragraph.part.document.styles[style_name]


def sync_review_template_styles(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    if len(table.rows) <= 13 or len(table.columns) < 4:
        return
    for row_idx in (12, 13):
        for col_idx in range(4):
            cell = table.cell(row_idx, col_idx)
            for paragraph in cell.paragraphs:
                stripped = paragraph.text.strip()
                if not stripped:
                    continue
                for rule in SECTION_LABEL_RULES:
                    if stripped.startswith(rule.prefix):
                        apply_section_label_style(paragraph, rule.style_name)
                        break


def sync_all_templates(template_paths: list[Path]) -> None:
    for path in template_paths:
        doc = Document(str(path))
        ensure_base_styles(doc)
        if path.name == "review_template.docx":
            sync_review_template_styles(doc)
        doc.save(str(path))
