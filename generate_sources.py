#!/usr/bin/env python3

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from docx_utils import add_hyperlink, apply_font_size_to_runs
from style_tokens import BODY_TEXT_SIZE_PT


HARDCODED_TIMESTAMP_LINE = "07:27-09:20 (1分53秒)"
DESCRIPTION_TIMESTAMP_RE = re.compile(r"^\s*(?P<mm>\d{1,2}):(?P<ss>\d{2})｜")
TIMECODE_ROW_RE = re.compile(
    r"^(?P<sh>\d{2}):(?P<sm>\d{2}):(?P<ss>\d{2}):(?P<sf>\d{2})\t"
    r"(?P<eh>\d{2}):(?P<em>\d{2}):(?P<es>\d{2}):(?P<ef>\d{2})\t(?P<text>.*)$"
)


def _safe_filename(text: str) -> str:
    safe = text
    for ch in '<>:"/\\|?*':
        safe = safe.replace(ch, " ")
    safe = re.sub(r"\s+", " ", safe).strip()
    return safe


def _find_subtitle_file(subtitles_dir: Path, ep_id: str) -> Path | None:
    if not ep_id:
        return None
    for candidate in sorted(subtitles_dir.glob("*.txt")):
        if ":Zone.Identifier" in candidate.name:
            continue
        if f"第{ep_id}集_ch_" in candidate.name or candidate.name.startswith(f"{ep_id}_"):
            return candidate
    return None


def _first_summary_line(description: str) -> str:
    for line in description.splitlines():
        stripped = line.strip()
        if stripped:
            return stripped
    return ""


def _format_minutes_seconds(total_seconds: int) -> str:
    minutes, seconds = divmod(max(0, total_seconds), 60)
    return f"{minutes}分{seconds}秒"


def _dynamic_timestamp_line(last_line: str, subtitle_lines: list[str]) -> str:
    match = DESCRIPTION_TIMESTAMP_RE.match(last_line.strip())
    if not match:
        return HARDCODED_TIMESTAMP_LINE
    end_total = int(match.group("mm")) * 60 + int(match.group("ss"))
    start = "00:00"
    start_total = 0
    duration = max(0, end_total - start_total)
    end = f"{int(match.group('mm')):02d}:{int(match.group('ss')):02d}"
    return f"{start}-{end} ({_format_minutes_seconds(duration)})"


def _to_seconds(hh: int, mm: int, ss: int) -> int:
    return hh * 3600 + mm * 60 + ss


def _to_ticks(hh: int, mm: int, ss: int, ff: int) -> int:
    # Keep frame precision for boundary comparisons without assuming exact frame rate.
    return (((hh * 60) + mm) * 60 + ss) * 100 + ff


def _mmss(total_seconds: int) -> str:
    minutes, seconds = divmod(max(0, total_seconds), 60)
    return f"{minutes:02d}:{seconds:02d}"


def _extract_star_range(subtitle_lines: list[str]) -> tuple[int, int] | None:
    marker_spans: list[tuple[int, int]] = []
    for line in subtitle_lines:
        if "*" not in line:
            continue
        match = TIMECODE_ROW_RE.match(line)
        if not match:
            continue
        start = _to_seconds(
            int(match.group("sh")),
            int(match.group("sm")),
            int(match.group("ss")),
        )
        end = _to_seconds(
            int(match.group("eh")),
            int(match.group("em")),
            int(match.group("es")),
        )
        marker_spans.append((start, end))
        if len(marker_spans) == 2:
            break
    if len(marker_spans) < 2:
        return None
    return marker_spans[0][0], marker_spans[1][1]


def _extract_star_range_ticks(subtitle_lines: list[str]) -> tuple[int, int] | None:
    marker_spans: list[tuple[int, int]] = []
    for line in subtitle_lines:
        if "*" not in line:
            continue
        match = TIMECODE_ROW_RE.match(line)
        if not match:
            continue
        start = _to_ticks(
            int(match.group("sh")),
            int(match.group("sm")),
            int(match.group("ss")),
            int(match.group("sf")),
        )
        end = _to_ticks(
            int(match.group("eh")),
            int(match.group("em")),
            int(match.group("es")),
            int(match.group("ef")),
        )
        marker_spans.append((start, end))
        if len(marker_spans) == 2:
            break
    if len(marker_spans) < 2:
        return None
    return marker_spans[0][0], marker_spans[1][1]


def _build_timestamp_line(last_line: str, subtitle_lines: list[str]) -> str:
    star_range = _extract_star_range(subtitle_lines)
    if star_range is not None:
        start, end = star_range
        return f"{_mmss(start)}-{_mmss(end)} ({_format_minutes_seconds(end - start)})"
    return _dynamic_timestamp_line(last_line, subtitle_lines)


def _strip_star_marker(line: str) -> str:
    return re.sub(r"\s*\*\s*$", "", line).rstrip()


def _line_span_ticks(line: str) -> tuple[int, int] | None:
    match = TIMECODE_ROW_RE.match(line)
    if not match:
        return None
    return (
        _to_ticks(
            int(match.group("sh")),
            int(match.group("sm")),
            int(match.group("ss")),
            int(match.group("sf")),
        ),
        _to_ticks(
            int(match.group("eh")),
            int(match.group("em")),
            int(match.group("es")),
            int(match.group("ef")),
        ),
    )


def _highlight_flags_for_lines(subtitle_lines: list[str]) -> list[bool]:
    star_range = _extract_star_range_ticks(subtitle_lines)
    if star_range is None:
        return [False] * len(subtitle_lines)
    range_start, range_end = star_range
    flags: list[bool] = []
    for line in subtitle_lines:
        span = _line_span_ticks(line)
        if span is None:
            flags.append(False)
            continue
        line_start, line_end = span
        # Highlight only lines fully inside the marker range, including the two marker lines.
        flags.append(line_start >= range_start and line_end <= range_end)
    return flags


def _read_text_with_fallback(path: Path) -> str:
    raw = path.read_bytes()
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")
    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16")
    for encoding in ("utf-8", "big5", "cp950", "gb18030", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def resolve_default_episodes_file(base_dir: Path = Path(".")) -> Path:
    return base_dir / "episodes.json"


def resolve_default_subtitles_dir(base_dir: Path = Path(".")) -> Path:
    return base_dir / "subtitles"


def _render_docx(
    template_path: Path,
    output_path: Path,
    title: str,
    youtube_url: str,
    summary: str,
    subtitle_lines: list[str],
    subtitle_highlights: list[bool],
    timestamp_line: str,
) -> None:
    doc = Document(str(template_path))
    if not doc.paragraphs:
        doc.add_paragraph("")
    first = doc.paragraphs[0]
    first.text = ""
    first.add_run(title)
    apply_font_size_to_runs(first, font_size_pt=BODY_TEXT_SIZE_PT)

    p_url = doc.add_paragraph("")
    add_hyperlink(p_url, youtube_url, youtube_url)
    apply_font_size_to_runs(p_url, font_size_pt=BODY_TEXT_SIZE_PT)

    p_time = doc.add_paragraph(timestamp_line)
    for run in p_time.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    apply_font_size_to_runs(p_time, font_size_pt=BODY_TEXT_SIZE_PT)

    doc.add_paragraph("")
    p_summary = doc.add_paragraph(summary)
    apply_font_size_to_runs(p_summary, font_size_pt=BODY_TEXT_SIZE_PT)

    if subtitle_lines:
        doc.add_paragraph("")
        for idx, line in enumerate(subtitle_lines):
            p_line = doc.add_paragraph(line)
            if idx < len(subtitle_highlights) and subtitle_highlights[idx]:
                for run in p_line.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            apply_font_size_to_runs(p_line, font_size_pt=BODY_TEXT_SIZE_PT)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def generate_sources(
    *,
    episodes_file: Path,
    template_path: Path,
    subtitles_dir: Path,
    output_dir: Path,
) -> dict[str, int]:
    episodes = json.loads(episodes_file.read_text(encoding="utf-8"))
    generated = 0
    skipped = 0
    errors = 0

    for item in episodes:
        ep_id = str(item.get("epId", "")).strip()
        subtitle_file = _find_subtitle_file(subtitles_dir, ep_id)
        if subtitle_file is None:
            skipped += 1
            continue

        try:
            title = str(item.get("youtubeTitle", "")).strip() or str(
                item.get("titleZh", "")
            ).strip()
            youtube_url = str(item.get("youtubeUrl", "")).strip()
            summary = _first_summary_line(str(item.get("youtubeDescription", "")))
            last_ts_line = str(item.get("descriptionLastTimestampLine", "")).strip()
            subtitle_lines = [
                line.rstrip("\n")
                for line in _read_text_with_fallback(subtitle_file).splitlines()
                if line.strip()
            ]
            timestamp_line = _build_timestamp_line(last_ts_line, subtitle_lines)
            subtitle_highlights = _highlight_flags_for_lines(subtitle_lines)
            subtitle_lines = [_strip_star_marker(line) for line in subtitle_lines]
            if not title or not youtube_url:
                errors += 1
                continue
            filename = f"{_safe_filename(title)}.docx"
            output_path = output_dir / filename
            _render_docx(
                template_path,
                output_path,
                title,
                youtube_url,
                summary,
                subtitle_lines,
                subtitle_highlights,
                timestamp_line,
            )
            generated += 1
        except Exception:
            errors += 1

    return {
        "generated": generated,
        "skipped": skipped,
        "errors": errors,
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate source DOCX files from episodes JSON and subtitle source files."
    )
    parser.add_argument(
        "--episodes-file",
        default="",
        help="Path to episodes.json. Defaults to ./episodes.json.",
    )
    parser.add_argument(
        "--template",
        default="templates/sources_template.docx",
        help="Path to sources template docx.",
    )
    parser.add_argument(
        "--subtitles-dir",
        default="",
        help="Directory containing subtitle txt files. Defaults to ./subtitles.",
    )
    parser.add_argument(
        "--output-dir",
        default="output",
        help="Directory to write generated source docx files.",
    )
    args = parser.parse_args()

    episodes_file = (
        Path(args.episodes_file)
        if args.episodes_file
        else resolve_default_episodes_file()
    )
    subtitles_dir = (
        Path(args.subtitles_dir)
        if args.subtitles_dir
        else resolve_default_subtitles_dir()
    )

    result = generate_sources(
        episodes_file=episodes_file,
        template_path=Path(args.template),
        subtitles_dir=subtitles_dir,
        output_dir=Path(args.output_dir),
    )
    print(
        f"generated={result['generated']} skipped={result['skipped']} errors={result['errors']}"
    )


if __name__ == "__main__":
    main()
