#!/usr/bin/env python3

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document

from docx_utils import (
    add_highlighted_run,
    add_hyperlink,
    apply_font_size_to_document_runs,
    apply_font_size_to_runs,
    clear_paragraph,
    ensure_blank_after_labels,
    get_default_tab_stop_inches,
    set_source_indent,
)
from prepare_posts import (
    apply_source_style,
    ensure_blank_after_reference_url,
    fetch_youtube_video_metadata,
    insert_video_section_spacing,
    insert_paragraph_after,
    normalize_empty_paragraphs,
    remove_obsolete_post_labels,
    remove_paragraph,
    replace_placeholders,
    sync_empty_paragraph_indents,
)
from style_tokens import (
    BODY_TEXT_SIZE_PT,
    REFERENCE_HIGHLIGHT_DEFAULT,
    REFERENCE_TEXT_SIZE_PT,
)


FIELD_RE = re.compile(r"^(?P<key>[A-Za-z_]+)\s*:\s*(?P<value>.*)$")
FIELD_KEYS = {"TITLE", "VIDEO_URL", "BODY", "SOURCES"}
HTTP_URL_RE = re.compile(r"^https?://\S+$", re.IGNORECASE)


def _join_content(lines: list[str]) -> str:
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines)


def parse_post_text(text: str) -> dict[str, str]:
    lines = text.lstrip("\ufeff").splitlines()
    fields: dict[str, str] = {}
    idx = 0
    while idx < len(lines):
        match = FIELD_RE.match(lines[idx].strip())
        if match is None or match.group("key").upper() not in FIELD_KEYS:
            idx += 1
            continue
        key = match.group("key").upper()
        collected = [match.group("value")] if match.group("value") else []
        idx += 1
        while idx < len(lines):
            next_match = FIELD_RE.match(lines[idx].strip())
            if next_match and next_match.group("key").upper() in FIELD_KEYS:
                break
            collected.append(lines[idx])
            idx += 1
        fields[key] = "\n".join(collected).strip()

    for key in ("TITLE", "VIDEO_URL", "BODY", "SOURCES"):
        if not fields.get(key, "").strip():
            raise ValueError(f"[error] Missing required field: {key}")

    video_url = fields["VIDEO_URL"].strip()
    if not HTTP_URL_RE.fullmatch(video_url):
        raise ValueError("[error] VIDEO_URL must be an HTTP or HTTPS URL.")

    content = fields["BODY"].splitlines()

    hashtag_indices = [
        idx for idx, line in enumerate(content) if line.strip().startswith("#")
    ]
    if len(hashtag_indices) != 2:
        raise ValueError("[error] BODY must contain exactly two hashtag lines.")
    en_hash_idx, zh_hash_idx = hashtag_indices[0], hashtag_indices[1]

    return {
        "title": fields["TITLE"].strip(),
        "video_url": video_url,
        "post_en": _join_content(content[:en_hash_idx]),
        "hashtags_en": content[en_hash_idx].strip(),
        "post_zh": _join_content(content[en_hash_idx + 1 : zh_hash_idx]),
        "hashtags_zh": content[zh_hash_idx].strip(),
        "sources": fields["SOURCES"],
    }


def parse_post_file(path: Path) -> dict[str, str]:
    return parse_post_text(path.read_text(encoding="utf-8-sig"))


def _remove_draft_scaffolding(doc: Document) -> None:
    remove_text = {
        "標題",
        "{{TITLE_LINE_1}}",
        "{{TITLE_LINE_2}}",
    }
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip() in remove_text:
            remove_paragraph(paragraph)
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        remove_paragraph(doc.paragraphs[0])


def _render_sources(doc: Document, sources: str, indent_inches: float) -> None:
    target = next(
        (
            paragraph
            for paragraph in doc.paragraphs
            if paragraph.text.strip() == "{{REF_TITLE}}"
        ),
        None,
    )
    if target is None:
        return

    lines = sources.splitlines()
    clear_paragraph(target)
    current = target
    for idx, line in enumerate(lines):
        if idx:
            current = insert_paragraph_after(current)
        stripped = line.strip()
        if not stripped:
            continue
        set_source_indent(current, indent_inches)
        if HTTP_URL_RE.fullmatch(stripped):
            add_hyperlink(current, stripped, stripped, highlight=True)
        else:
            add_highlighted_run(
                current,
                line,
                font_size_pt=REFERENCE_TEXT_SIZE_PT,
                highlight_color=REFERENCE_HIGHLIGHT_DEFAULT,
            )
        apply_source_style(current)


def _render_multiline_placeholder(doc: Document, placeholder: str, text: str) -> None:
    target = next(
        (paragraph for paragraph in doc.paragraphs if placeholder in paragraph.text),
        None,
    )
    if target is None:
        return

    paragraph_properties = deepcopy(target._p.pPr)
    lines = text.split("\n") if text else [""]
    clear_paragraph(target)
    current = target
    for idx, line in enumerate(lines):
        if idx:
            current = insert_paragraph_after(current)
            if paragraph_properties is not None:
                current._p.insert(0, deepcopy(paragraph_properties))
        if line:
            current.add_run(line)
            apply_font_size_to_runs(current, font_size_pt=BODY_TEXT_SIZE_PT)


def _resolve_template_path(path: Path) -> Path:
    if path.is_absolute() or path.exists():
        return path
    return Path(__file__).resolve().parent / path


def generate_post(
    *,
    input_path: Path,
    template_path: Path,
    output_path: Path,
) -> Path:
    post = parse_post_file(input_path)
    video_url = post["video_url"]
    video_desc_en = ""
    video_desc_zh = ""
    video_title = ""
    if video_url:
        video_title, _, video_desc_zh = fetch_youtube_video_metadata(video_url)
    video_title = video_title or post["title"]

    doc = Document(str(_resolve_template_path(template_path)))
    remove_obsolete_post_labels(doc)
    apply_font_size_to_document_runs(doc, font_size_pt=BODY_TEXT_SIZE_PT)
    _remove_draft_scaffolding(doc)
    default_tab_stop = get_default_tab_stop_inches(doc)
    mapping = {
        "{{HEADER_TITLE}}": post["title"],
        "{{HEADER_URL}}": video_url,
        "{{HASHTAGS_EN}}": post["hashtags_en"],
        "{{HASHTAGS_ZH}}": post["hashtags_zh"],
        "{{REF_URL}}": "",
        "{{REF_SUMMARY_ZH}}": "",
        "{{REF_TITLE_EN}}": "",
        "{{REF_SUMMARY_EN}}": "",
        "{{VIDEO_URL}}": video_url,
        "{{VIDEO_TITLE}}": video_title,
        "{{VIDEO_DESC_EN}}": video_desc_en,
        "{{VIDEO_DESC_ZH}}": video_desc_zh,
    }
    _render_multiline_placeholder(doc, "{{POST_EN}}", post["post_en"])
    _render_multiline_placeholder(doc, "{{POST_ZH}}", post["post_zh"])
    _render_sources(doc, post["sources"], default_tab_stop)
    replace_placeholders(doc, mapping, default_tab_stop)
    insert_video_section_spacing(
        doc,
        video_title=video_title,
        video_desc_en=video_desc_en,
        video_desc_zh=video_desc_zh,
        indent_inches=default_tab_stop,
    )
    ensure_blank_after_labels(doc, {"參考資料：", "要用的影片："})
    normalize_empty_paragraphs(doc)
    sync_empty_paragraph_indents(doc)
    ensure_blank_after_reference_url(
        doc,
        ref_url="",
        indent_inches=default_tab_stop,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a finished post DOCX from completed post text."
    )
    parser.add_argument("--input", required=True, help="Completed post TXT file.")
    parser.add_argument(
        "--template",
        default="templates/post_template.docx",
        help="Post DOCX template.",
    )
    parser.add_argument("--output", default="", help="Output DOCX path.")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")
    generate_post(
        input_path=input_path,
        template_path=Path(args.template),
        output_path=output_path,
    )


if __name__ == "__main__":
    main()
