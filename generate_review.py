#!/usr/bin/env python3

from __future__ import annotations

import argparse
import json
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from docx_utils import apply_highlight_to_runs, clear_paragraph
from style_tokens import REVIEW_NOTES_TEXT_SIZE_PT, REVIEW_TEXT_SIZE_PT


GOAL_LABEL_TEXT = "本月精進目標:"
MONTH_KEY = "MONTH"
HEADER_FONT_SIZE_PT = REVIEW_TEXT_SIZE_PT
REVIEWER_NAME = "陳威穎"


def resolve_template_path(template_path: Path) -> Path:
    if template_path.is_absolute() or template_path.exists():
        return template_path
    return Path(__file__).resolve().parent / template_path


def _format_year_month_text(value: str) -> str:
    text = value.strip()
    if not text:
        return ""
    if len(text) == 7 and text[4] == "-":
        year_text, month_text = text.split("-", 1)
        if year_text.isdigit() and month_text.isdigit():
            return f"{int(year_text)}年{int(month_text)}月"
    return text


def parse_tasks_payload(path: Path) -> list[dict]:
    raw = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(raw, list):
        raise ValueError("tasks.json must be a top-level JSON array.")
    return raw


def _first_stage(task: dict) -> dict:
    stages = task.get("stages")
    if not isinstance(stages, list):
        return {}
    for stage in reversed(stages):
        if isinstance(stage, dict):
            return stage
    return {}


def _task_stage_value(task: dict, key: str):
    stage = _first_stage(task)
    return stage.get(key)


def _task_value(task: dict, key: str):
    value = task.get(key)
    if value not in (None, ""):
        return value
    return _task_stage_value(task, key)


def _task_type(task: dict) -> str:
    return str(_task_value(task, "type") or "").strip().lower()


def _task_extensions(task: dict) -> list[dict]:
    extensions = _task_stage_value(task, "extensions")
    if not isinstance(extensions, list):
        return []
    return [item for item in extensions if isinstance(item, dict)]


def _task_descendants(task: dict) -> list[dict]:
    descendants: list[dict] = []

    children = task.get("children")
    if isinstance(children, list):
        descendants.extend(item for item in children if isinstance(item, dict))

    descendants.extend(_task_extensions(task))
    return descendants


def _parse_iso_datetime(value: object) -> datetime | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00"))
    except ValueError:
        return None


def _task_start_datetime(task: dict) -> datetime | None:
    return _parse_iso_datetime(_task_value(task, "startAt"))


def _derive_target_month(tasks: list[dict]) -> tuple[int, int] | None:
    # tasks.json is append-only in this workflow, so the last task determines
    # the review month shown in the header.
    for task in reversed(tasks):
        if not isinstance(task, dict):
            continue
        start_at = _task_start_datetime(task)
        if start_at is None:
            continue
        return (start_at.year, start_at.month)
    return None


def derive_month_from_tasks(tasks: list[dict]) -> str:
    target_month = _derive_target_month(tasks)
    if target_month is not None:
        year, month = target_month
        return f"{year}年{month}月"
    return ""


def _is_parent_subs_task(task: dict) -> bool:
    task_type = _task_type(task)
    return task_type in ("", "subs")


def _task_month_relation(task: dict, target_month: tuple[int, int]) -> str | None:
    start_at = _task_start_datetime(task)
    if start_at is None:
        return None
    task_month = (start_at.year, start_at.month)
    if task_month == target_month:
        return "current"
    if task_month < target_month:
        return "previous"
    return "future"


def _partition_parent_subs_tasks(
    tasks: list[dict],
    target_month: tuple[int, int] | None,
) -> tuple[list[dict], list[dict]]:
    if target_month is None:
        return ([], [])

    current_tasks: list[dict] = []
    previous_tasks: list[dict] = []
    for task in tasks:
        if not isinstance(task, dict) or not _is_parent_subs_task(task):
            continue
        relation = _task_month_relation(task, target_month)
        if relation == "current":
            current_tasks.append(task)
        elif relation == "previous":
            previous_tasks.append(task)
    return current_tasks, previous_tasks


def _format_month_day(iso_text: str) -> str:
    if not iso_text:
        return ""
    text = iso_text.replace("Z", "+00:00")
    dt = datetime.fromisoformat(text)
    return f"{dt.month}/{dt.day}"


def _format_work_minutes(work_minutes: int | str | None) -> str:
    if work_minutes in (None, ""):
        return ""
    minutes = int(work_minutes)
    hours, remain = divmod(minutes, 60)
    parts: list[str] = []
    if hours:
        parts.append(f"{hours}時")
    if remain:
        parts.append(f"{remain}分")
    return "".join(parts) if parts else "0分"


def _format_content_seconds(content_seconds: int | str | None) -> str:
    if content_seconds in (None, ""):
        return ""
    total = int(content_seconds)
    hours, remainder = divmod(total, 3600)
    minutes, seconds = divmod(remainder, 60)
    parts: list[str] = []
    if hours:
        parts.append(f"{hours}時")
    if minutes:
        parts.append(f"{minutes}分")
    if seconds:
        parts.append(f"{seconds}秒")
    return "".join(parts) if parts else "0秒"


def _sum_parent_content_seconds(tasks: list[dict]) -> int:
    total = 0
    for item in tasks:
        if not isinstance(item, dict):
            continue
        value = _task_value(item, "contentSeconds")
        if value not in (None, ""):
            try:
                total += int(value)
            except (TypeError, ValueError):
                pass
    return total


def _set_cell_lines(cell, lines: list[str], *, font_size_pt: int | None = None) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    if lines:
        for idx, line in enumerate(lines):
            run = paragraph.add_run(line)
            if font_size_pt is not None:
                run.font.size = Pt(font_size_pt)
            if idx < len(lines) - 1:
                run.add_break()
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)


def _extract_feedback_lines(task: dict) -> list[str]:
    raw_notes = task.get("notes")
    if isinstance(raw_notes, list):
        return [f"• {str(note).strip()}" for note in raw_notes if str(note).strip()]
    return []


def _find_regular_translation_rows(table) -> list[int]:
    rows: list[int] = []
    for idx in range(1, len(table.rows)):
        first = table.cell(idx, 0).text.strip()
        second = table.cell(idx, 1).text.strip()
        if first == "日期" and "審稿" in second:
            break
        rows.append(idx)
    return rows


def _find_block(table, start_col0: str, start_col1_contains: str, end_col0_prefix: str) -> tuple[int, int]:
    start_idx = None
    next_section_idx = None
    for idx in range(len(table.rows)):
        first = table.cell(idx, 0).text.strip()
        second = table.cell(idx, 1).text.strip()
        if start_idx is None and first == start_col0 and start_col1_contains in second:
            start_idx = idx + 1
            continue
        if start_idx is not None and first.startswith(end_col0_prefix):
            next_section_idx = idx
            break
    if start_idx is None or next_section_idx is None or start_idx >= next_section_idx:
        raise ValueError("Unable to locate expected row block in template.")
    return start_idx, next_section_idx


def _remove_row(table, row_idx: int) -> None:
    table.rows[row_idx]._tr.getparent().remove(table.rows[row_idx]._tr)


def _remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _insert_cloned_row_before(table, source_row_idx: int, before_row_idx: int) -> None:
    cloned = deepcopy(table.rows[source_row_idx]._tr)
    if before_row_idx >= len(table.rows):
        table.rows[-1]._tr.addnext(cloned)
        return
    table.rows[before_row_idx]._tr.addprevious(cloned)


def _find_regular_translation_block(table) -> tuple[int, int]:
    return _find_block(table, "日期", "字幕翻譯", "日期")


def _ensure_regular_translation_row_count(table, desired_count: int) -> list[int]:
    start_idx, next_section_idx = _find_regular_translation_block(table)
    current_count = next_section_idx - start_idx
    target = max(1, desired_count)

    while current_count < target:
        _insert_cloned_row_before(table, start_idx, next_section_idx)
        next_section_idx += 1
        current_count += 1

    while current_count > target:
        _remove_row(table, next_section_idx - 1)
        next_section_idx -= 1
        current_count -= 1

    return list(range(start_idx, start_idx + target))


def _find_temp_work_block(table) -> tuple[int, int]:
    return _find_block(table, "日期", "臨時工作", "本月工作心得")


def _ensure_temp_work_row_count(table, desired_count: int) -> list[int]:
    start_idx, next_section_idx = _find_temp_work_block(table)
    current_count = next_section_idx - start_idx
    target = max(1, desired_count)

    while current_count < target:
        _insert_cloned_row_before(table, start_idx, next_section_idx)
        next_section_idx += 1
        current_count += 1

    while current_count > target:
        _remove_row(table, next_section_idx - 1)
        next_section_idx -= 1
        current_count -= 1

    return list(range(start_idx, start_idx + target))


def _find_previous_work_block(table) -> tuple[int, int] | None:
    start_idx = None
    for idx in range(len(table.rows)):
        first = table.cell(idx, 0).text.strip()
        second = table.cell(idx, 1).text.strip()
        if first == "日期" and second == "工作項目" and idx > 0:
            previous_header = table.cell(idx - 1, 0).text.strip()
            if previous_header == "之前工作紀錄":
                start_idx = idx + 1
                break
    if start_idx is None:
        return None
    return start_idx, len(table.rows)


def _ensure_previous_work_row_count(table, desired_count: int) -> list[int]:
    block = _find_previous_work_block(table)
    if block is None:
        return []
    start_idx, next_section_idx = block
    current_count = next_section_idx - start_idx
    target = max(1, desired_count)

    while current_count < target:
        _insert_cloned_row_before(table, start_idx, next_section_idx)
        next_section_idx += 1
        current_count += 1

    while current_count > target:
        _remove_row(table, next_section_idx - 1)
        next_section_idx -= 1
        current_count -= 1

    return list(range(start_idx, start_idx + target))


def _collect_temp_posts(tasks: list[dict]) -> list[dict]:
    posts: list[dict] = []
    for task in tasks:
        for child in _task_descendants(task):
            if _task_type(child) != "posts":
                continue
            posts.append(child)
    return posts


def _count_news_children(tasks: list[dict]) -> int:
    count = 0
    for task in tasks:
        for child in _task_descendants(task):
            if _task_type(child) == "news":
                count += 1
    return count


def fill_previous_work_table(doc: Document, tasks: list[dict]) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    row_indexes = _ensure_previous_work_row_count(table, len(tasks))
    if not row_indexes:
        return

    for slot, row_idx in enumerate(row_indexes):
        task = tasks[slot] if slot < len(tasks) else None
        if not task:
            _set_cell_lines(table.cell(row_idx, 0), [])
            _set_cell_lines(table.cell(row_idx, 1), [])
            _set_cell_lines(table.cell(row_idx, 2), [])
            _set_cell_lines(table.cell(row_idx, 3), [])
            continue

        _set_cell_lines(
            table.cell(row_idx, 0),
            [_format_month_day(str(_task_value(task, "startAt") or "").strip())],
        )
        item_lines = [f"{slot + 1}.", str(task.get("name", "")).strip()]
        length_text = _format_content_seconds(_task_value(task, "contentSeconds"))
        if length_text:
            item_lines.append(f"長度:{length_text}")
        work_text = _format_work_minutes(_task_value(task, "workMinutes"))
        if work_text:
            item_lines.append(f"實際作業時間:{work_text}")
        _set_cell_lines(table.cell(row_idx, 1), [line for line in item_lines if line])
        _set_cell_lines(
            table.cell(row_idx, 2),
            _extract_feedback_lines(task),
            font_size_pt=REVIEW_NOTES_TEXT_SIZE_PT,
        )
        _set_cell_lines(table.cell(row_idx, 3), [])


def remove_subtitle_review_section(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    start_idx = None
    end_idx = None

    for idx in range(len(table.rows)):
        first = table.cell(idx, 0).text.strip()
        second = table.cell(idx, 1).text.strip()
        if start_idx is None and first == "日期" and "字幕審稿" in second:
            start_idx = idx
            continue
        if start_idx is not None and first == "日期" and "臨時工作" in second:
            end_idx = idx
            break

    if start_idx is None or end_idx is None or start_idx >= end_idx:
        return

    for idx in range(end_idx - 1, start_idx - 1, -1):
        _remove_row(table, idx)


def remove_subtitle_review_summary_block(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    heading_prefix = "本月總審稿時數(字幕):"

    for row in table.rows:
        for cell in row.cells:
            idx = 0
            while idx < len(cell.paragraphs):
                text = cell.paragraphs[idx].text.strip()
                if not text.startswith(heading_prefix):
                    idx += 1
                    continue

                _remove_paragraph(cell.paragraphs[idx])
                removed = 0
                while idx < len(cell.paragraphs) and removed < 2:
                    line = cell.paragraphs[idx].text.strip()
                    if line.startswith("中翻英:") or line.startswith("英翻中:"):
                        _remove_paragraph(cell.paragraphs[idx])
                        removed += 1
                        continue
                    break
                if idx < len(cell.paragraphs) and not cell.paragraphs[idx].text.strip():
                    _remove_paragraph(cell.paragraphs[idx])
                continue


def remove_translation_english_to_chinese_summary_line(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    heading_prefix = "本月總翻譯時數(字幕):"

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for idx, paragraph in enumerate(paragraphs):
                if not paragraph.text.strip().startswith(heading_prefix):
                    continue
                for j in range(idx + 1, len(cell.paragraphs)):
                    if cell.paragraphs[j].text.strip().startswith("英翻中:"):
                        _remove_paragraph(cell.paragraphs[j])
                        break
                break


def set_other_work_news_count_line(doc: Document, tasks: list[dict]) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    news_count = _count_news_children(tasks)
    prefix = "英文新聞:"
    replacement = f"英文新聞: {news_count}篇"

    for row in table.rows:
        for cell in row.cells:
            found_other_work_idx = None
            for idx, paragraph in enumerate(cell.paragraphs):
                stripped = paragraph.text.strip()
                if stripped.startswith(prefix):
                    paragraph.text = replacement
                    return
                if stripped.startswith("其他工作:"):
                    found_other_work_idx = idx

            # Fallback: if template has 「其他工作:」 but no 「英文新聞:」 line yet,
            # write the count into the first non-empty line after 「其他工作:」.
            if found_other_work_idx is not None:
                for idx in range(found_other_work_idx + 1, len(cell.paragraphs)):
                    if not cell.paragraphs[idx].text.strip():
                        continue
                    cell.paragraphs[idx].text = replacement
                    return


def remove_work_notes_meeting_lines(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    target_prefixes = ("同心圓會議:", "部門內部會議:")

    for row in table.rows:
        for cell in row.cells:
            idx = 0
            while idx < len(cell.paragraphs):
                text = cell.paragraphs[idx].text.strip()
                if any(text.startswith(prefix) for prefix in target_prefixes):
                    _remove_paragraph(cell.paragraphs[idx])
                    continue
                idx += 1


def normalize_translation_summary_heading_spacing(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    old = "本月總翻譯時數(字幕): (影片長度總和 非工作時數)"
    new = "本月總翻譯時數(字幕): (影片長度總和非工作時數)"

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if old not in paragraph.text:
                    continue
                if paragraph.runs:
                    for run in paragraph.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                else:
                    paragraph.text = paragraph.text.replace(old, new)


def set_translation_total_length_line(doc: Document, tasks: list[dict]) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    heading_prefix = "本月總翻譯時數(字幕):"
    total_text = f"長度:{_format_content_seconds(_sum_parent_content_seconds(tasks))}"

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for idx, paragraph in enumerate(paragraphs):
                if not paragraph.text.strip().startswith(heading_prefix):
                    continue
                for j in range(idx + 1, len(cell.paragraphs)):
                    line = cell.paragraphs[j].text.strip()
                    if line.startswith("長度:") or line.startswith("中翻英:"):
                        cell.paragraphs[j].text = total_text
                        return
                return


def fill_regular_translation_table(doc: Document, tasks: list[dict]) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    row_indexes = _ensure_regular_translation_row_count(table, len(tasks))
    for slot, row_idx in enumerate(row_indexes):
        task = tasks[slot] if slot < len(tasks) else None
        if not task:
            _set_cell_lines(table.cell(row_idx, 0), [])
            _set_cell_lines(table.cell(row_idx, 1), [])
            _set_cell_lines(table.cell(row_idx, 2), [])
            _set_cell_lines(table.cell(row_idx, 3), [])
            continue

        _set_cell_lines(
            table.cell(row_idx, 0),
            [
                _format_month_day(
                    str(_task_value(task, "startAt") or "").strip()
                )
            ],
        )
        item_lines = [
            f"{slot + 1}.",
            str(task.get("name", "")).strip(),
        ]
        length_text = _format_content_seconds(_task_value(task, "contentSeconds"))
        if length_text:
            item_lines.append(f"長度:{length_text}")
        work_text = _format_work_minutes(_task_value(task, "workMinutes"))
        if work_text:
            item_lines.append(f"實際作業時間:{work_text}")
        _set_cell_lines(table.cell(row_idx, 1), [line for line in item_lines if line])

        _set_cell_lines(
            table.cell(row_idx, 2),
            _extract_feedback_lines(task),
            font_size_pt=REVIEW_NOTES_TEXT_SIZE_PT,
        )
        _set_cell_lines(table.cell(row_idx, 3), [])


def fill_temp_work_table(doc: Document, tasks: list[dict]) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    posts = _collect_temp_posts(tasks)
    row_indexes = _ensure_temp_work_row_count(table, len(posts))

    for slot, row_idx in enumerate(row_indexes):
        post = posts[slot] if slot < len(posts) else None
        if not post:
            _set_cell_lines(table.cell(row_idx, 0), [])
            _set_cell_lines(table.cell(row_idx, 1), [])
            _set_cell_lines(table.cell(row_idx, 2), [])
            _set_cell_lines(table.cell(row_idx, 3), [])
            continue

        _set_cell_lines(
            table.cell(row_idx, 0),
            [_format_month_day(str(_task_value(post, "startAt") or "").strip())],
        )
        item_lines = [f"{slot + 1}.", str(post.get("name", "")).strip()]
        length_text = _format_content_seconds(_task_value(post, "contentSeconds"))
        if length_text:
            item_lines.append(f"長度:{length_text}")
        work_text = _format_work_minutes(_task_value(post, "workMinutes"))
        if work_text:
            item_lines.append(f"實際作業時間:{work_text}")
        _set_cell_lines(table.cell(row_idx, 1), [line for line in item_lines if line])
        _set_cell_lines(
            table.cell(row_idx, 2),
            _extract_feedback_lines(post),
            font_size_pt=REVIEW_NOTES_TEXT_SIZE_PT,
        )
        _set_cell_lines(table.cell(row_idx, 3), [])


def replace_placeholders(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
        updated = text
        for key, value in mapping.items():
            updated = updated.replace(f"{{{{{key}}}}}", value)
        if updated != text:
            paragraph.text = updated


def apply_review_highlights(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == GOAL_LABEL_TEXT:
            apply_highlight_to_runs(
                paragraph,
                highlight_color=WD_COLOR_INDEX.YELLOW,
            )


def apply_header_font_size(doc: Document, size_pt: int = HEADER_FONT_SIZE_PT) -> None:
    for idx, paragraph in enumerate(doc.paragraphs[:4]):
        if not paragraph.runs:
            paragraph.add_run(paragraph.text)
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)


def generate_review(
    template_path: Path,
    output_path: Path,
    tasks_path: Path,
) -> None:
    tasks = parse_tasks_payload(tasks_path)
    target_month = _derive_target_month(tasks)
    current_subs_tasks, previous_subs_tasks = _partition_parent_subs_tasks(
        tasks, target_month
    )
    data = {"NAME": REVIEWER_NAME, MONTH_KEY: derive_month_from_tasks(tasks)}
    doc = Document(str(resolve_template_path(template_path)))
    replace_placeholders(doc, data)
    apply_header_font_size(doc)
    apply_review_highlights(doc)
    fill_regular_translation_table(doc, current_subs_tasks)
    remove_subtitle_review_section(doc)
    remove_subtitle_review_summary_block(doc)
    remove_translation_english_to_chinese_summary_line(doc)
    remove_work_notes_meeting_lines(doc)
    normalize_translation_summary_heading_spacing(doc)
    set_translation_total_length_line(doc, current_subs_tasks)
    set_other_work_news_count_line(doc, current_subs_tasks)
    fill_temp_work_table(doc, current_subs_tasks)
    fill_previous_work_table(doc, previous_subs_tasks)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Render review DOCX from template and text input."
    )
    parser.add_argument(
        "--template",
        default="templates/review_template.docx",
        help="Path to the review DOCX template.",
    )
    parser.add_argument(
        "--output",
        default="output/review_output.docx",
        help="Path to write the rendered review DOCX.",
    )
    parser.add_argument(
        "--tasks-json",
        default="tasks.json",
        help="Path to tasks JSON array.",
    )
    args = parser.parse_args()

    generate_review(
        Path(args.template),
        Path(args.output),
        Path(args.tasks_json),
    )


if __name__ == "__main__":
    main()
