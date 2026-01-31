#!/usr/bin/env python3

from __future__ import annotations

import argparse
import re
import unicodedata
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches

from docx_utils import (
    add_highlighted_run,
    add_hyperlink,
    clear_paragraph,
    get_default_tab_stop_inches,
    set_source_indent,
)


PLACEHOLDER_KEYS = [
    "TITLE",
    "URL",
    "SUMMARY",
    "YT_TITLE_SUGGESTED",
    "TITLE_SUGGESTED",
    "INTRO",
    "THUMBNAIL",
    "TIME_RANGE",
    "BODY",
]
PLACEHOLDER_KEY_SET = set(PLACEHOLDER_KEYS)
SOURCE_URL_RE = re.compile(r"^https?://\S+")
TIME_RANGE_LINE_RE = re.compile(
    r"^\d{2}:\d{2}:\d{2}:\d{2}\t\d{2}:\d{2}:\d{2}:\d{2}\t"
)
SYMBOL_FONT_NAME = ""
HIGHLIGHT_MARKER_RE = re.compile(r"\*([^*]+)\*")
SOURCE_HIGHLIGHT_DEFAULT = WD_COLOR_INDEX.TURQUOISE
SOURCE_HIGHLIGHT_MARKED = WD_COLOR_INDEX.BRIGHT_GREEN
SOURCE_HYPERLINK_HIGHLIGHT_MARKED = "brightGreen"


def parse_input(path: Path) -> dict[str, str]:
    data: dict[str, str] = {}
    lines = path.read_text(encoding="utf-8").splitlines()
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        if ":" not in raw_line:
            idx += 1
            continue

        key, value = raw_line.split(":", 1)
        key = key.strip().upper()
        value = value.lstrip()

        if key not in PLACEHOLDER_KEY_SET:
            idx += 1
            continue

        if key in {"BODY", "INTRO", "SUMMARY"}:
            collected: list[str] = []
            if value:
                collected.append(value)
            idx += 1
            while idx < len(lines):
                next_line = lines[idx]
                if ":" in next_line:
                    next_key = next_line.split(":", 1)[0].strip().upper()
                    if next_key in PLACEHOLDER_KEY_SET:
                        break
                collected.append(next_line)
                idx += 1
            data[key] = "\n".join(collected).rstrip()
            continue

        data[key] = value
        idx += 1

    data.setdefault("BODY", "")
    data.setdefault("INTRO", "")
    data.setdefault("SUMMARY", "")

    return data


def _run_contains_symbol(text: str) -> bool:
    return any(unicodedata.category(char).startswith("S") for char in text)


def _set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def apply_symbol_font(run) -> None:
    if not SYMBOL_FONT_NAME:
        return
    if run.text and _run_contains_symbol(run.text):
        _set_run_font(run, SYMBOL_FONT_NAME)


def apply_symbol_fonts_in_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        apply_symbol_font(run)


def replace_placeholder(paragraph, placeholder: str, value: str) -> bool:
    if placeholder not in paragraph.text:
        return False

    paragraph.text = paragraph.text.replace(placeholder, value)
    apply_symbol_fonts_in_paragraph(paragraph)
    return True


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    apply_symbol_font(run)
    return new_para


def replace_body_paragraph(
    paragraph, body_text: str, source_indent_inches: float
) -> None:
    lines = body_text.splitlines() if body_text else []
    paragraph.text = ""
    if not lines:
        return
    while lines and not lines[0].strip():
        lines.pop(0)
    if not lines:
        return

    current = paragraph
    in_source_block = False

    def _add_source_runs(target, text: str) -> None:
        parts: list[tuple[str, bool]] = []
        last_idx = 0
        for match in HIGHLIGHT_MARKER_RE.finditer(text):
            if match.start() > last_idx:
                parts.append((text[last_idx : match.start()], False))
            parts.append((match.group(1), True))
            last_idx = match.end()
        if last_idx < len(text):
            parts.append((text[last_idx:], False))

        if not parts:
            run = add_highlighted_run(
                target, text, highlight_color=SOURCE_HIGHLIGHT_DEFAULT
            )
            apply_symbol_font(run)
            return

        for part_text, marked in parts:
            if not part_text:
                continue
            run = add_highlighted_run(
                target,
                part_text,
                highlight_color=(
                    SOURCE_HIGHLIGHT_MARKED if marked else SOURCE_HIGHLIGHT_DEFAULT
                ),
            )
            apply_symbol_font(run)

    def write_line(target, text: str, source_line: bool, is_url: bool) -> None:
        if not text:
            return
        if source_line:
            set_source_indent(target, source_indent_inches)
            if is_url:
                add_hyperlink(
                    target,
                    text,
                    text,
                    highlight=True,
                    highlight_color="cyan",
                )
                return
            _add_source_runs(target, text)
        else:
            run = target.add_run(text)
            apply_symbol_font(run)

    for idx, line in enumerate(lines):
        if idx > 0:
            current = insert_paragraph_after(current, "")

        if TIME_RANGE_LINE_RE.match(line):
            in_source_block = False

        cleaned_line = HIGHLIGHT_MARKER_RE.sub(r"\1", line)
        is_url = SOURCE_URL_RE.match(cleaned_line)
        if is_url:
            in_source_block = True

        write_line(
            current,
            cleaned_line if is_url else line,
            in_source_block and not TIME_RANGE_LINE_RE.match(line),
            bool(is_url),
        )


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def ensure_time_range_style(doc: Document):
    style_name = "TimeRange"
    styles = doc.styles
    if style_name in [s.name for s in styles]:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    return style_name


def ensure_hyperlink_style(doc: Document):
    style_name = "Hyperlink"
    styles = doc.styles
    if style_name in [s.name for s in styles]:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    style.font.underline = True
    return style_name


def _get_section_metrics(doc: Document):
    section = doc.sections[0]
    return {
        "page_width": section.page_width,
        "left_margin": section.left_margin,
        "right_margin": section.right_margin,
        "usable_width": section.page_width - section.left_margin - section.right_margin,
    }


def apply_default_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def _normalize_document_namespace(xml_text: str) -> str:
    match = re.search(r"<w:document[^>]*>", xml_text)
    if not match:
        return xml_text
    tag = match.group(0)
    if "ns1:Ignorable" not in tag:
        return xml_text
    new_tag = (
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
        'xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'mc:Ignorable="w14 w15 w16se wp14">'
    )
    return xml_text.replace(tag, new_tag, 1)


def fix_docx_namespaces(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        xml_text = zin.read("word/document.xml").decode("utf-8")
        fixed_text = _normalize_document_namespace(xml_text)
        if fixed_text == xml_text:
            return

        tmp_path = path.with_suffix(path.suffix + ".tmp")
        with zipfile.ZipFile(tmp_path, "w") as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    data = fixed_text.encode("utf-8")
                zout.writestr(info, data)

    tmp_path.replace(path)


def fill_template(template_path: Path, input_path: Path, output_path: Path) -> None:
    data = parse_input(input_path)
    input_base = input_path.parent
    doc = Document(str(template_path))
    apply_default_margins(doc)
    time_range_style = ensure_time_range_style(doc)
    ensure_hyperlink_style(doc)
    source_indent_inches = get_default_tab_stop_inches(doc)
    metrics = _get_section_metrics(doc)

    def _next_paragraph(paragraph: Paragraph) -> Paragraph | None:
        next_elm = paragraph._p.getnext()
        while next_elm is not None and next_elm.tag != qn("w:p"):
            next_elm = next_elm.getnext()
        if next_elm is None:
            return None
        return Paragraph(next_elm, paragraph._parent)

    def _ensure_blank_after_labels(doc_obj: Document, labels: set[str]) -> None:
        for para in list(doc_obj.paragraphs):
            if para.text.strip() not in labels:
                continue
            next_para = _next_paragraph(para)
            if next_para is None or next_para.text.strip():
                insert_paragraph_after(para, "")

    for paragraph in list(doc.paragraphs):
        if "{{INTRO}}" in paragraph.text:
            intro = data.get("INTRO", "")
            if not intro and paragraph.text.strip() == "{{INTRO}}":
                remove_paragraph(paragraph)
                continue
            replace_body_paragraph(paragraph, intro, source_indent_inches)
            continue
        if "{{SUMMARY}}" in paragraph.text:
            summary = data.get("SUMMARY", "")
            if not summary and paragraph.text.strip() == "{{SUMMARY}}":
                remove_paragraph(paragraph)
                continue
            replace_body_paragraph(paragraph, summary, source_indent_inches)
            continue
        if "{{BODY}}" in paragraph.text:
            body = data.get("BODY", "")
            if not body and paragraph.text.strip() == "{{BODY}}":
                remove_paragraph(paragraph)
                continue
            replace_body_paragraph(paragraph, body, source_indent_inches)
            continue

        for key in PLACEHOLDER_KEYS:
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                value = data.get(key, "")
                if not value and paragraph.text.strip() == placeholder:
                    remove_paragraph(paragraph)
                    break
                if key == "URL" and value:
                    clear_paragraph(paragraph)
                    add_hyperlink(paragraph, value, value)
                elif key == "THUMBNAIL" and value:
                    thumbnail_path = Path(value)
                    if not thumbnail_path.is_absolute():
                        thumbnail_path = input_base / thumbnail_path
                    if thumbnail_path.exists():
                        clear_paragraph(paragraph)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.left_indent = 0
                        paragraph.paragraph_format.right_indent = 0
                        paragraph.paragraph_format.first_line_indent = 0
                        run = paragraph.add_run()
                        run.add_picture(str(thumbnail_path), width=metrics["usable_width"])
                    else:
                        replace_placeholder(paragraph, placeholder, value)
                elif key == "TIME_RANGE" and value:
                    clear_paragraph(paragraph)
                    run = paragraph.add_run(value)
                    run.style = time_range_style
                    insert_paragraph_after(paragraph, "")
                else:
                    replace_placeholder(paragraph, placeholder, value)
                break
    _ensure_blank_after_labels(doc, {"簡介：", "簡介:", "字幕：", "字幕:"})

    doc.save(str(output_path))
    fix_docx_namespaces(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill a DOCX template from a text file.")
    parser.add_argument(
        "--template",
        default="templates/template.docx",
        help="Path to the DOCX template.",
    )
    parser.add_argument(
        "--input",
        default="input.txt",
        help="Path to the input text file.",
    )
    parser.add_argument(
        "--output",
        default="outputs/output.docx",
        help="Path to write the filled DOCX.",
    )
    args = parser.parse_args()

    fill_template(Path(args.template), Path(args.input), Path(args.output))


if __name__ == "__main__":
    main()
