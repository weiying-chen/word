#!/usr/bin/env python3

from __future__ import annotations

import argparse
from copy import deepcopy
import re
import unicodedata
import warnings
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from docx_utils import (
    add_hyperlink,
    apply_highlight_to_runs,
    apply_font_size_to_runs,
    apply_font_size_to_document_runs,
    clear_paragraph,
    ensure_blank_after_labels,
    get_default_tab_stop_inches,
    set_source_indent,
)
from style_tokens import (
    BODY_TEXT_SIZE_PT,
    REFERENCE_HIGHLIGHT_DEFAULT,
    REFERENCE_HIGHLIGHT_MARKED,
    REFERENCE_LINK_RGB,
    REFERENCE_TEXT_SIZE_PT,
)


PLACEHOLDER_KEYS = [
    "YT_TITLE_SUGGESTED",
    "TITLE_SUGGESTED",
    "INTRO",
    "THUMBNAIL",
]
PLACEHOLDER_KEY_SET = set(PLACEHOLDER_KEYS)
INPUT_KEY_SET = PLACEHOLDER_KEY_SET | {"BODY"}
GENERIC_PLACEHOLDER_RE = re.compile(r"^\{\{([A-Z_]+)\}\}$")
SOURCE_LINK_RE = re.compile(r"^https?://\S+")
SUBTITLE_LINE_RE = re.compile(
    r"^(?:[^\t]+\t)?\d{2}:\d{2}:\d{2}:\d{2}\t\d{2}:\d{2}:\d{2}:\d{2}\t"
)
PARENTHESIZED_LINE_RE = re.compile(r"^[（(].*[）)]$")
SYMBOL_FONT_NAME = "Segoe UI Symbol"
CJK_FONT_NAME = "新細明體"
CJK_MIDDLE_DOT = "\u2027"
HIGHLIGHT_MARKER_RE = re.compile(r"\*([^*]+)\*")
CPS_IGNORE_MARKER_RE = re.compile(r"\s+#\s*$")
SOURCE_HIGHLIGHT_DEFAULT = REFERENCE_HIGHLIGHT_DEFAULT
SOURCE_HIGHLIGHT_MARKED = REFERENCE_HIGHLIGHT_MARKED
SOURCE_HYPERLINK_HIGHLIGHT_MARKED = "brightGreen"
BOX_DRAWING_HORIZONTAL = "\u2500"
SPACED_HYPHEN_MINUS = " - "
DASH_LIKE_CHARS = "\u2500\u2013\u2014"
DASH_LIKE_RE = re.compile(rf"\s*[{re.escape(DASH_LIKE_CHARS)}]\s*")
SUBS_OUTPUT_SUFFIX = "_al"
SUBTITLE_LABELS = {"字幕：", "字幕:"}
SECTION_LABELS = {
    "建議YT標題：",
    "建議YT標題:",
    "建議標題：",
    "建議標題:",
    "簡介：",
    "簡介:",
    "選圖：",
    "選圖:",
    *SUBTITLE_LABELS,
}
DEFAULT_THUMBNAIL_CREDIT = "Image created with ChatGPT."
THUMBNAIL_CREDIT_MARKER = "*"


def _parse_thumbnail_value(raw_value: str) -> tuple[str, bool]:
    text = normalize_input_text(raw_value).strip()
    if not text:
        return "", False
    has_credit = False
    if text.endswith(f" {THUMBNAIL_CREDIT_MARKER}"):
        text = text[: -len(f" {THUMBNAIL_CREDIT_MARKER}")].rstrip()
        has_credit = True
    return text, has_credit


def normalize_input_text(text: str) -> str:
    if not text:
        return text
    return text.replace(BOX_DRAWING_HORIZONTAL, SPACED_HYPHEN_MINUS)


def _normalized_paragraph_text(text: str) -> str:
    return text.lstrip("\ufeff").lstrip()


def normalize_title_text(text: str) -> str:
    if not text:
        return text
    normalized = normalize_input_text(text)
    normalized = DASH_LIKE_RE.sub(SPACED_HYPHEN_MINUS, normalized)
    return re.sub(r"[ \t]+", " ", normalized).strip()


def strip_cps_ignore_marker(text: str) -> str:
    return CPS_IGNORE_MARKER_RE.sub("", text).rstrip()


def _decode_input_text(path: Path) -> tuple[str, str, bool]:
    raw = path.read_bytes()
    tried_encodings: list[str] = []

    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig"), "utf-8-sig", False

    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16"), "utf-16", False

    for encoding in ("utf-8", "big5", "cp950", "gb18030", "cp1252"):
        tried_encodings.append(encoding)
        try:
            return raw.decode(encoding), encoding, encoding != "utf-8"
        except UnicodeDecodeError:
            continue

    tried = ", ".join(tried_encodings)
    raise UnicodeError(
        f"Unable to decode input file '{path}' with supported encodings: {tried}"
    )


def parse_input(path: Path) -> dict[str, str]:
    data: dict[str, str] = {}
    thumbnail_values: list[str] = []
    thumbnail_credit_flags: list[str] = []
    text, encoding_used, used_fallback = _decode_input_text(path)
    if used_fallback:
        warnings.warn(
            f"Using fallback encoding '{encoding_used}' for {path}; rewriting as UTF-8.",
            stacklevel=2,
        )
        path.write_text(text, encoding="utf-8")

    lines = text.splitlines()
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        if ":" not in raw_line:
            idx += 1
            continue

        key, value = raw_line.split(":", 1)
        key = key.lstrip("\ufeff").strip().upper()
        value = value.lstrip()

        if key not in INPUT_KEY_SET:
            idx += 1
            continue

        if key in {"INTRO", "BODY"}:
            collected: list[str] = []
            if value:
                collected.append(value)
            idx += 1
            while idx < len(lines):
                next_line = lines[idx]
                if ":" in next_line:
                    next_key = next_line.split(":", 1)[0].strip().upper()
                    if next_key in INPUT_KEY_SET:
                        break
                collected.append(next_line)
                idx += 1
            data[key] = normalize_input_text("\n".join(collected).rstrip())
            continue

        if key in {"TITLE_SUGGESTED", "YT_TITLE_SUGGESTED"}:
            data[key] = normalize_title_text(value)
        else:
            data[key] = normalize_input_text(value)
            if key == "THUMBNAIL" and value.strip():
                thumbnail_path, has_credit = _parse_thumbnail_value(value)
                if thumbnail_path:
                    thumbnail_values.append(thumbnail_path)
                    thumbnail_credit_flags.append("1" if has_credit else "0")
        idx += 1

    data.setdefault("INTRO", "")
    data.setdefault("BODY", "")
    if thumbnail_values:
        data["__THUMBNAIL_VALUES__"] = "\n".join(thumbnail_values)
        data["__THUMBNAIL_CREDIT_FLAGS__"] = "\n".join(thumbnail_credit_flags)

    return data


def _validate_thumbnail_paths(data: dict[str, str], input_base: Path) -> None:
    raw_values = data.get("__THUMBNAIL_VALUES__", "")
    if not raw_values and data.get("THUMBNAIL", "").strip():
        raw_values = data.get("THUMBNAIL", "")
    if not raw_values:
        return
    missing: list[Path] = []
    for raw_value in raw_values.splitlines():
        value = raw_value.strip()
        if not value:
            continue
        thumbnail_path = Path(value)
        if not thumbnail_path.is_absolute():
            thumbnail_path = input_base / thumbnail_path
        if not thumbnail_path.is_file():
            missing.append(thumbnail_path)
    if missing:
        raise FileNotFoundError(
            "\n".join(f"THUMBNAIL file not found: {path}" for path in missing)
        )


def _thumbnail_paths_from_data(data: dict[str, str], input_base: Path) -> list[Path]:
    raw_values = data.get("__THUMBNAIL_VALUES__", "")
    if not raw_values and data.get("THUMBNAIL", "").strip():
        raw_values = data.get("THUMBNAIL", "")

    paths: list[Path] = []
    for raw_value in raw_values.splitlines():
        value = raw_value.strip()
        if not value:
            continue
        thumbnail_path = Path(value)
        if not thumbnail_path.is_absolute():
            thumbnail_path = input_base / thumbnail_path
        paths.append(thumbnail_path)
    return paths


def _thumbnail_credit_flags_from_data(data: dict[str, str], count: int) -> list[bool]:
    raw_flags = data.get("__THUMBNAIL_CREDIT_FLAGS__", "")
    if not raw_flags:
        return [False] * count
    flags = [line.strip() == "1" for line in raw_flags.splitlines()]
    if len(flags) < count:
        flags.extend([False] * (count - len(flags)))
    return flags[:count]


def _run_contains_symbol(text: str) -> bool:
    return any(unicodedata.category(char) == "So" for char in text)


def _set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def apply_symbol_font(run) -> None:
    if not SYMBOL_FONT_NAME:
        return
    if run.text and _run_contains_symbol(run.text):
        _set_run_font(run, SYMBOL_FONT_NAME)


def apply_cjk_middle_dot_font(run) -> None:
    if not CJK_FONT_NAME:
        return
    if run.text and CJK_MIDDLE_DOT in run.text:
        _set_run_font(run, CJK_FONT_NAME)


def apply_symbol_fonts_in_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        apply_symbol_font(run)
        apply_cjk_middle_dot_font(run)


def replace_placeholder(paragraph, placeholder: str, value: str) -> bool:
    if placeholder not in paragraph.text:
        return False

    paragraph.text = paragraph.text.replace(placeholder, value)
    apply_font_size_to_runs(paragraph, font_size_pt=BODY_TEXT_SIZE_PT)
    apply_symbol_fonts_in_paragraph(paragraph)
    return True


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        apply_symbol_font(run)
        apply_cjk_middle_dot_font(run)
    return new_para


def _split_marked_parts(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    last_idx = 0

    for match in HIGHLIGHT_MARKER_RE.finditer(text):
        if match.start() > last_idx:
            parts.append((text[last_idx : match.start()], False))
        parts.append((match.group(1), True))
        last_idx = match.end()

    if last_idx < len(text):
        parts.append((text[last_idx:], False))

    if not parts:
        parts.append((text, False))

    return parts


def _split_symbol_chunks(text: str) -> list[tuple[str, bool]]:
    if not text:
        return []

    parts: list[tuple[str, bool]] = []
    start = 0
    current_is_symbol = unicodedata.category(text[0]) == "So"

    for idx in range(1, len(text)):
        is_symbol = unicodedata.category(text[idx]) == "So"
        if is_symbol != current_is_symbol:
            parts.append((text[start:idx], current_is_symbol))
            start = idx
            current_is_symbol = is_symbol
    parts.append((text[start:], current_is_symbol))
    return parts


def _add_text_runs(
    paragraph,
    text: str,
    *,
    run_style: str | None = None,
    highlight_color=None,
    font_size_pt: int | None = BODY_TEXT_SIZE_PT,
) -> None:
    for chunk, is_symbol in _split_symbol_chunks(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if run_style:
            run.style = run_style
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
        if highlight_color is not None:
            run.font.highlight_color = highlight_color
        if is_symbol:
            apply_symbol_font(run)
        else:
            apply_cjk_middle_dot_font(run)


def _add_marked_runs(
    paragraph,
    text: str,
    *,
    default_highlight=None,
    marked_highlight=None,
    run_style: str | None = None,
    apply_default_size: bool = True,
) -> None:
    for part_text, marked in _split_marked_parts(text):
        if not part_text:
            continue

        highlight_color = marked_highlight if marked else default_highlight
        _add_text_runs(
            paragraph,
            part_text,
            run_style=run_style,
            highlight_color=highlight_color,
            font_size_pt=(
                REFERENCE_TEXT_SIZE_PT if apply_default_size else BODY_TEXT_SIZE_PT
            ),
        )


def replace_body_paragraph(
    paragraph,
    body_text: str,
    source_indent_inches: float,
    timing_style: str | None = None,
    marked_highlight=None,
) -> None:
    lines = body_text.splitlines() if body_text else []
    clear_paragraph(paragraph)
    if not lines:
        return
    while lines and not lines[0].strip():
        lines.pop(0)
    if not lines:
        return

    current = paragraph
    in_source_block = False
    previous_was_subtitle_line = False
    in_parenthesized_super_block = False

    def _add_source_runs(target, text: str) -> None:
        _add_marked_runs(
            target,
            text,
            default_highlight=SOURCE_HIGHLIGHT_DEFAULT,
            marked_highlight=SOURCE_HIGHLIGHT_MARKED,
        )

    def write_line(target, text: str, source_line: bool, is_url: bool) -> None:
        if not text:
            return
        if source_line:
            set_source_indent(target, source_indent_inches)
            if is_url:
                add_hyperlink(
                    target,
                    text,
                    text,
                    highlight=True,
                    highlight_color="cyan",
                )
                return
            _add_source_runs(target, text)
        else:
            if timing_style or marked_highlight is not None:
                _add_marked_runs(
                    target,
                    text,
                    marked_highlight=marked_highlight,
                    run_style=timing_style,
                    apply_default_size=False,
                )
            else:
                _add_text_runs(target, text)

    for idx, line in enumerate(lines):
        if idx > 0:
            current = insert_paragraph_after(current, "")

        normalized_line = _normalized_paragraph_text(line)
        if not normalized_line.strip():
            in_source_block = False
            previous_was_subtitle_line = False
            in_parenthesized_super_block = False
            continue

        is_subtitle_line = bool(SUBTITLE_LINE_RE.match(normalized_line))
        if is_subtitle_line:
            in_source_block = False
            in_parenthesized_super_block = False

        stripped_line = line.strip()
        is_parenthesized_line = bool(PARENTHESIZED_LINE_RE.match(stripped_line))
        next_stripped_line = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
        subtitle_has_super_block = is_subtitle_line and bool(
            PARENTHESIZED_LINE_RE.match(next_stripped_line)
        )
        if is_parenthesized_line and (previous_was_subtitle_line or in_parenthesized_super_block):
            in_parenthesized_super_block = True
        elif not is_parenthesized_line:
            in_parenthesized_super_block = False

        line_without_cps_marker = strip_cps_ignore_marker(line)
        cleaned_line = HIGHLIGHT_MARKER_RE.sub(r"\1", line_without_cps_marker)
        is_link = SOURCE_LINK_RE.match(cleaned_line)
        if is_link:
            in_source_block = True

        write_line(
            current,
            cleaned_line if is_link else line_without_cps_marker,
            in_source_block and not is_subtitle_line,
            bool(is_link),
        )

        if subtitle_has_super_block or in_parenthesized_super_block:
            apply_highlight_to_runs(current, highlight_color=WD_COLOR_INDEX.YELLOW)

        previous_was_subtitle_line = is_subtitle_line


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def ensure_annotation_style(doc: Document):
    style_name = "Annotation"
    styles = doc.styles
    if style_name in [s.name for s in styles]:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    style.font.color.rgb = REFERENCE_LINK_RGB
    return style_name


def ensure_hyperlink_style(doc: Document):
    style_name = "Hyperlink"
    styles = doc.styles
    if style_name in [s.name for s in styles]:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    style.font.color.rgb = REFERENCE_LINK_RGB
    style.font.underline = True
    return style_name


def _get_section_metrics(doc: Document):
    section = doc.sections[0]
    return {
        "page_width": section.page_width,
        "left_margin": section.left_margin,
        "right_margin": section.right_margin,
        "usable_width": section.page_width - section.left_margin - section.right_margin,
    }


def apply_default_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def _normalize_document_namespace(xml_text: str) -> str:
    match = re.search(r"<w:document[^>]*>", xml_text)
    if not match:
        return xml_text
    tag = match.group(0)
    if "ns1:Ignorable" not in tag:
        return xml_text
    new_tag = (
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
        'xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'mc:Ignorable="w14 w15 w16se wp14">'
    )
    return xml_text.replace(tag, new_tag, 1)


def fix_docx_namespaces(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        xml_text = zin.read("word/document.xml").decode("utf-8")
        fixed_text = _normalize_document_namespace(xml_text)
        if fixed_text == xml_text:
            return

        tmp_path = path.with_suffix(path.suffix + ".tmp")
        with zipfile.ZipFile(tmp_path, "w") as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    data = fixed_text.encode("utf-8")
                zout.writestr(info, data)

    tmp_path.replace(path)


def with_subs_output_suffix(path: Path) -> Path:
    if path.stem.endswith(SUBS_OUTPUT_SUFFIX):
        return path
    return path.with_name(f"{path.stem}{SUBS_OUTPUT_SUFFIX}{path.suffix}")


def _extract_source_paragraphs(
    source_docx_path: Path,
) -> tuple[Document, list[Paragraph], list[Paragraph]]:
    source_doc = Document(str(source_docx_path))
    paragraphs = list(source_doc.paragraphs)

    for idx, paragraph in enumerate(paragraphs):
        if SUBTITLE_LINE_RE.match(_normalized_paragraph_text(paragraph.text)):
            return source_doc, paragraphs[:idx], paragraphs[idx:]

    raise ValueError("source.docx must contain at least one subtitle timestamp paragraph.")


def _remap_relationship_ids(source_part, target_part, paragraph_element) -> None:
    for element in paragraph_element.iter():
        relation_id = element.get(qn("r:id"))
        if not relation_id:
            continue
        relation = source_part.rels[relation_id]
        if relation.reltype != RT.HYPERLINK:
            continue
        new_relation_id = target_part.relate_to(
            relation.target_ref,
            RT.HYPERLINK,
            is_external=True,
        )
        element.set(qn("r:id"), new_relation_id)


def _remap_hyperlink_run_styles(target_paragraph: Paragraph, paragraph_element) -> None:
    try:
        hyperlink_style_id = target_paragraph.part.document.styles["Hyperlink"].style_id
    except KeyError:
        return

    def set_run_style(run_element) -> None:
        run_properties = run_element.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            run_element.insert(0, run_properties)
        run_style = run_properties.find(qn("w:rStyle"))
        if run_style is None:
            run_style = OxmlElement("w:rStyle")
            run_properties.insert(0, run_style)
        run_style.set(qn("w:val"), hyperlink_style_id)

    for hyperlink in paragraph_element.iter(qn("w:hyperlink")):
        for run in hyperlink.iter(qn("w:r")):
            set_run_style(run)

    # Also handle field-code hyperlinks in the form:
    # fldChar(begin) + instrText(HYPERLINK ...) + fldChar(separate) + display runs + fldChar(end)
    in_field = False
    in_result = False
    hyperlink_field = False
    instruction_text = ""
    for child in paragraph_element:
        if child.tag != qn("w:r"):
            continue

        fld_char = child.find(qn("w:fldChar"))
        if fld_char is not None:
            fld_char_type = fld_char.get(qn("w:fldCharType"))
            if fld_char_type == "begin":
                in_field = True
                in_result = False
                hyperlink_field = False
                instruction_text = ""
            elif fld_char_type == "separate" and in_field:
                in_result = True
                hyperlink_field = "HYPERLINK" in instruction_text.upper()
            elif fld_char_type == "end":
                in_field = False
                in_result = False
                hyperlink_field = False
                instruction_text = ""
            continue

        if not in_field:
            continue

        if in_result:
            if hyperlink_field:
                set_run_style(child)
            continue

        for instr_text in child.findall(qn("w:instrText")):
            instruction_text += instr_text.text or ""


def _style_ids_with_bold(source_part) -> set[str]:
    bold_style_ids: set[str] = set()
    styles_element = source_part.document.styles.element

    for style in styles_element.iter(qn("w:style")):
        style_id = style.get(qn("w:styleId"))
        if not style_id:
            continue
        run_properties = style.find(qn("w:rPr"))
        if run_properties is None:
            continue
        if run_properties.find(qn("w:b")) is not None or run_properties.find(
            qn("w:bCs")
        ) is not None:
            bold_style_ids.add(style_id)

    return bold_style_ids


def _preserve_bold_runs_from_source_styles(source_paragraph: Paragraph, paragraph_element) -> None:
    bold_style_ids = _style_ids_with_bold(source_paragraph.part)
    if not bold_style_ids:
        return

    source_runs = list(source_paragraph._p.iter(qn("w:r")))
    cloned_runs = list(paragraph_element.iter(qn("w:r")))

    for source_run, cloned_run in zip(source_runs, cloned_runs):
        source_run_properties = source_run.find(qn("w:rPr"))
        if source_run_properties is None:
            continue

        has_direct_bold = (
            source_run_properties.find(qn("w:b")) is not None
            or source_run_properties.find(qn("w:bCs")) is not None
        )
        if has_direct_bold:
            continue

        run_style = source_run_properties.find(qn("w:rStyle"))
        if run_style is None:
            continue
        style_id = run_style.get(qn("w:val"))
        if style_id not in bold_style_ids:
            continue

        cloned_run_properties = cloned_run.find(qn("w:rPr"))
        if cloned_run_properties is None:
            cloned_run_properties = OxmlElement("w:rPr")
            cloned_run.insert(0, cloned_run_properties)

        if cloned_run_properties.find(qn("w:b")) is None:
            cloned_run_properties.append(OxmlElement("w:b"))
        if cloned_run_properties.find(qn("w:bCs")) is None:
            cloned_run_properties.append(OxmlElement("w:bCs"))


def _clone_paragraph_before(source_paragraph: Paragraph, target_paragraph: Paragraph) -> None:
    cloned = deepcopy(source_paragraph._p)
    _preserve_bold_runs_from_source_styles(source_paragraph, cloned)
    _remap_relationship_ids(source_paragraph.part, target_paragraph.part, cloned)
    _remap_hyperlink_run_styles(target_paragraph, cloned)
    target_paragraph._p.addprevious(cloned)


def _clone_paragraph_after(source_paragraph: Paragraph, target_paragraph: Paragraph) -> Paragraph:
    cloned = deepcopy(source_paragraph._p)
    _preserve_bold_runs_from_source_styles(source_paragraph, cloned)
    _remap_relationship_ids(source_paragraph.part, target_paragraph.part, cloned)
    _remap_hyperlink_run_styles(target_paragraph, cloned)
    target_paragraph._p.addnext(cloned)
    return Paragraph(cloned, target_paragraph._parent)


def _find_subtitle_target_paragraph(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() not in SUBTITLE_LABELS:
            continue
        next_elm = paragraph._p.getnext()
        while next_elm is not None and next_elm.tag != qn("w:p"):
            next_elm = next_elm.getnext()
        if next_elm is None:
            return insert_paragraph_after(paragraph, "")
        return Paragraph(next_elm, paragraph._parent)
    return None


def generate_subs(
    template_path: Path,
    source_docx_path: Path,
    input_path: Path,
    output_path: Path,
) -> None:
    data = parse_input(input_path)
    input_base = input_path.parent
    _validate_thumbnail_paths(data, input_base)
    _, source_header, source_body = _extract_source_paragraphs(source_docx_path)
    doc = Document(str(template_path))
    apply_default_margins(doc)
    annotation_style = ensure_annotation_style(doc)
    ensure_hyperlink_style(doc)
    source_indent_inches = get_default_tab_stop_inches(doc)
    metrics = _get_section_metrics(doc)

    if doc.paragraphs and source_header:
        anchor = doc.paragraphs[0]
        for paragraph in source_header:
            _clone_paragraph_before(paragraph, anchor)

    # Normalize document text to the shared body size before specialized
    # source/annotation replacements apply 10pt where needed.
    apply_font_size_to_document_runs(doc, font_size_pt=BODY_TEXT_SIZE_PT)

    for paragraph in list(doc.paragraphs):
        if "{{INTRO}}" in paragraph.text:
            intro = data.get("INTRO", "")
            if not intro and paragraph.text.strip() == "{{INTRO}}":
                remove_paragraph(paragraph)
                continue
            replace_body_paragraph(paragraph, intro, source_indent_inches)
            continue
        if "{{BODY}}" in paragraph.text:
            body = data.get("BODY", "")
            if not body and paragraph.text.strip() == "{{BODY}}":
                remove_paragraph(paragraph)
                continue
            replace_body_paragraph(paragraph, body, source_indent_inches)
            continue
        stripped = paragraph.text.strip()
        match = GENERIC_PLACEHOLDER_RE.fullmatch(stripped)
        if match and match.group(1) not in (INPUT_KEY_SET | {"INTRO", "BODY"}):
            remove_paragraph(paragraph)
            continue

        for key in PLACEHOLDER_KEYS:
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                value = data.get(key, "")
                if not value and paragraph.text.strip() == placeholder:
                    remove_paragraph(paragraph)
                    break
                if key == "THUMBNAIL" and value:
                    thumbnail_paths = _thumbnail_paths_from_data(data, input_base)
                    if not thumbnail_paths:
                        parsed_path, _ = _parse_thumbnail_value(value)
                        thumbnail_paths = [Path(parsed_path or value)]
                        if not thumbnail_paths[0].is_absolute():
                            thumbnail_paths[0] = input_base / thumbnail_paths[0]
                    credit_flags = _thumbnail_credit_flags_from_data(
                        data, len(thumbnail_paths)
                    )

                    clear_paragraph(paragraph)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.left_indent = 0
                    paragraph.paragraph_format.right_indent = 0
                    paragraph.paragraph_format.first_line_indent = 0
                    run = paragraph.add_run()
                    run.add_picture(str(thumbnail_paths[0]), width=metrics["usable_width"])

                    current = paragraph
                    if credit_flags and credit_flags[0]:
                        credit_paragraph = insert_paragraph_after(current, "")
                        credit_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        credit_paragraph.paragraph_format.left_indent = 0
                        credit_paragraph.paragraph_format.right_indent = 0
                        credit_paragraph.paragraph_format.first_line_indent = 0
                        _add_text_runs(
                            credit_paragraph,
                            DEFAULT_THUMBNAIL_CREDIT,
                            run_style=annotation_style,
                        )
                        current = credit_paragraph

                    for idx_thumb, thumbnail_path in enumerate(
                        thumbnail_paths[1:], start=1
                    ):
                        current = insert_paragraph_after(current, "")
                        image_paragraph = insert_paragraph_after(current, "")
                        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        image_paragraph.paragraph_format.left_indent = 0
                        image_paragraph.paragraph_format.right_indent = 0
                        image_paragraph.paragraph_format.first_line_indent = 0
                        image_run = image_paragraph.add_run()
                        image_run.add_picture(
                            str(thumbnail_path), width=metrics["usable_width"]
                        )
                        current = image_paragraph

                        if credit_flags[idx_thumb]:
                            credit_paragraph = insert_paragraph_after(current, "")
                            credit_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            credit_paragraph.paragraph_format.left_indent = 0
                            credit_paragraph.paragraph_format.right_indent = 0
                            credit_paragraph.paragraph_format.first_line_indent = 0
                            _add_text_runs(
                                credit_paragraph,
                                DEFAULT_THUMBNAIL_CREDIT,
                                run_style=annotation_style,
                            )
                            current = credit_paragraph
                else:
                    replace_placeholder(paragraph, placeholder, value)
                break
    ensure_blank_after_labels(doc, SECTION_LABELS)
    subtitle_target = _find_subtitle_target_paragraph(doc)
    body_text = data.get("BODY", "").strip()
    if body_text and subtitle_target is not None:
        # Keep one blank line between "字幕：" and the first timestamp line.
        clear_paragraph(subtitle_target)
        body_paragraph = insert_paragraph_after(subtitle_target, "")
        replace_body_paragraph(body_paragraph, body_text, source_indent_inches)
    elif subtitle_target is not None:
        current = subtitle_target
        for paragraph in source_body:
            current = _clone_paragraph_after(paragraph, current)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            continue
        apply_font_size_to_runs(paragraph, font_size_pt=BODY_TEXT_SIZE_PT)

    doc.save(str(output_path))
    fix_docx_namespaces(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill a DOCX template from a text file.")
    parser.add_argument(
        "--template",
        default="templates/subs_template.docx",
        help="Path to the DOCX template.",
    )
    parser.add_argument(
        "--source-docx",
        required=True,
        help="Original source DOCX whose header and subtitles are preserved.",
    )
    parser.add_argument(
        "--input",
        default="input.txt",
        help="Path to the input text file.",
    )
    parser.add_argument(
        "--output",
        default="output/output.docx",
        help="Path to write the filled DOCX.",
    )
    args = parser.parse_args()

    try:
        generate_subs(
            Path(args.template),
            Path(args.source_docx),
            Path(args.input),
            with_subs_output_suffix(Path(args.output)),
        )
    except FileNotFoundError as exc:
        lines = [line.strip() for line in str(exc).splitlines() if line.strip()]
        if not lines:
            raise SystemExit("[error] File not found")
        raise SystemExit("\n".join(f"[error] {line}" for line in lines))


if __name__ == "__main__":
    main()
